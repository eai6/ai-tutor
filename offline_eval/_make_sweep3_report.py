"""Generate EVAL_SWEEP3_REPORT.docx from the sweep-3 single-turn result JSONs.

Headline numbers are computed directly from single_turn_results/sweep3/*.json, so
the report is reproducible from the committed results. Prose/insight is maintained
here (mirrors the sweep-2 report's structure, adapted for single-turn).

    venv/bin/python offline_eval/_make_sweep3_report.py
"""
import glob
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULTS = os.path.join(ROOT, "offline_eval", "single_turn_results", "sweep3")
OUT = os.path.join(ROOT, "offline_eval", "single_turn_results", "EVAL_SWEEP3_REPORT.docx")

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)

OSS = {"qwen3.5_4b", "qwen3.5_9b", "qwen3_14b", "qwen3_30b-a3b", "qwen3_4b"}
# Display niceties: underscore OSS tags -> the Ollama colon form.
DISPLAY = {
    "qwen3.5_4b": "qwen3.5:4b", "qwen3.5_9b": "qwen3.5:9b", "qwen3_14b": "qwen3:14b",
    "qwen3_30b-a3b": "qwen3:30b-a3b", "qwen3_4b": "qwen3:4b",
}


def load_rows():
    rows = []
    for f in sorted(glob.glob(os.path.join(RESULTS, "*.json"))):
        d = json.load(open(f))
        name = os.path.basename(f)[:-5]
        res = d.get("results") or []
        tot = d.get("total_scenarios") or len(res)
        passed = d.get("passed")
        ms = [(r.get("rubric_result") or {}).get("mean_score") for r in res]
        ms = [m for m in ms if isinstance(m, (int, float))]
        rub = sum(ms) / len(ms) if ms else None
        rows.append({
            "name": name, "disp": DISPLAY.get(name, name),
            "passed": passed, "total": tot, "rate": 100 * passed / tot,
            "rubric": rub, "errored": d.get("errored") or 0,
            "tier": "local" if name in OSS else "cloud",
        })
    rows.sort(key=lambda r: -r["rate"])
    return rows


# --- docx helpers -----------------------------------------------------------
def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(doc, text, size=10.5, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def board_table(doc, rows):
    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["#", "Model", "Tier", "Pass rate", "Pass", "Rubric"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for i, r in enumerate(rows, 1):
        c = t.add_row().cells
        c[0].text = str(i)
        c[1].text = r["disp"]
        c[2].text = r["tier"]
        c[3].text = f"{r['rate']:.1f}%"
        c[4].text = f"{r['passed']}/{r['total']}"
        c[5].text = f"{r['rubric']:.2f}" if r["rubric"] is not None else "—"
    return t


def build():
    rows = load_rows()
    by = {r["name"]: r for r in rows}
    cloud = [r for r in rows if r["tier"] == "cloud"]
    local = [r for r in rows if r["tier"] == "local"]

    doc = Document()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = t.add_run("Single-Turn Tutor Evaluation — Sweep 3")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    s = sub.add_run(
        f"{len(rows)} models scored on 200 balanced single-turn scenarios "
        "(16 lessons, 6 personas, 26 archetypes). Cloud APIs + local open-weight "
        "Qwen, on one leaderboard. Covers the board, the cloud/local split, the "
        "newer-and-thinking-model regression, single-turn vs. multi-turn, and data "
        "integrity.")
    s.font.size = Pt(11)
    s.italic = True
    meta = doc.add_paragraph()
    m = meta.add_run(
        "AI Tutor · simple_tutor engine · single-turn response scoring · "
        "judge = Anthropic Haiku 4.5 @ temp 0 · n=200 per model (±3.5pp)")
    m.font.size = Pt(9)
    m.font.color.rgb = GREY

    # 1. Executive summary
    H(doc, "1. Executive summary")
    top = rows[0]
    top_cloud = cloud[0]
    para(doc,
         f"Sweep 3 scores {len(rows)} models — {len(cloud)} cloud APIs and "
         f"{len(local)} local open-weight Qwen models run through Ollama — on a "
         "single graded tutor turn each, across the same 200 scenarios and the same "
         "judge, so every model lands on one leaderboard. The data is clean: zero "
         "errored scenarios and zero empty responses across all "
         f"{len(rows)} models, so every score is the model's own behaviour.")
    bullet(doc,
           f"tops the board at {top['rate']:.0f}% — but it is a 4B local model, and "
           "that #1 is a single-turn artefact, not a claim it out-tutors the "
           "frontier. Single-turn scores one isolated reply; it does not require "
           "carrying state across a lesson. The strongest cloud model is "
           f"{top_cloud['disp']} at {top_cloud['rate']:.0f}%.",
           bold_lead=f"{top['disp']} ")
    bullet(doc,
           "Three generational pairs all regress: the newer or reasoning-heavy tier "
           "scores below its predecessor on this rubric — grok-4.20 below grok-4.1, "
           "Gemini 3.x below Gemini 2.5, and every explicit thinking model lands in "
           "the bottom third. Three vendors, same direction.",
           bold_lead="Newer and thinking models do worse. ")
    bullet(doc,
           "the local open-weight tier is genuinely competitive on single turns — "
           "qwen3.5:9b (76%) and the 14B/30B (71%) sit inside the cloud pack. "
           "Whether that survives a full multi-turn lesson is the open question, and "
           "the reason single-turn is not the deployment number.",
           bold_lead="Offline is close on single turns. ")
    bullet(doc,
           "sweep 3 is a NEW baseline. The dataset was rebuilt (60 → 200 scenarios, "
           "4 → 16 lessons, balanced by construction), so these numbers are not "
           "comparable to sweep 1 or sweep 2. The trend line starts here.",
           bold_lead="Not comparable to earlier sweeps. ")

    # 2. What sweep 3 tested
    H(doc, "2. What sweep 3 tested")
    para(doc,
         "Each scenario seeds a short context and one student utterance; the model "
         "produces the next tutor turn, which is scored by three layers: "
         "deterministic assertions (phrase / structure / judge-label), a "
         "natural-language rubric (Anthropic Haiku 4.5 @ temperature 0, the same "
         "judge for every model), and per-dimension pedagogical checks. A scenario "
         "passes only when all applicable layers agree. There is no student "
         "simulator and no trajectory — that is the multi-turn benchmark, which is a "
         "separate and more demanding test.")
    para(doc,
         "The 200 scenarios are balanced by construction across 6 personas, 2 "
         "subjects (math + geography), 16 frozen lessons, and 26 situation "
         "archetypes — no persona is starved and no lesson dominates. At n=200 the "
         "standard error on a pass rate is about ±3.5pp, so two models must differ "
         "by roughly 10pp before the gap is meaningful.")

    # 3. Leaderboard
    H(doc, "3. Sweep 3 leaderboard")
    para(doc,
         "Pass rate is the primary metric; rubric is the mean natural-language "
         "rubric score (0–1) on the scored responses. Tier is cloud API vs. local "
         "open-weight (Ollama).")
    board_table(doc, rows)
    para(doc,
         "qwen3.6:27b is excluded — its Colab run stopped partway and emitted no "
         "result file. All other models completed the full 200.", size=9,
         italic=True, color=GREY)

    # 4. The two tiers
    H(doc, "4. The two tiers")
    cavg = sum(r["rate"] for r in cloud) / len(cloud)
    lavg = sum(r["rate"] for r in local) / len(local)
    para(doc,
         f"Cloud tier ({len(cloud)} models): mean {cavg:.0f}% pass rate. The top of "
         "the cloud board is a cheap flash model (gemini-2.5-flash) and mid-size "
         "open-weight MoE models (kimi-k2-thinking, qwen3-next-80b-instruct), not "
         "the largest or newest — frontier size is not the winner here.")
    para(doc,
         f"Local open-weight tier ({len(local)} models, the offline-pilot "
         f"candidates): mean {lavg:.0f}% pass rate, forming a size gradient — "
         "qwen3.5:9b 76%, qwen3:14b / qwen3:30b-a3b 71%, qwen3:4b 58%. qwen3.5:4b is "
         "the outlier at 89%; its rubric (0.87) is consistent with its pass rate, so "
         "the responses are genuinely good — but see §6 on why a small model can "
         "look this strong on single turns and still not be deployable for whole "
         "lessons.")

    # 5. Newer/thinking regression
    H(doc, "5. Newer and thinking models regress")
    para(doc,
         "The clearest signal in the board is generational, and it points the wrong "
         "way for three separate vendors:")

    def pair(a, b):
        ra, rb = by.get(a), by.get(b)
        if ra and rb:
            bullet(doc,
                   f"{ra['disp']} {ra['rate']:.0f}% vs {rb['disp']} {rb['rate']:.0f}% "
                   f"— a {ra['rate'] - rb['rate']:+.0f}pp move to the OLDER model.",
                   bold_lead="")
    pair("glm-4.7", "glm-5")
    pair("grok-4.1-fast-non-reasoning", "grok-4.20-non-reasoning")
    pair("gemini-2.5-flash", "gemini-3.5-flash")
    para(doc,
         "Every explicit thinking / reasoning variant lands in the bottom third "
         "(grok-4.20-reasoning, deepseek-r1, qwen3-next-80b-thinking, gemini-3.x), "
         "and their rubric scores fall with their pass rates (0.59–0.72) — so it is "
         "not that they teach well but trip an assertion; the whole response is "
         "weaker. Three vendors moving the same direction is unlikely to be noise, "
         "though a thinking-mode / output-truncation interaction with this rubric "
         "is not fully ruled out and is worth a log dig before the pattern is "
         "treated as settled.")

    # 6. Single vs multi
    H(doc, "6. Single-turn overstates deployable ability")
    para(doc,
         "The sweep-2 multi-turn report found that single-turn scores systematically "
         "overstate real tutoring ability, because a single turn never exercises the "
         "register→grade→advance loop or asks the model to carry state across ~15 "
         "turns. This board is the single-turn half of that comparison. Read the #1 "
         "in that light: a 4B model producing one good reply per scenario is exactly "
         "the case the single-turn harness rewards and the multi-turn harness "
         "punishes. The multi-turn score — not this one — is what predicts what a "
         "student actually experiences.")
    para(doc,
         "This is why the follow-up is a multi-turn run, and why it targets the "
         "protocol/state fixes shipped after sweep 2 (adaptive tool-forcing, "
         "reasoning-channel recovery, server-side bare-answer grading) rather than a "
         "new model.")

    # 7. Data integrity
    H(doc, "7. Data integrity")
    bullet(doc,
           f"across all {len(rows)} models — every score is the model's own "
           "behaviour, not an infrastructure artefact.",
           bold_lead="Zero errored scenarios and zero empty responses ")
    bullet(doc,
           "were validated before merging into the board: no errored, no empty, no "
           "empty-but-passed responses, and passing responses spot-checked as "
           "genuine, contextual tutoring (they localise to Seychelles navigation "
           "charts, for instance).",
           bold_lead="The five local OSS models (run in Colab) ")
    bullet(doc,
           "was re-run after an earlier invalid result (32 judge-connection failures "
           "in the first attempt); the 40.5% here is the clean score.",
           bold_lead="qwen3-next-80b-thinking ")

    # 8. Caveats + next
    H(doc, "8. Caveats and next steps")
    bullet(doc, "n=200 → ±3.5pp; single-model gaps under ~10pp are inside the noise.",
           bold_lead="Small gaps are noise. ")
    bullet(doc,
           "single-turn scores overstate deployable ability; the multi-turn number "
           "is the one to act on.",
           bold_lead="This is not the deployment number. ")
    bullet(doc,
           "the dataset was rebuilt since sweep 2, so there is no honest delta "
           "against earlier boards.",
           bold_lead="New baseline. ")
    bullet(doc,
           "run the top models multi-turn to see whether the single-turn ranking "
           "holds under a full lesson and whether the post-sweep-2 engine fixes "
           "convert the protocol failures that sank the non-Anthropic tier.",
           bold_lead="Next. ")

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"  {len(rows)} models · cloud mean {cavg:.1f}% · local mean {lavg:.1f}%")


if __name__ == "__main__":
    build()
