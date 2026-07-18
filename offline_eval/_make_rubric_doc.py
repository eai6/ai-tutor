"""Generate MULTI_TURN_RUBRIC_EXPLAINED.docx — the multi-turn eval rubric + rationale.

The rubric ITEM TEXT is pulled straight from the source of truth
(evals/gen_multi_turn.py, evals/matrix.py) so the document can never drift from
what the harness actually scores. Prose/rationale is maintained here.

    venv/bin/python offline_eval/_make_rubric_doc.py
"""
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)  # so `evals` imports resolve when run as a script

# Rubric source of truth (pure imports — no Django needed).
from evals.gen_multi_turn import BEA_SESSION_RUBRIC, SHAPE_RUBRIC, PERSONA_NOTE
from evals.matrix import MULTI_TURN
OUT = os.path.join(ROOT, "offline_eval", "MULTI_TURN_RUBRIC_EXPLAINED.docx")

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)

# --- docx helpers (mirrors offline_eval/_make_report.py) -------------------
def h(doc, text, level):
    return doc.add_heading(text, level=level)


def para(doc, text, *, italic=False, size=None, color=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    for i, seg in enumerate(text.split("**")):
        if seg == "":
            continue
        r = p.add_run(seg)
        r.bold = (i % 2 == 1)
        r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p


def bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    for i, seg in enumerate(text.split("**")):
        if seg == "":
            continue
        r = p.add_run(seg)
        r.bold = (i % 2 == 1)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    for i, seg in enumerate(text.split("**")):
        if seg == "":
            continue
        r = p.add_run(seg)
        r.bold = (i % 2 == 1)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, htext in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    return t


# --- content: per-BEA-item rationale (keyed by list index) ------------------
BEA_RATIONALE = [
    "The core tutoring signal. Catches both failure directions: a false ACCEPT "
    "(affirming a wrong answer) and a false HEDGE (doubting or re-opening a "
    "correct one). This is exactly the item the math MCQ-fabrication bug drove "
    "down — a tutor that mis-graded correct answers scored low here.",
    "Localization, not just detection. A tutor may notice a slip but respond with "
    "a generic 'not quite'; that teaches little. The item rewards naming the "
    "specific step, number, or misconception. Empirically the hardest item for "
    "models to satisfy — they under-follow it even when instructed to.",
    "Hint-vs-reveal discipline. Under repeated pressure (or its own confusion) a "
    "tutor tends to hand over the answer. Scaffolding hints are fine; naming the "
    "canonical value or letter for an unsolved question is not.",
    "Guidance must be both correct and calibrated to the student's level — a right "
    "explanation pitched too high, or a confident wrong one, both fail here.",
    "Momentum. Every turn should hand the conversational floor back with a concrete "
    "next action. Passive endings ('let me know when you're ready') stall the "
    "session and waste the turn budget — a common pacing failure.",
    "A trajectory-level property a single-turn eval cannot see: the tutor must not "
    "contradict its own earlier turns, assume unestablished facts, or ignore what "
    "the student just said. Self-contradiction confuses the student more than the "
    "original mistake.",
    "Affective quality — warm and encouraging without being condescending, honest "
    "without being harsh. Tone drives whether a struggling student stays engaged.",
    "Human-likeness — natural conversational teaching rather than robotic, "
    "templated replies or filler openers ('Great question!', 'Let me think about "
    "this carefully...') that pad the turn without teaching.",
]

BEA_SHORT = [
    "Mistake recognition + honest affirmation",
    "Specific error localization",
    "Never reveal the final answer",
    "Correct, calibrated guidance",
    "Actionability (clear next step)",
    "Logical consistency across the session",
    "Warm, encouraging tone",
    "Natural, non-robotic teaching",
]

# =====================================================================
doc = Document()
# normal style
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# --- title -----------------------------------------------------------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("The Multi-Turn Tutoring Evaluation Rubric")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = ACCENT
sub = doc.add_paragraph()
r = sub.add_run("What each session is scored on, how the score is computed, and "
                "the design rationale behind every choice")
r.font.size = Pt(11)
r.font.color.rgb = GREY
meta = doc.add_paragraph()
r = meta.add_run("AI Tutor · offline evaluation harness · rubric text pulled "
                 "verbatim from evals/gen_multi_turn.py and evals/matrix.py")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

# --- 1. What this measures ------------------------------------------
h(doc, "1. What a multi-turn scenario evaluates", 1)
para(doc, "A multi-turn scenario runs a whole simulated tutoring session and "
     "asks: **did the tutor teach well across the entire arc?** Unlike a "
     "single-turn check, it tests trajectory-level behaviour — coherence over "
     "20–40 turns, adapting to a student who refuses or disengages, not spiralling "
     "on repeated errors, reaching a clean end. A scenario is defined by exactly "
     "four things: a **persona**, a **lesson**, a **turn budget**, and the "
     "**trajectory properties the tutor must satisfy** (its rubric).")
para(doc, "Three models are in the loop each session:")
bullet(doc, "**Tutor** — the model under test (e.g. gemini-2.5-flash, "
       "kimi-k2-thinking, qwen3-next-80b). This is the only thing that changes "
       "between runs of the same scenario.")
bullet(doc, "**Student simulator** — Anthropic Haiku, role-playing one of six "
       "personas, producing the student turns live (no scripted history).")
bullet(doc, "**Rubric judge** — Anthropic Sonnet 4.6 at temperature 0, which "
       "reads the finished transcript and scores it against the rubric.")

# --- 2. Pass criterion ----------------------------------------------
h(doc, "2. How a session passes — two independent gates", 1)
para(doc, "A session is marked PASS only when **both** gates clear "
     "(`passed = deterministic_passed AND rubric_passed`). Splitting the check "
     "keeps the expensive LLM judge for the things that genuinely need judgement, "
     "and lets cheap, reproducible code catch mechanical failures.")
para(doc, "**Gate A — deterministic assertions** (regex / structural, no LLM). "
     "Every scenario carries:")
bullet(doc, "**expected_reason** — the session must end for an allowed reason "
       "(completed, exit_ticket, or max_turns), not a crash.")
bullet(doc, "**max_turn_count** — a quality target: reach the exit ticket within "
       "N turns (N = the scenario's turn budget). A session can complete but still "
       "miss this efficiency target.")
bullet(doc, "**no_repeated_tutor_phrase_within_window** — the tutor may not recycle "
       "the same phrase inside a sliding window (window/threshold scale with the "
       "budget: 4-turn/2× under a tight cap, 5-turn/3× otherwise).")
bullet(doc, "**no_tool_syntax_in_any_turn** — no raw tool-call or thinking syntax "
       "may leak into a student-visible turn.")
para(doc, "**Gate B — the rubric.** The judge scores each rubric item 0.0–1.0; the "
     "session passes Gate B when the **mean over applicable items ≥ 0.60** "
     "(`pass_threshold`, set per scenario). The 0.60 bar is the calibrated "
     "session-level threshold; single-turn evals use a stricter 0.70.")

# --- 3. The judge ----------------------------------------------------
h(doc, "3. The judge — Anthropic Sonnet 4.6 @ temperature 0", 1)
para(doc, "**Temperature 0** so scores are reproducible run-to-run. **Sonnet, not "
     "Haiku**: judging a 20–40 turn transcript is a harder, longer-context task "
     "than the single-turn eval. A Haiku-vs-Sonnet A/B (2026-07-06) found Haiku "
     "**too lenient** — it passed 7 of 8 sessions where Sonnet passed 2, agreement "
     "with Sonnet was Cohen's kappa ≈ 0.09 (near chance), and Haiku missed a real "
     "tutor self-contradiction that Sonnet caught. Single-turn scenarios keep Haiku "
     "for continuity with the frozen earlier board; multi-turn upgrades to Sonnet "
     "where judge quality matters most. A scenario can still override the judge via "
     "its own rubric_judge block.")

# --- 4. Scoring mechanics -------------------------------------------
h(doc, "4. Scoring mechanics", 1)
para(doc, "**Continuous 0.0–1.0, not binary.** Session-level behaviour is a matter "
     "of degree — an item that held on every relevant turn scores 1.0, one that "
     "lapsed on some scores lower. The judge is anchored: 1.0 fully satisfies, 0.7 "
     "minor issue, 0.5 partial with clear gaps, 0.3 significant issue, 0.0 fails.")
para(doc, "**Whole-session, not turn-by-turn.** For each item the judge considers "
     "every tutor turn where the item is relevant and scores how well it held "
     "overall — an 'if the student made a mistake' item applies only on turns where "
     "the student actually erred; a 'reveals the answer' item applies only while a "
     "question is unresolved.")
para(doc, "**Conditional items and the n/a rule.** Many items are conditional "
     "('If the student made a mistake…'). When the condition never holds anywhere "
     "in the session, the judge returns \"n/a\" and the item is **excluded from the "
     "mean** — it neither helps nor hurts. Without this, a spurious 0.0 on an "
     "inapplicable item would unfairly sink an otherwise-passing session. If every "
     "item is n/a, the layer is a vacuous pass.")
para(doc, "**End-reason note.** The judge is told how the session ended. The "
     "simulator deliberately ends the tutoring phase at the exit-ticket hand-off — "
     "the graded quiz and any remediation happen afterward and are not in the "
     "transcript. Without the note the judge misreads that clean hand-off as an "
     "abrupt cutoff and penalises 'terminates cleanly' or scores exit-ticket items "
     "as failures instead of n/a. The note names the ending as successful "
     "(exit_ticket / completed) or unclean (max_turns / deadlock / error).")

# --- 5. Universal rubric --------------------------------------------
h(doc, "5. The universal rubric — 8 BEA-aligned session items", 1)
para(doc, "These eight items are appended, byte-identical, to **every** multi-turn "
     "scenario, so the judge applies one consistent standard across all 200 "
     "scenarios. They are aligned with the BEA (Building Educational Applications) "
     "pedagogical-ability dimensions — anchoring on a community standard for tutor "
     "evaluation rather than an ad-hoc list makes the results defensible and "
     "comparable. Each is scored across the whole session.")
for i, item in enumerate(BEA_SESSION_RUBRIC):
    para(doc, f"**{i+1}. {BEA_SHORT[i]}.**", after=2)
    para(doc, item, italic=True, size=9.5, after=2)
    para(doc, BEA_RATIONALE[i], size=9.5, color=GREY, after=8)

# --- 6. Shape-specific rubric ---------------------------------------
h(doc, "6. The shape-specific rubric — 12 trajectory templates", 1)
para(doc, "Each scenario is also one of twelve **shapes** — a recurring "
     "full-session trajectory that stresses a specific tutoring skill. A shape "
     "contributes three scenario-specific rubric items on top of the eight "
     "universal ones (≈11 items per scenario). Shape items make the eval "
     "**diagnostic**: they show which behaviour failed, not just that the session "
     "did. Scenarios are generated from these twelve templates rather than "
     "hand-authored — twelve reviewable shapes beat 200 near-identical YAMLs that "
     "would drift. Each shape is eligible only for the personas it makes sense with "
     "(a speedrun needs a fast, correct student; an error_cascade needs an "
     "error-prone one).")

summary_by_key = {s.key: s for s in MULTI_TURN}
table(doc,
      ["Shape", "What it stresses", "Eligible personas", "Turn budgets"],
      [[s.key,
        s.summary,
        ", ".join(s.eligible) if len(s.eligible) < 6 else "all six",
        ", ".join(str(b) for b in s.budgets)]
       for s in MULTI_TURN],
      widths=[1.4, 2.5, 1.9, 0.9])

para(doc, "", after=2)
para(doc, "The three items each shape adds:", after=4)
for key, spec in SHAPE_RUBRIC.items():
    para(doc, f"**{key}**", after=1)
    for item in spec["rubric"]:
        bullet(doc, item)

# --- 7. Personas -----------------------------------------------------
h(doc, "7. The six student personas", 1)
para(doc, "The student simulator role-plays one persona per session. Personas are "
     "what make the rubric items bite — 'change strategy after a refusal' only has "
     "teeth against a probe-resistant student; 'don't give up on the student' only "
     "against a non-responder.")
for name, note in PERSONA_NOTE.items():
    # strip the leading article the source uses ("A STRUGGLER (...)" -> body)
    para(doc, f"**{name}** — {note}", after=3)

# --- 8. Composition --------------------------------------------------
h(doc, "8. How a scenario's rubric is assembled", 1)
para(doc, "For a given (persona, lesson, shape, turn budget), the harness builds "
     "the scenario's rubric mechanically:")
numbered(doc, "Start with the **three shape-specific items** for that trajectory.")
numbered(doc, "Append the **eight universal BEA items**, verbatim.")
numbered(doc, "Attach the deterministic **assertions** (expected_reason, "
         "max_turn_count = the budget, phrase-repetition window, no tool syntax).")
numbered(doc, "Set **pass_threshold = 0.60**.")
para(doc, "The judge scores all ~11 items across the transcript; the applicable-item "
     "mean must clear 0.60 (Gate B) and all assertions must hold (Gate A).")

# --- 9. Design rationale --------------------------------------------
h(doc, "9. Design rationale — the why, consolidated", 1)
for head, body in [
    ("Trajectory-level, LLM-judged",
     "Multi-turn tutoring quality lives in the arc, not any single turn. A tutor "
     "can be locally fine yet globally incoherent — contradict itself, drift off "
     "the objective, loop the same probe. Only reading the whole transcript catches "
     "those. And the questions that matter ('did the explanation address the "
     "misconception?', 'was the tone right for a struggler?') need a model in the "
     "loop; regexes can't answer them."),
    ("Two layers, cheap-first",
     "Mechanical failures (tool syntax leaking, phrase repetition, running out of "
     "turns) are caught by deterministic assertions — reproducible and free. The "
     "LLM judge is reserved for the pedagogical dimensions that actually need "
     "judgement, keeping cost and variance down."),
    ("A community-standard universal core",
     "The eight BEA-aligned items give every session one fixed, defensible yardstick "
     "held byte-identical across all 200 scenarios, so scores are comparable and not "
     "an author's ad-hoc taste."),
    ("Shape-specific items for diagnosis",
     "The universal core says whether tutoring was good; the shape items say which "
     "specific skill (scaffolding under help-requests, changing strategy under "
     "refusal, crediting a self-correction, not over-probing a correct student) "
     "held or broke. That turns a pass/fail into an actionable signal."),
    ("Continuous scoring + n/a exclusion",
     "Degrees, not booleans, preserve the 'mostly good with one slip' signal. "
     "Excluding inapplicable conditional items keeps the mean honest — an item can "
     "neither inflate a score by being vacuously satisfied nor sink it by being "
     "scored 0 when it never applied."),
    ("Sonnet @ temp 0, with an end-reason note",
     "Temperature 0 for reproducibility; Sonnet because the A/B proved Haiku too "
     "lenient on long transcripts; the end-reason note so a deliberate exit-ticket "
     "hand-off is not mistaken for a crash."),
    ("Generated from shapes, persona-masked",
     "Twelve reviewable templates instead of hundreds of drifting hand-written "
     "files; a persona-eligibility mask so every generated scenario is internally "
     "coherent."),
]:
    para(doc, f"**{head}.** {body}", after=6)

# --- 10. Limitations -------------------------------------------------
h(doc, "10. Known limitations", 1)
bullet(doc, "**Small-sample noise.** A representative run is n=20 per model "
       "(Wilson 95% CI ≈ ±11pp overall, wider per subject at n=10). Treat the "
       "aggregate pass rate as noise-dominated; per-item rubric means and "
       "end-reason distributions are the cleaner signals.")
bullet(doc, "**Two-gate coupling.** A session can earn a high rubric mean yet fail "
       "overall on a strict deterministic assertion (e.g. reaching the exit ticket "
       "in 17 turns against a 12-turn max_turn_count target). Read the two gates "
       "separately when diagnosing.")
bullet(doc, "**Single-judge dependency.** One judge model defines 'good'. Temp 0 + "
       "the Sonnet upgrade mitigate leniency and variance, but the judge is still a "
       "single point; its own blind spots become the eval's.")
bullet(doc, "**Simulated students.** The persona sim (Haiku) approximates real "
       "learners; it is internally consistent and reproducible, but a simulated "
       "student is not a real one.")

doc.save(OUT)
print("wrote", OUT)
print("BEA items:", len(BEA_SESSION_RUBRIC), "| shapes:", len(SHAPE_RUBRIC),
      "| personas:", len(PERSONA_NOTE))
