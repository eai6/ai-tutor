"""Generate Offline_Model_Evaluation_Report.docx from leaderboard_combined.csv.

The combined 44-model table is built directly from `leaderboard_combined.csv`
(itself produced from `aggregate.py`), so the report's headline numbers are
reproducible from the committed result JSONs. Prose/insights are maintained here.

    venv/bin/python offline_eval/_make_report.py
"""
import csv
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CSV = os.path.join(ROOT, "offline_eval", "leaderboard_combined.csv")
OUT = os.path.join(ROOT, "offline_eval", "Offline_Model_Evaluation_Report.docx")

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)

# --- auxiliary metadata not in the CSV -------------------------------------
MODE = {
    "grok-4.1-fast-reasoning": "reasoning", "grok-4.1-fast-non-reasoning": "non-reasoning",
    "grok-4.20-reasoning": "reasoning", "grok-4.20-non-reasoning": "non-reasoning",
    "qwen3-coder-480b": "instruct", "qwen3-next-80b-instruct": "instruct",
    "qwen3-235b-instruct": "instruct", "qwen3-next-80b-thinking": "thinking",
    "glm-4.7": "—", "glm-5": "—",
    "deepseek-v3.2": "non-thinking", "deepseek-v3.1": "non-thinking",
    "deepseek-r1": "thinking", "kimi-k2-thinking": "thinking",
}
OSS = {  # csv_name: (params, device tier)
    "qwen2.5_14b": ("14B", "GPU laptop / server"), "mistral-nemo_12b": ("12B", "GPU laptop / server"),
    "qwen2.5_7b": ("7B", "GPU laptop"), "qwen2.5_3b": ("3B", "phone / tablet"),
    "glm4_9b": ("9B", "GPU laptop"), "granite3.1-dense_8b": ("8B", "GPU laptop"),
    "llama3.1_8b": ("8B", "GPU laptop"), "mistral_7b": ("7B", "GPU laptop"),
    "llama3.2_3b": ("3B", "phone / tablet"), "qwen2.5_1.5b": ("1.5B", "phone / tablet"),
    "hermes3_8b": ("8B", "GPU laptop"), "llama3-groq-tool-use_8b": ("8B", "GPU laptop"),
    "command-r7b": ("7B", "GPU laptop"), "granite3.1-dense_2b": ("2B", "phone / tablet"),
    "granite3.1-moe_3b": ("3B", "phone / tablet"), "aya-expanse_8b": ("8B", "GPU laptop"),
    "llama3.2_1b": ("1B", "phone / tablet"), "nemotron-mini": ("~4B", "phone / tablet"),
    "qwen2.5_0.5b": ("0.5B", "phone / tablet"), "hermes3_3b": ("3B", "phone / tablet"),
    "falcon3_10b": ("10B", "—"), "gemma2_2b": ("2B", "phone / tablet"), "phi4": ("14B", "—"),
}
MARK = {"gemini-2.5-pro": "*", "deepseek-r1": "†", "qwen3-next-80b-thinking": "†",
        "falcon3_10b": "‡", "gemma2_2b": "‡", "phi4": "‡"}


def disp(name):
    """Display name: OSS result-file names use '_' for the ollama ':' tag."""
    if name in OSS:
        return name.replace("_", ":", 1) if "_" in name else name
    return name


def load():
    with open(CSV) as f:
        return list(csv.DictReader(f))


# --- docx helpers ----------------------------------------------------------
def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def para(doc, text, *, italic=False, size=None, color=None, after=6):
    """Add a paragraph, rendering **bold** spans."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if seg == "":
            continue
        r = p.add_run(seg)
        r.bold = (i % 2 == 1)
        r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if seg == "":
            continue
        r = p.add_run(seg)
        r.bold = (i % 2 == 1)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    parts = text.split("**")
    for i, seg in enumerate(parts):
        if seg == "":
            continue
        r = p.add_run(seg)
        r.bold = (i % 2 == 1)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, htext in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    return t


def mono(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p


# --- build -----------------------------------------------------------------
rows = load()
doc = Document()
# base font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# Title block
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("Offline & Cloud Model Evaluation for the AI Tutor")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = ACCENT
sub = doc.add_paragraph()
r = sub.add_run("Ranking 44 tutor models — local open-source, cloud open-weight (Vertex Model Garden), and proprietary — on the real production tutoring engine")
r.font.size = Pt(11.5)
r.italic = True
r.font.color.rgb = GREY
meta = doc.add_paragraph()
r = meta.add_run("Pixel Design Labs LLC — AI Tutor   ·   44 models scored   ·   June 2026")
r.font.size = Pt(9.5)
r.font.color.rgb = GREY

# 1. Executive Summary
h(doc, "1. Executive Summary", 1)
para(doc, "We evaluated **44 candidate tutor models** on the AI Tutor's real production "
     "tutoring engine, using one identical harness of 60 single-turn lesson scenarios. "
     "Every model drives the tutor end-to-end (posing questions, grading answers, and "
     "advancing through a lesson via tool-calls); scoring is held constant on Anthropic "
     "models and is cross-family, so no model grades itself. **0 scenarios errored** "
     "across the final dataset.")
para(doc, "The frontier ceiling is **Claude Opus 4.7 at 90%**. The strongest non-Anthropic "
     "model is **Grok 4.1-fast (reasoning) at 72%** — served as open-weight Model-as-a-Service "
     "through Google Vertex AI Model Garden — which beats every Gemini variant and trails "
     "only the three Claude models. A dense band of **cloud-hosted open-weight** models follows: "
     "**Qwen3-Coder 480B (68%)**, **GLM-4.7 (67%)**, **Qwen3-Next-80B-Instruct (65%)**, and "
     "**Qwen3-235B (63%)** — all of which beat the best locally-run open model.")
para(doc, "For **on-device** deployment in low-connectivity schools, the picks are "
     "**qwen2.5:3b (45%)** for phone/tablet and **qwen2.5:7b (52%, rubric 0.71)** for a "
     "laptop or school server. These are stock models with no tutor-specific tuning — a "
     "baseline, and the gap to the ceiling is the target for prompt/fine-tuning.")

# 2. Objective & Context
h(doc, "2. Objective & Context", 1)
para(doc, "The pilot runs on a hosted Anthropic model. For **data-residency** and **offline use** "
     "in low-connectivity schools (Mozambique, Tanzania), we need to know which models can drive "
     "the tutor well — both models small enough to run **on-device** and large open-weight models "
     "we could host ourselves or reach cheaply via cloud Model-as-a-Service (MaaS). This study "
     "ranks them on how well they run our **real tutoring engine**, not a toy benchmark.")
para(doc, "Three serving modes are compared head-to-head on the same harness:")
bullet(doc, "**Proprietary** — hosted vendor APIs (Anthropic Claude, Google Gemini).")
bullet(doc, "**Cloud MaaS** — open-weight models served pay-per-token through Google Vertex AI "
       "Model Garden's OpenAI-compatible endpoint (DeepSeek, Kimi, Qwen3, Grok, GLM). No GPU "
       "endpoints are deployed; billing is per-token only.")
bullet(doc, "**Local OSS** — open-source models run locally via Ollama on an 8 GB CPU laptop "
       "(≤9B) or a free Google Colab T4 GPU (7–14B), the realistic offline deployment targets.")

# 3. Methodology
h(doc, "3. Methodology", 1)
bullet(doc, "**Tutor under test:** each model drives the production `simple_tutor` engine, "
       "controlling pedagogy through tool-calls (pose question → grade answer → advance).")
bullet(doc, "**Scorers held constant on Anthropic** (our trusted reference): the rubric **judge** "
       "(claude-haiku-4-5) and the **student-simulator** are the same models used in production, so "
       "every tutor model is graded on an identical, high-quality yardstick.")
bullet(doc, "**Cross-family grading:** the pass/fail grader excludes the tutor's own vendor, so a "
       "model never grades itself.")
bullet(doc, "**Test set:** 60 single-turn lesson scenarios (math + reading, multiple student "
       "personas), each scored **pass/fail** plus a **0–1 rubric** quality score.")
bullet(doc, "**Provider abstraction:** the engine routes any provider through a pluggable client "
       "layer (`apps/llm/client.py`). Cloud MaaS models reach Vertex through "
       "`VertexModelGardenClient` (OpenAI-compatible endpoint, Application Default Credentials, "
       "per-region routing); local models go through the Ollama client.")
bullet(doc, "**Hardware:** ≤9B local models on an 8 GB CPU laptop; 7–14B local models on a free "
       "Colab T4 GPU; cloud + MaaS models via API. Same scenarios throughout → directly comparable.")

# 4. Combined leaderboard
h(doc, "4. Combined Leaderboard — all 44 models", 1)
para(doc, "Single ranking by pass rate. **Pass** = cross-family pass rate (primary metric). "
     "**Rubric** = mean 0–1 teaching-quality score (Anthropic-Haiku judge). **Top failure** = the "
     "single biggest bottleneck category for that model, with its count. Every model ran the "
     "identical 60-scenario harness with **0 errors**.")
comb_rows = []
for r0 in rows:
    name = r0["model"]
    mark = MARK.get(name, "")
    comb_rows.append([
        r0["rank"], disp(name), f'{r0["type"]} · {r0["vendor"]}',
        f'{r0["pass_rate_pct"]}%{mark}', r0["rubric"],
        f'{r0["top_failure"].replace("_"," ")} ({r0["top_failure_count"]})',
    ])
table(doc, ["#", "Model", "Type · Vendor", "Pass", "Rubric", "Top failure"], comb_rows,
      widths=[0.35, 1.95, 1.7, 0.7, 0.6, 1.4])
para(doc, "", after=2)
para(doc, "* gemini-2.5-pro suspect (Pro scores below Flash — likely a thinking-mode harness "
     "interaction, not true capability).", italic=True, size=9, color=GREY)
para(doc, "† deepseek-r1 / qwen3-next-thinking: reasoning-token truncation before the tool call "
     "(task-fit artifact; their instruct siblings score far higher), not raw capability.",
     italic=True, size=9, color=GREY)
para(doc, "‡ falcon3 / gemma2 / phi4: tool-protocol failures (no/garbled tool-calls), not "
     "teaching failures.", italic=True, size=9, color=GREY)

# 5. Results by tier
h(doc, "5. Results by Tier", 1)

h(doc, "5.1 Proprietary benchmark ceiling", 2)
prop = [r0 for r0 in rows if r0["type"] == "Proprietary"]
table(doc, ["Model", "Vendor", "Pass", "Rubric"],
      [[disp(r0["model"]), r0["vendor"], f'{r0["pass_rate_pct"]}%{MARK.get(r0["model"],"")}', r0["rubric"]] for r0 in prop],
      widths=[2.2, 1.4, 0.9, 0.9])
para(doc, "Claude Opus 4.7 (90%) sets the ceiling; Haiku 4.5 (82%) edging Sonnet 4.6 (78%) is "
     "plausible for this strong recent small model. **Gemini Pro anomaly:** the Pro models score "
     "below the Flash models (2.5-flash 65% > 3.1-pro 58% > 2.5-pro 43%), which is backwards — the "
     "Pro models returned empty responses in id-probing (thinking-mode), so the Pro rows are "
     "**suspect pending a diagnostic** rather than a true capability ranking.")

h(doc, "5.2 Cloud open-weight via Vertex Model Garden (MaaS)", 2)
para(doc, "Open-weight models served pay-per-token through Vertex AI Model Garden's "
     "OpenAI-compatible endpoint. All ran with 0 harness errors.")
maas = [r0 for r0 in rows if r0["type"] == "Cloud MaaS"]
table(doc, ["Model", "Vendor", "Mode", "Pass", "Rubric"],
      [[disp(r0["model"]), r0["vendor"], MODE.get(r0["model"], "—"),
        f'{r0["pass_rate_pct"]}%{MARK.get(r0["model"],"")}', r0["rubric"]] for r0 in maas],
      widths=[2.2, 1.2, 1.2, 0.8, 0.8])

h(doc, "5.3 Local open-source (on-device deployment candidates)", 2)
para(doc, "Run locally via Ollama — the realistic offline targets. Device tier indicates where the "
     "model fits.")
local = [r0 for r0 in rows if r0["type"] == "Local OSS"]
table(doc, ["Model", "Params", "Device tier", "Pass", "Rubric"],
      [[disp(r0["model"]), OSS.get(r0["model"], ("?", "?"))[0], OSS.get(r0["model"], ("?", "?"))[1],
        f'{r0["pass_rate_pct"]}%{MARK.get(r0["model"],"")}', r0["rubric"]] for r0 in local],
      widths=[2.0, 0.8, 1.6, 0.8, 0.8])

# 6. Key findings & insights
h(doc, "6. Key Findings & Insights", 1)
numbered(doc, "**The frontier ceiling is Claude Opus 4.7 at 90%.** The best locally-run open model "
         "(qwen2.5:14b, 55%) reaches ~61% of that — a real but closeable gap, especially given no "
         "tutor-specific tuning has been applied yet.")
numbered(doc, "**Grok 4.1-fast (reasoning, 72%) is the best non-Anthropic model in the benchmark** — "
         "ahead of every Gemini, DeepSeek, Qwen, and GLM, and behind only the three Claude models. "
         "As cloud MaaS it is a strong frontier option when on-device isn't required.")
numbered(doc, "**A cloud open-weight tier rivals proprietary Flash models.** Qwen3-Coder 480B (68%), "
         "GLM-4.7 (67%), Qwen3-Next-80B-Instruct (65%), and Qwen3-235B (63%) all beat Gemini 2.5 "
         "Flash (65%) or sit alongside it — and beat every locally-run open model. Hostable open "
         "weights can reach the Flash tier.")
numbered(doc, "**Instruct/non-thinking beats pure-thinking for this tool-driven tutor.** The biggest "
         "single effect in the study: qwen3-next-80b-**thinking** scored 2% versus 65% for its "
         "**instruct** sibling, and deepseek-r1 (reasoning-only) scored 22%. Reasoning tokens "
         "exhaust the response budget before the model emits its tool-call. **Default to "
         "instruct/non-thinking variants** for the tutor.")
numbered(doc, "**Reasoning mode is model-dependent.** It helps Grok 4.1-fast (72% reasoning vs 57% "
         "non-reasoning) but hurts Grok 4.20 (45% vs 48%). Test both modes per model rather than "
         "assuming reasoning always helps.")
numbered(doc, "**Newer is not always better.** GLM-4.7 (67%) beats GLM-5 (57%), and Grok 4.1-fast "
         "beats Grok 4.20 on both modes. Always re-measure new versions against the harness.")
numbered(doc, "**Model family matters more than size.** mistral-nemo:12b (53%) beats llama3.1:8b "
         "(33%) by 20 points; picking the right family beats simply going bigger.")
numbered(doc, "**Universal weak spots: math reasoning and persona/tone adaptation** are the top two "
         "failure categories across nearly the entire field — both addressable with targeted prompt "
         "or fine-tuning on the chosen model.")
numbered(doc, "**Three local models scored 0% on a tool-protocol failure, not a teaching failure** "
         "(falcon3:10b, gemma2:2b, phi4) — they emit no or malformed tool-calls. phi4 (a strong 14B) "
         "is worth recovering with a format-specific parser, as was done for GLM-4.")

# 7. Data integrity & reliability
h(doc, "7. Data Integrity & Reliability", 1)
para(doc, "The Qwen3/Grok/GLM batch was first run during a **client-side internet outage**. The "
     "tutoring itself completed (tool-calls verified working), but the Anthropic rubric-judge calls "
     "failed with `APIConnectionError` for 7 of 10 models, which the harness recorded as scenario "
     "**failures** — making several capable models appear to score 0%.")
para(doc, "This was caught by inspecting per-model judge-failure counts rather than trusting raw pass "
     "rates: the three models that ran during connected windows had 0 judge failures, while the "
     "others had 56–60 of 60. Once connectivity returned, the affected models were re-run; **every "
     "number in this report comes from a run with `judge_failed = 0`.** The episode is a reminder to "
     "verify judge health before trusting a sweep — and a flagged follow-up is to harden the harness "
     "so a judge connectivity failure is recorded as an **errored** scenario (or retried) rather than "
     "silently scored as a model failure.")

# 8. Cost & footprint
h(doc, "8. Cost & Footprint", 1)
bullet(doc, "**Local OSS:** ≤9B models ran on an 8 GB CPU laptop (zero infra cost); 7–14B on a free "
       "Colab T4. Fully offline-capable — the deployment goal.")
bullet(doc, "**Cloud MaaS (Vertex Model Garden):** pay-per-token only, no hourly endpoint charge "
       "(no GPU endpoints deployed). A 60-scenario single-turn sweep is roughly US$0.50–$2 per model.")
bullet(doc, "**Scoring:** Anthropic judge + student-simulator calls are constant across every run and "
       "are a small, fixed per-sweep cost.")

# 9. Limitations & next steps
h(doc, "9. Limitations & Next Steps", 1)
bullet(doc, "**Single-turn only.** This sweep uses the 60 single-turn scenarios for fast, comparable "
       "signal; multi-turn evaluation of the finalists is the next measurement step.")
bullet(doc, "**Stock models, no tuning.** All scores are out-of-the-box. The gap to the ceiling is the "
       "target for prompt-engineering and fine-tuning on the leading candidates.")
bullet(doc, "**Suspect rows to diagnose:** the Gemini-Pro anomaly and the 0% tool-protocol models "
       "(phi4, falcon3) are under-measured by harness artifacts, not ability — worth re-scoring.")
bullet(doc, "**Recommended path:** (1) prompt-tune the leading on-device candidate (qwen2.5:7b / :3b) "
       "on math + persona; (2) decide the deployment tier (3B on-device vs 7–14B on a school server "
       "vs cloud MaaS for connected sites); (3) re-measure finalists multi-turn.")

# 10. Reproducibility
h(doc, "10. Reproducibility", 1)
para(doc, "The entire benchmark is scripted and the headline table in this report is generated "
     "directly from committed data. To reproduce or extend it:")

h(doc, "10.1 What lives where", 2)
bullet(doc, "`offline_eval/cloud_models.txt` — the model matrix: `provider/model  safe_name  [region]` "
       "(third column = Vertex region; differs per model — most are `global`).")
bullet(doc, "`offline_eval/run_cloud.sh` — runs the harness once per model, swapping only the tutor.")
bullet(doc, "`offline_eval/_probe_cloud_models.py` — 1-token reachability probe; validates IDs/regions "
       "before a full run.")
bullet(doc, "`offline_eval/aggregate.py` — ranks `results/*.json` into the leaderboard.")
bullet(doc, "`offline_eval/leaderboard_combined.csv` — the combined table as data (this report's §4 "
       "is built from it via `offline_eval/_make_report.py`).")
bullet(doc, "`offline_eval/results/*.json` — per-model run records (a curated set is committed for "
       "reproducibility).")
bullet(doc, "Integration: `apps/llm/client.py::VertexModelGardenClient`. Design + plan: "
       "`docs/superpowers/specs/2026-06-17-vertex-model-garden-eval-design.md` and the matching "
       "plan under `docs/superpowers/plans/`.")

h(doc, "10.2 How the swap works", 2)
para(doc, "The harness sets the `TUTOR_MODEL_OVERRIDE=\"provider/model\"` environment variable, which "
     "routes only the tutor through the chosen model. The judge and student-simulator stay on "
     "Anthropic, so the yardstick is identical for every model and the grader is always cross-family.")

h(doc, "10.3 Google Cloud setup (Vertex Model Garden MaaS)", 2)
numbered(doc, "Authenticate Application Default Credentials and pin the project. Auth is isolated in a "
         "dedicated config directory so it never touches a personal gcloud profile:")
mono(doc, "export CLOUDSDK_CONFIG=\"$HOME/.config/gcloud-pixeldesignlabs\"")
mono(doc, "gcloud auth application-default login")
mono(doc, "gcloud auth application-default set-quota-project ai-tutor-499714")
numbered(doc, "Enable the Vertex AI API:")
mono(doc, "gcloud services enable aiplatform.googleapis.com --project=ai-tutor-499714")
numbered(doc, "In Model Garden (console.cloud.google.com/agent-platform/model-garden), enable each "
         "model — choose the “… API Service” (MaaS, pay-per-token) cards, not the “Serve with …” / "
         "GPU-endpoint cards. The card shows the exact model ID and region.")
numbered(doc, "`.env` carries `CLOUDSDK_CONFIG` + `GOOGLE_CLOUD_PROJECT` so the harness picks up the "
         "isolated credentials automatically (it calls `load_dotenv()`).")

h(doc, "10.4 Run a sweep", 2)
mono(doc, "# validate model IDs + regions are reachable (cheap):")
mono(doc, "CLOUDSDK_CONFIG=$HOME/.config/gcloud-pixeldesignlabs \\")
mono(doc, "  venv/bin/python offline_eval/_probe_cloud_models.py")
mono(doc, "")
mono(doc, "# run the full sweep (skips models that already have results/*.json):")
mono(doc, "CLOUDSDK_CONFIG=$HOME/.config/gcloud-pixeldesignlabs \\")
mono(doc, "  bash offline_eval/run_cloud.sh")
mono(doc, "")
mono(doc, "# rebuild the leaderboard + CSV + this report:")
mono(doc, "venv/bin/python offline_eval/aggregate.py")
mono(doc, "venv/bin/python offline_eval/_make_report.py")

h(doc, "10.5 Verify judge health (do this before trusting a sweep)", 2)
para(doc, "A network/judge outage during a run can silently score capable models as failures (see §7). "
     "After any sweep, confirm each model's judge-failure count is zero:")
mono(doc, "grep -c \"judge call failed\\|APIConnectionError\" offline_eval/results/<safe>.log   # expect 0")
para(doc, "Re-run any affected model once connectivity returns:")
mono(doc, "CLOUD_MODELS_FILE=<subset.txt> FORCE=1 bash offline_eval/run_cloud.sh")

h(doc, "10.6 Add a new model", 2)
para(doc, "Append one row to `cloud_models.txt` (`provider/model  safe_name  region`), add the same "
     "candidate to `_probe_cloud_models.py`, then run §10.4. No code changes are needed for a new "
     "Vertex MaaS model — the provider client is generic.")

doc.save(OUT)
print(f"wrote {OUT}")
print(f"  {len(rows)} models in combined table; "
      f"{len(prop)} proprietary, {len(maas)} cloud MaaS, {len(local)} local OSS")
