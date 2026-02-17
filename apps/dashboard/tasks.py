"""
Curriculum Processing Tasks

Handles:
1. Parsing uploaded curriculum documents (PDF, DOCX)
2. Using AI to extract structure (units, lessons, objectives)
3. Auto-generating lesson content and exit tickets
"""

import os
import json
import logging
from django.utils import timezone

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: str) -> str:
    """Extract text content from uploaded file."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return extract_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_from_docx(file_path)
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except ImportError:
        # Fallback to pdfplumber
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            return text
        except ImportError:
            raise ImportError("Please install PyMuPDF or pdfplumber: pip install PyMuPDF pdfplumber")


def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    except ImportError:
        raise ImportError("Please install python-docx: pip install python-docx")


def parse_curriculum_with_ai(text: str, subject_name: str, grade_level: str) -> dict:
    """Use AI to parse curriculum text into structured format."""
    from apps.llm.models import ModelConfig
    from apps.llm.client import get_llm_client
    
    # Get active model config
    model_config = ModelConfig.objects.filter(is_active=True).first()
    if not model_config:
        raise ValueError("No active LLM model configured")
    
    client = get_llm_client(model_config)
    
    system_prompt = """You are a curriculum parsing assistant. Extract the structure from curriculum documents.

Output ONLY valid JSON with this structure:
{
    "subject": "Subject Name",
    "grade_level": "S1/S2/S3/etc",
    "description": "Brief subject description",
    "units": [
        {
            "name": "Unit Name",
            "description": "Unit description",
            "order": 1,
            "lessons": [
                {
                    "title": "Lesson Title",
                    "objective": "What students will learn",
                    "estimated_minutes": 30,
                    "order": 1,
                    "topics": ["Topic 1", "Topic 2"],
                    "key_vocabulary": ["term1", "term2"]
                }
            ]
        }
    ]
}

Be thorough - extract ALL units and lessons mentioned in the curriculum.
Use the exact names/titles from the document when available.
"""

    user_message = f"""Parse this {subject_name} curriculum for {grade_level or 'secondary school'}:

{text[:15000]}  # Limit to avoid token limits

Extract all units and lessons. Return ONLY the JSON structure."""

    response = client.generate(
        messages=[{"role": "user", "content": user_message}],
        system_prompt=system_prompt
    )
    
    # Parse JSON from response
    content = response.content.strip()
    
    # Try to extract JSON
    if content.startswith('```'):
        # Remove markdown code blocks
        content = content.split('```')[1]
        if content.startswith('json'):
            content = content[4:]
        content = content.strip()
    
    try:
        return json.loads(content)
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse AI response as JSON: {e}")
        logger.error(f"Response was: {content[:500]}")
        raise ValueError(f"AI returned invalid JSON: {e}")


def create_curriculum_from_structure(structure: dict, institution, upload) -> dict:
    """Create Course, Units, and Lessons from parsed structure."""
    from apps.curriculum.models import Course, Unit, Lesson, LessonStep
    
    # Create Course directly (no Subject model)
    grade = structure.get('grade_level', upload.grade_level) or 'S1'
    subject_name = structure.get('subject', upload.subject_name)
    course_name = f"{subject_name} {grade}"
    
    course, created = Course.objects.get_or_create(
        institution=institution,
        title=course_name,
        defaults={
            'grade_level': grade,
            'description': structure.get('description', ''),
            'is_published': False,  # Start unpublished for review
        }
    )
    upload.created_course = course
    upload.add_log(f"{'Created' if created else 'Using existing'} course: {course.title}")
    
    lessons_created = 0
    
    # Create Units and Lessons
    for unit_data in structure.get('units', []):
        unit, created = Unit.objects.get_or_create(
            course=course,
            title=unit_data.get('name', 'Unnamed Unit'),
            defaults={
                'description': unit_data.get('description', ''),
                'order_index': unit_data.get('order', 0),
            }
        )
        upload.add_log(f"  {'Created' if created else 'Using existing'} unit: {unit.title}")
        
        for lesson_data in unit_data.get('lessons', []):
            lesson, created = Lesson.objects.get_or_create(
                unit=unit,
                title=lesson_data.get('title', 'Unnamed Lesson'),
                defaults={
                    'objective': lesson_data.get('objective', ''),
                    'estimated_minutes': lesson_data.get('estimated_minutes', 30),
                    'order_index': lesson_data.get('order', 0),
                    'is_published': False,
                }
            )
            
            if created:
                lessons_created += 1
                
                # Create a basic TEACH step for conversational mode
                LessonStep.objects.create(
                    lesson=lesson,
                    step_type='teach',
                    order_index=0,
                    teacher_script=f"Teach about: {lesson.objective}",
                )
                
                upload.add_log(f"    Created lesson: {lesson.title}")
            else:
                upload.add_log(f"    Existing lesson: {lesson.title}")
    
    upload.lessons_created = lessons_created
    upload.save()
    
    return {
        'course_id': course.id,
        'lessons_created': lessons_created,
    }


def generate_lesson_content(lesson_id: int) -> dict:
    """Generate detailed lesson content using AI."""
    from apps.curriculum.models import Lesson, LessonStep
    from apps.llm.models import ModelConfig
    from apps.llm.client import get_llm_client
    
    lesson = Lesson.objects.get(id=lesson_id)
    
    model_config = ModelConfig.objects.filter(is_active=True).first()
    if not model_config:
        raise ValueError("No active LLM model configured")
    
    client = get_llm_client(model_config)
    
    system_prompt = """You are an expert teacher creating lesson content for secondary school students in Seychelles.

Create engaging, culturally relevant content that includes:
1. Clear explanations with local examples
2. Step-by-step worked examples
3. Practice questions with hints

Output JSON with this structure:
{
    "teaching_content": "Main explanation content with Seychelles examples...",
    "key_points": ["Point 1", "Point 2", "Point 3"],
    "worked_example": {
        "problem": "Example problem",
        "steps": ["Step 1", "Step 2", "Step 3"],
        "answer": "Final answer"
    },
    "practice_questions": [
        {
            "question": "Question text",
            "options": ["A) ...", "B) ...", "C) ...", "D) ..."],
            "correct": "A",
            "hints": ["Hint 1", "Hint 2"]
        }
    ]
}
"""
    
    user_message = f"""Create detailed lesson content for:

Course: {lesson.unit.course.title}
Unit: {lesson.unit.title}
Lesson: {lesson.title}
Objective: {lesson.objective}

Make it engaging for Seychelles secondary students. Include local examples and context."""
    
    response = client.generate(
        messages=[{"role": "user", "content": user_message}],
        system_prompt=system_prompt
    )
    
    # Parse response
    content = response.content.strip()
    if content.startswith('```'):
        content = content.split('```')[1]
        if content.startswith('json'):
            content = content[4:]
        content = content.strip()
    
    try:
        lesson_content = json.loads(content)
    except json.JSONDecodeError:
        logger.error(f"Failed to parse lesson content: {content[:500]}")
        return {'error': 'Failed to generate content'}
    
    # Update lesson step with teaching content
    teach_step = lesson.steps.filter(step_type='teach').first()
    if teach_step:
        teach_step.teacher_script = lesson_content.get('teaching_content', '')
        teach_step.save()
    
    return lesson_content


def generate_exit_ticket_for_lesson(lesson) -> bool:
    """Generate an exit ticket for a single lesson."""
    from apps.tutoring.exit_ticket_models import ExitTicket, ExitTicketQuestion
    from apps.llm.models import ModelConfig
    from apps.llm.client import get_llm_client
    
    # Skip if already has exit ticket
    if ExitTicket.objects.filter(lesson=lesson).exists():
        return False
    
    model_config = ModelConfig.objects.filter(is_active=True).first()
    if not model_config:
        logger.warning("No active model config for exit ticket generation")
        return False
    
    llm_client = get_llm_client(model_config)
    
    subject = lesson.unit.course.title if lesson.unit and lesson.unit.course else "General"
    
    prompt = f"""Generate exactly 10 multiple choice questions for a summative assessment on this lesson.

LESSON: {lesson.title}
OBJECTIVE: {lesson.objective}
SUBJECT: {subject}

REQUIREMENTS:
1. Generate EXACTLY 10 questions
2. Questions should progress from easy to hard
3. Each question must have exactly 4 options (A, B, C, D)
4. Use context relevant to Seychelles secondary school students

OUTPUT FORMAT (JSON array only, no other text):
[
    {{
        "question": "What is...?",
        "option_a": "First option",
        "option_b": "Second option", 
        "option_c": "Third option",
        "option_d": "Fourth option",
        "correct": "B",
        "explanation": "Brief explanation",
        "difficulty": "easy"
    }}
]

Generate the 10 questions now:"""

    try:
        response = llm_client.generate(
            messages=[{"role": "user", "content": prompt}],
            system_prompt="You are an expert educational assessment designer. Output only valid JSON."
        )
        
        content = response.content
        start = content.find('[')
        end = content.rfind(']') + 1
        
        if start == -1 or end <= start:
            logger.error(f"No JSON array in exit ticket response for lesson {lesson.id}")
            return False
        
        questions = json.loads(content[start:end])
        
        if len(questions) < 10:
            logger.warning(f"Only {len(questions)} questions generated for lesson {lesson.id}")
        
        # Create exit ticket
        exit_ticket = ExitTicket.objects.create(
            lesson=lesson,
            passing_score=8,
            time_limit_minutes=15,
            instructions=f"Answer all 10 questions about {lesson.title}. You need 8 correct to pass."
        )
        
        for i, q in enumerate(questions[:10]):
            ExitTicketQuestion.objects.create(
                exit_ticket=exit_ticket,
                question_text=q['question'],
                option_a=q['option_a'],
                option_b=q['option_b'],
                option_c=q['option_c'],
                option_d=q['option_d'],
                correct_answer=q['correct'],
                explanation=q.get('explanation', ''),
                difficulty=q.get('difficulty', 'medium'),
                order_index=i,
            )
        
        return True
        
    except Exception as e:
        logger.error(f"Exit ticket generation failed for lesson {lesson.id}: {e}")
        return False


def process_curriculum_upload(upload_id: int) -> dict:
    """Main function to process a curriculum upload."""
    from apps.dashboard.models import CurriculumUpload
    
    upload = CurriculumUpload.objects.get(id=upload_id)
    
    try:
        upload.status = 'processing'
        upload.add_log("Starting curriculum processing...")
        
        # Step 1: Extract text from file
        upload.add_log("Extracting text from document...")
        text = extract_text_from_file(upload.file_path)
        upload.add_log(f"Extracted {len(text)} characters")
        
        # Step 2: Parse with AI
        upload.add_log("Parsing curriculum structure with AI...")
        structure = parse_curriculum_with_ai(
            text,
            upload.subject_name,
            upload.grade_level
        )
        upload.add_log(f"Found {len(structure.get('units', []))} units")
        
        # Step 3: Create database records
        upload.add_log("Creating curriculum records...")
        result = create_curriculum_from_structure(
            structure,
            upload.institution,
            upload
        )
        
        # Step 4: Generate exit tickets for new lessons
        upload.add_log("Generating exit tickets...")
        from apps.curriculum.models import Lesson
        
        course = upload.created_course
        if course:
            lessons = Lesson.objects.filter(unit__course=course)
            tickets_created = 0
            for lesson in lessons:
                upload.add_log(f"  Generating exit ticket for: {lesson.title}...")
                if generate_exit_ticket_for_lesson(lesson):
                    tickets_created += 1
                    upload.add_log(f"    ✓ Created")
                else:
                    upload.add_log(f"    ⏭️ Skipped (exists or failed)")
            upload.add_log(f"Generated {tickets_created} exit tickets")
        
        # Mark complete
        upload.status = 'completed'
        upload.completed_at = timezone.now()
        upload.add_log(f"✓ Completed! Created {result['lessons_created']} lessons.")
        upload.save()
        
        return result
        
    except Exception as e:
        logger.exception(f"Curriculum processing failed: {e}")
        upload.status = 'failed'
        upload.error_message = str(e)
        upload.add_log(f"✗ Error: {e}")
        upload.save()
        raise
