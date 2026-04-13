"""
Curriculum Pipeline Processor

This module implements the 5-step curriculum processing pipeline:

1. PARSE: Extract text from PDF/DOCX
2. VECTORIZE: Chunk and embed into vector database
3. GENERATE LESSONS: Query DB to structure curriculum into units/lessons
4. GENERATE CONTENT: Generate tutoring steps and media with curriculum context
5. TUTORING: Provide live context during tutoring sessions

Each step builds on the previous, creating a rich, context-aware tutoring system.
"""

import os
import json
import logging
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict

from pydantic import BaseModel, Field
from django.utils import timezone

logger = logging.getLogger(__name__)


# =============================================================================
# STRUCTURED OUTPUT SCHEMAS
# =============================================================================

class LessonSchema(BaseModel):
    """A single lesson in a unit."""
    title: str = Field(description="Short lesson title (3-8 words)")
    objective: str = Field(description="Students will be able to [specific, measurable outcome]")
    key_concepts: List[str] = Field(default_factory=list, description="Key concepts covered")
    enabling_objectives: List[str] = Field(
        default_factory=list,
        description="Specific enabling objectives (granular teaching steps) for this lesson"
    )


class UnitSchema(BaseModel):
    """A unit containing multiple lessons."""
    title: str = Field(description="Clear unit title")
    description: str = Field(default="", description="Brief description of what this unit covers")
    grade_level: str = Field(default="", description="Target grade level(s), e.g. 'S1' or 'S1,S2'")
    terminal_objectives: List[str] = Field(
        default_factory=list,
        description="Unit-level terminal objectives — what students must achieve by end of unit"
    )
    enabling_objectives: List[str] = Field(
        default_factory=list,
        description="Unit-level enabling objectives — all granular teaching steps"
    )
    lessons: List[LessonSchema] = Field(description="Lessons in this unit")


class LessonStructureResult(BaseModel):
    """Complete curriculum structure with units and lessons."""
    units: List[UnitSchema] = Field(description="4-8 logical units based on major topics")


class ExitTicketQuestion(BaseModel):
    """An exit ticket question."""
    question: str
    correct_answer: str
    distractors: List[str] = Field(default_factory=list)


class TutoringStep(BaseModel):
    """A single step in a tutoring session."""
    phase: str = Field(description="5E phase: engage, explore, explain, practice, or evaluate")
    step_type: str = Field(description="Step type: teach, question, or reflect")
    content: str = Field(description="What the tutor says/shows")
    question: Optional[str] = Field(default=None, description="Question to ask (for question type)")
    expected_answer: Optional[str] = Field(default=None, description="Correct answer")
    hints: Optional[List[str]] = Field(default=None, description="Hints if student struggles")
    media_suggestion: Optional[str] = Field(default=None, description="Description of helpful image/diagram")


class LessonContentResult(BaseModel):
    """Generated tutoring content for a lesson."""
    steps: List[TutoringStep] = Field(description="Tutoring session steps")
    exit_ticket: List[ExitTicketQuestion] = Field(default_factory=list)
    key_vocabulary: List[str] = Field(default_factory=list)
    teaching_tips: List[str] = Field(default_factory=list)


# ============================================================================
# STEP 1: PARSE
# ============================================================================

def extract_text_from_file(file_path: str) -> Tuple[str, str]:
    """
    Extract text from curriculum document.
    
    Supports: PDF, DOCX, TXT, MD
    
    Returns: (text, file_type)
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in ['.txt', '.md']:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return text, 'text'
    
    elif ext == '.docx':
        return _extract_from_docx(file_path), 'docx'
    
    elif ext == '.pdf':
        return _extract_from_pdf(file_path), 'pdf'

    elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.tiff', '.tif']:
        from apps.curriculum.curriculum_parser import extract_from_image
        return extract_from_image(file_path), 'image'

    else:
        # Try reading as text
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return text, 'text'
        except:
            raise ValueError(f"Unsupported file type: {ext}")


def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"python-docx failed: {e}, trying as text")
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF file (delegates to shared implementation with LLM vision fallback)."""
    from apps.curriculum.curriculum_parser import extract_from_pdf
    return extract_from_pdf(file_path)


# ============================================================================
# STEP 2: VECTORIZE (via Knowledge Base)
# ============================================================================

def vectorize_curriculum(
    file_path: str,
    subject: str,
    grade_level: str,
    institution_id: int,
    upload_id: int = None
) -> Dict:
    """
    Vectorize curriculum document into the knowledge base.
    
    Args:
        file_path: Path to curriculum document
        subject: Subject name
        grade_level: Grade level
        institution_id: Institution ID
        upload_id: Optional CurriculumUpload ID
    
    Returns:
        Dict with indexing statistics
    """
    from apps.curriculum.knowledge_base import CurriculumKnowledgeBase
    
    kb = CurriculumKnowledgeBase(institution_id=institution_id)
    
    return kb.index_curriculum_document(
        file_path=file_path,
        subject=subject,
        grade_level=grade_level,
        curriculum_upload_id=upload_id
    )


def _get_instructor_client():
    """Get an instructor-wrapped client for structured curriculum generation."""
    try:
        import instructor
        from apps.llm.models import ModelConfig

        config = ModelConfig.get_for('generation')
        if not config:
            return None, None
        PROVIDER_MAP = {
            'anthropic': 'anthropic',
            'openai': 'openai',
            'google': 'google',
            'local_ollama': 'ollama',
        }
        provider = PROVIDER_MAP.get(config.provider, config.provider)
        client = instructor.from_provider(
            f"{provider}/{config.model_name}",
            api_key=config.get_api_key(),
        )
        return client, config
    except Exception as e:
        logger.warning(f"Could not create instructor client for pipeline: {e}")
        return None, None


# ============================================================================
# STEP 2b: KB-BASED COMPETENCY EXTRACTION (format-agnostic)
# ============================================================================

def extract_competencies_from_kb(
    subject: str,
    grade_level: str,
    institution_id: int,
    unit_title: str = "",
) -> List[Dict]:
    """
    Extract standardized competencies from the KB using LLM analysis.

    This is format-agnostic — it queries the vectorized KB for relevant chunks
    from ANY uploaded document (curriculum, textbook, worksheet, exam paper)
    and uses an LLM to extract normalized competency statements.

    Returns a list of competency dicts:
    [{"text": "Define population density", "type": "knowledge|skill|attitude",
      "bloom_level": "remember", "source_code": "K408"}]
    """
    from apps.curriculum.knowledge_base import CurriculumKnowledgeBase
    from apps.llm.models import ModelConfig
    from apps.llm.client import get_llm_client

    kb = CurriculumKnowledgeBase(institution_id=institution_id)
    result = kb.query_for_competency_extraction(
        subject=subject,
        grade_level=grade_level,
        unit_title=unit_title,
        n_results=30,
    )

    if not result.chunks:
        return []

    # Build context from KB chunks
    kb_context = "\n\n".join(
        f"[Source: {c.get('section', 'Document')}]\n{c.get('content', '')}"
        for c in result.chunks[:15]
        if c.get('content', '').strip()
    )

    if not kb_context.strip():
        return []

    model_config = ModelConfig.get_for('generation')
    if not model_config:
        logger.warning("No LLM model configured for competency extraction")
        return []

    llm_client = get_llm_client(model_config)

    prompt = f"""Analyze the following curriculum/teaching material and extract ALL specific learning competencies (objectives, skills, knowledge items) that students must master.

SUBJECT: {subject}
GRADE LEVEL: {grade_level}
{f'UNIT/TOPIC: {unit_title}' if unit_title else ''}

=== SOURCE MATERIAL ===
{kb_context}
=== END MATERIAL ===

Extract every specific, measurable competency. For each one, provide:
- "text": The competency statement (action verb + specific content, e.g., "Define population density")
- "type": "knowledge", "skill", or "attitude"
- "bloom_level": "remember", "understand", "apply", "analyze", "evaluate", or "create"
- "source_code": Original curriculum code if visible (e.g., "K408", "S401"), otherwise ""
- "strand": The topic/strand this belongs to (e.g., "Number", "Population", "Map Skills")

Return a JSON array of competency objects. Extract as many as you can find — do NOT summarize or merge.
Return ONLY valid JSON."""

    try:
        response = llm_client.generate(
            messages=[{"role": "user", "content": prompt}],
            system_prompt="You are a curriculum analysis expert. Extract learning competencies precisely. Return only valid JSON.",
            max_tokens=8000,
        )

        from apps.llm.json_utils import parse_llm_json
        competencies = parse_llm_json(response.content, expect_array=True)
        if competencies and isinstance(competencies, list):
            logger.info(f"Extracted {len(competencies)} competencies from KB for {subject} {grade_level}")
            return competencies
    except Exception as e:
        logger.warning(f"KB competency extraction failed: {e}")

    return []


def _merge_competencies_into_structure(structure: Dict, competencies: List[Dict]):
    """Merge KB-extracted competencies into the lesson structure as enabling_objectives.

    Uses keyword matching to assign each competency to the most relevant unit/lesson.
    Competencies that don't match any lesson are added to the unit level.
    """
    units = structure.get('units', [])
    if not units or not competencies:
        return

    for comp in competencies:
        comp_text = comp.get('text', '')
        comp_full = f"{comp.get('source_code', '')}: {comp_text}".strip(': ')
        comp_lower = comp_text.lower()
        best_lesson = None
        best_score = 0

        for unit in units:
            for lesson in unit.get('lessons', []):
                # Score by keyword overlap between competency and lesson title/objective
                lesson_text = f"{lesson.get('title', '')} {lesson.get('objective', '')}".lower()
                words = set(comp_lower.split()) & set(lesson_text.split())
                score = len(words)
                if score > best_score:
                    best_score = score
                    best_lesson = lesson

        if best_lesson and best_score >= 2:
            best_lesson.setdefault('enabling_objectives', [])
            if comp_full not in best_lesson['enabling_objectives']:
                best_lesson['enabling_objectives'].append(comp_full)
        else:
            # Add to the first unit's enabling_objectives as unassigned
            units[0].setdefault('enabling_objectives', [])
            if comp_full not in units[0]['enabling_objectives']:
                units[0]['enabling_objectives'].append(comp_full)


# ============================================================================
# STEP 3: GENERATE LESSONS (Query KB + LLM)
# ============================================================================

def generate_lesson_structure(
    subject: str,
    grade_level: str,
    institution_id: int,
    extracted_text: str = None
) -> Dict:
    """
    Generate structured lessons from the curriculum.

    Uses instructor for structured output. Falls back to raw LLM + JSON
    repair if instructor is unavailable.

    Args:
        subject: Subject name
        grade_level: Grade level
        institution_id: Institution ID
        extracted_text: Optional raw text (if KB not yet indexed)

    Returns:
        Dict with units and lessons structure
    """
    from apps.curriculum.knowledge_base import CurriculumKnowledgeBase
    from apps.llm.models import ModelConfig
    from apps.llm.client import get_llm_client

    # For pilot subjects: try dedicated parsers first (more accurate extraction)
    print(f"[generate_lesson_structure] subject={subject} grade={grade_level} text_len={len(extracted_text or '')}", flush=True)
    if extracted_text:
        subject_lower = subject.lower()
        dedicated_result = None

        if subject_lower in ('mathematics', 'maths', 'math'):
            try:
                from apps.curriculum.curriculum_parser import parse_mathematics_curriculum
                print(f"[generate_lesson_structure] Trying math parser...", flush=True)
                result = parse_mathematics_curriculum(extracted_text, grade_level)
                if result.units and sum(len(u.get('lessons', [])) for u in result.units) >= 5:
                    dedicated_result = result
                    print(f"[generate_lesson_structure] Math parser: {len(result.units)} strands", flush=True)
            except Exception as e:
                print(f"[generate_lesson_structure] Math parser failed: {e}", flush=True)

        elif subject_lower in ('geography', 'geo'):
            try:
                from apps.curriculum.curriculum_parser import parse_geography_curriculum
                print(f"[generate_lesson_structure] Trying geography parser...", flush=True)
                result = parse_geography_curriculum(extracted_text, grade_level)
                units_count = len(result.units) if result.units else 0
                lessons_count = sum(len(u.get('lessons', [])) for u in result.units) if result.units else 0
                print(f"[generate_lesson_structure] Geography parser: {units_count} units, {lessons_count} lessons", flush=True)
                if result.units and lessons_count >= 3:
                    dedicated_result = result
            except Exception as e:
                print(f"[generate_lesson_structure] Geography parser failed: {e}", flush=True)
                import traceback; traceback.print_exc()

        if dedicated_result:
            print(f"[generate_lesson_structure] Using dedicated parser result", flush=True)
            return {
                'subject': subject,
                'grade_level': grade_level,
                'units': dedicated_result.units,
                'total_lessons': sum(len(u.get('lessons', [])) for u in dedicated_result.units),
            }

    # Fall through to LLM-based parsing
    print(f"[generate_lesson_structure] Dedicated parser insufficient, using LLM...", flush=True)

    # Query knowledge base for context
    kb = CurriculumKnowledgeBase(institution_id=institution_id)
    kb_results = kb.query_for_lesson_generation(
        subject=subject,
        grade_level=grade_level,
        n_results=30
    )

    # Build context from KB results
    kb_context = ""
    if kb_results.chunks:
        kb_context = "\n\n".join([
            f"[{c.get('section', 'Content')}]\n{c.get('content', '')}"
            for c in kb_results.chunks[:20]
        ])

    # If no KB context, use raw text
    if not kb_context and extracted_text:
        kb_context = extracted_text[:40000]

    if not kb_context:
        raise ValueError("No curriculum content available")

    prompt = f"""Analyze this {subject} curriculum for {grade_level} students and create a well-organized lesson structure.

=== CURRICULUM CONTENT ===
{kb_context}
=== END CONTENT ===

Create a structured curriculum with UNITS and LESSONS.

REQUIREMENTS:
1. Create 4-8 logical UNITS based on major topics/strands in the curriculum
2. Each unit should have 8-20 LESSONS covering specific skills or concepts
3. Lesson titles should be SHORT (3-8 words), clear, action-oriented
4. Each lesson should teach ONE specific concept that can be covered in ~40 minutes
5. Each unit MUST include a "grade_level" field (e.g. "S1", "S2") indicating the target grade
6. Group lessons by grade level first, then by topic within each grade
7. Unit titles should include the grade prefix (e.g. "S1: Map Skills", "S2: Algebra")

OBJECTIVES EXTRACTION (CRITICAL):
8. Each unit MUST have "terminal_objectives" — broad outcomes students must achieve by end of unit.
   Extract these from the curriculum document (look for numbered objectives, learning outcomes, or
   "students should be able to" statements at the unit/topic level).
9. Each unit MUST have "enabling_objectives" — the granular, specific teaching steps from the
   Teaching/Learning Scheme table. These are more detailed than terminal objectives (typically
   20-40 per unit). Look for action verbs: define, state, explain, describe, identify, compare.
10. Each lesson MUST have "enabling_objectives" — the specific enabling objectives this lesson
    covers (a subset of the unit's enabling objectives). Each enabling objective should appear
    in exactly one lesson.

For {subject}, organize by these strands where applicable:
- Mathematics: Number, Algebra, Geometry, Measurement, Data/Statistics
- Geography: Physical Geography, Human Geography, Map Skills, Regional Studies
- Science: Life Science, Physical Science, Earth Science"""

    system_prompt = "You are a curriculum design expert. Create well-structured educational content."

    # LLM-based lesson structure generation
    # Strategy: chunk large documents and process each chunk, then merge
    CHUNK_SIZE = 15000  # chars per LLM call — safe for most models
    all_content = kb_context or (extracted_text[:40000] if extracted_text else "")

    if len(all_content) <= CHUNK_SIZE:
        chunks_to_process = [all_content]
    else:
        # Split into chunks at paragraph boundaries
        chunks_to_process = []
        remaining = all_content
        while remaining:
            if len(remaining) <= CHUNK_SIZE:
                chunks_to_process.append(remaining)
                break
            # Find a good split point (paragraph break near chunk size)
            split_at = remaining.rfind('\n\n', CHUNK_SIZE // 2, CHUNK_SIZE)
            if split_at == -1:
                split_at = remaining.rfind('\n', CHUNK_SIZE // 2, CHUNK_SIZE)
            if split_at == -1:
                split_at = CHUNK_SIZE
            chunks_to_process.append(remaining[:split_at])
            remaining = remaining[split_at:].lstrip()

        logger.info(f"Document split into {len(chunks_to_process)} chunks for LLM processing")

    all_units = []
    last_error = None
    client, config = _get_instructor_client()

    for chunk_idx, chunk_text in enumerate(chunks_to_process):
        chunk_prompt = f"""Analyze this {subject} curriculum for {grade_level} students and create a well-organized lesson structure.

=== CURRICULUM CONTENT (Part {chunk_idx + 1} of {len(chunks_to_process)}) ===
{chunk_text}
=== END CONTENT ===

{prompt.split('=== END CONTENT ===')[1] if '=== END CONTENT ===' in prompt else ''}"""

        for attempt in range(2):
            try:
                if client:
                    create_kwargs = dict(
                        response_model=LessonStructureResult,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": chunk_prompt},
                        ],
                        max_retries=2,
                    )
                    if config and config.provider == 'google':
                        create_kwargs['generation_config'] = {'max_tokens': 8192}
                    else:
                        create_kwargs['max_tokens'] = 8192

                    result = client.chat.completions.create(**create_kwargs)
                    chunk_units = result.model_dump().get('units', [])
                    all_units.extend(chunk_units)
                    logger.info(f"Chunk {chunk_idx + 1}: extracted {len(chunk_units)} units")
                    break
                else:
                    model_config = ModelConfig.get_for('generation')
                    if not model_config:
                        raise ValueError("No active LLM model configured")
                    llm_client = get_llm_client(model_config)

                    response = llm_client.generate(
                        messages=[{"role": "user", "content": chunk_prompt + "\n\nReturn ONLY valid JSON."}],
                        system_prompt=system_prompt,
                        max_tokens=8192,
                    )
                    content = _clean_json_response(response.content.strip())
                    try:
                        parsed = json.loads(content)
                    except json.JSONDecodeError:
                        parsed = _try_fix_json(content)

                    if parsed and parsed.get('units'):
                        all_units.extend(parsed['units'])
                        logger.info(f"Chunk {chunk_idx + 1}: extracted {len(parsed['units'])} units")
                    break

            except Exception as e:
                last_error = str(e)
                error_lower = str(e).lower()
                if any(kw in error_lower for kw in ['timeout', 'stream', 'deadline', 'timed out']):
                    logger.warning(f"Chunk {chunk_idx + 1} attempt {attempt + 1} timed out: {e}")
                    if attempt == 0:
                        # Retry with even shorter chunk
                        chunk_text = chunk_text[:CHUNK_SIZE // 2]
                        continue
                elif attempt == 0:
                    logger.warning(f"Chunk {chunk_idx + 1} attempt {attempt + 1} failed: {e}")
                    continue
                logger.error(f"Chunk {chunk_idx + 1} failed permanently: {e}")
                break

    if not all_units:
        # Final fallback: dedicated parsers
        if extracted_text:
            logger.warning(f"LLM parsing produced no units ({last_error}), trying dedicated parsers")
            try:
                from apps.curriculum.curriculum_parser import (
                    parse_mathematics_curriculum, parse_geography_curriculum, parse_generic_curriculum
                )
                sl = subject.lower()
                if 'math' in sl:
                    result = parse_mathematics_curriculum(extracted_text, grade_level)
                elif 'geo' in sl:
                    result = parse_geography_curriculum(extracted_text, grade_level)
                else:
                    result = parse_generic_curriculum(extracted_text, subject, grade_level)
                if result and result.units:
                    all_units = result.units
            except Exception:
                pass

        if not all_units:
            raise ValueError(f"Lesson structure generation failed: {last_error}")

    structure = {'units': all_units}

    # Validate and clean
    return _validate_lesson_structure(structure, subject, grade_level)


def _clean_json_response(content: str, truncated: bool = False) -> str:
    """
    Clean LLM response to get valid JSON.

    If the JSON is truncated (brace_count never reaches 0), returns the raw
    content so that _repair_truncated_json / _try_fix_json can handle it.
    """
    # Remove markdown code blocks
    if '```' in content:
        parts = content.split('```')
        for part in parts:
            part = part.strip()
            if part.startswith('json'):
                part = part[4:].strip()
            if part.startswith('{'):
                content = part
                break

    # Find the JSON object boundaries
    if content.startswith('{'):
        brace_count = 0
        end_pos = 0
        for i, char in enumerate(content):
            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0:
                    end_pos = i + 1
                    break
        if end_pos > 0:
            content = content[:end_pos]
        elif brace_count > 0:
            # JSON is truncated — braces never balanced
            logger.warning(f"JSON appears truncated (unclosed braces: {brace_count})")

    return content.strip()


def _repair_truncated_json(content: str) -> Optional[Dict]:
    """
    Attempt to repair JSON that was truncated mid-stream (e.g. stop_reason='max_tokens').

    Uses a stack to track bracket order and closes them correctly.
    """
    import re

    text = content.rstrip()

    # If we're inside an unclosed string (odd unescaped quotes), strip back
    quote_count = 0
    for i, ch in enumerate(text):
        if ch == '"' and (i == 0 or text[i-1] != '\\'):
            quote_count += 1
    if quote_count % 2 != 0:
        for pos in range(len(text) - 1, 0, -1):
            if text[pos] == '"' and (pos == 0 or text[pos-1] != '\\'):
                text = text[:pos]
                break

    # Remove trailing comma, colon, or partial key-value
    text = re.sub(r'[,:]\s*$', '', text.rstrip())
    # Strip dangling key without value (e.g. `, "title"` or `{ "title"` at end)
    text = re.sub(r'([{,])\s*"[^"]*"\s*$', r'\1', text.rstrip())
    text = re.sub(r'[{,]\s*$', '', text.rstrip())

    # Build bracket stack to know closing order
    stack = []
    in_string = False
    escape_next = False
    for ch in text:
        if escape_next:
            escape_next = False
            continue
        if ch == '\\' and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            stack.append('}')
        elif ch == '[':
            stack.append(']')
        elif ch in ('}', ']') and stack and stack[-1] == ch:
            stack.pop()

    # Close open brackets in reverse order
    text += ''.join(reversed(stack))

    # Clean trailing commas before closing brackets
    text = re.sub(r',(\s*[}\]])', r'\1', text)

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return None


def _try_fix_json(content: str) -> Optional[Dict]:
    """Try to fix common JSON issues and parse."""
    import re

    # Try 0: Repair truncated JSON (most common failure mode)
    repaired = _repair_truncated_json(content)
    if repaired:
        return repaired

    # Try 1: Remove trailing commas before } or ]
    fixed = re.sub(r',(\s*[}\]])', r'\1', content)
    try:
        return json.loads(fixed)
    except:
        pass

    # Try 2: Fix unescaped quotes in strings (common LLM issue)
    try:
        return json.loads(content.replace('\\"', '"').replace('"', '\\"').replace('\\\\"', '"'))
    except:
        pass

    # Try 3: Extract just the units array if the outer structure is broken
    units_match = re.search(r'"units"\s*:\s*\[(.*)\]', content, re.DOTALL)
    if units_match:
        try:
            units_json = '[' + units_match.group(1) + ']'
            units_json = re.sub(r',(\s*[}\]])', r'\1', units_json)
            units = json.loads(units_json)
            return {"units": units}
        except:
            pass

    # Try 4: Use ast.literal_eval as fallback
    try:
        import ast
        return ast.literal_eval(content)
    except:
        pass

    return None


def _validate_lesson_structure(structure: Dict, subject: str, grade_level: str) -> Dict:
    """Validate and clean the lesson structure."""
    validated_units = []
    
    for unit in structure.get('units', []):
        unit_title = unit.get('title', '').strip()
        if not unit_title or len(unit_title) < 3:
            continue
        
        validated_lessons = []
        for lesson in unit.get('lessons', []):
            lesson_title = lesson.get('title', '').strip()
            if not lesson_title or len(lesson_title) < 3:
                continue
            
            # Clean up title if needed
            if len(lesson_title) > 60:
                lesson_title = lesson_title[:57] + "..."
            
            validated_lessons.append({
                "title": lesson_title,
                "objective": lesson.get('objective', lesson_title),
                "key_concepts": lesson.get('key_concepts', []),
            })
        
        if validated_lessons:
            validated_units.append({
                "title": unit_title,
                "description": unit.get('description', ''),
                "grade_level": unit.get('grade_level', ''),
                "lessons": validated_lessons,
            })
    
    return {
        "subject": subject,
        "grade_level": grade_level,
        "units": validated_units,
        "total_lessons": sum(len(u['lessons']) for u in validated_units),
    }


# ============================================================================
# STEP 4: GENERATE CONTENT (Tutoring Steps + Media)
# ============================================================================

def generate_lesson_content(
    lesson,  # Lesson model instance
    institution_id: int
) -> Dict:
    """
    Generate tutoring content for a lesson using curriculum context.
    
    This queries the knowledge base for:
    - Teaching strategies from the curriculum
    - Related content and examples
    - Assessment methods
    
    Then generates:
    - Tutoring steps (5E model or similar)
    - Media suggestions
    - Exit ticket questions
    
    Args:
        lesson: Lesson model instance
        institution_id: Institution ID
    
    Returns:
        Dict with generated content
    """
    from apps.curriculum.knowledge_base import CurriculumKnowledgeBase
    from apps.llm.models import ModelConfig
    from apps.llm.client import get_llm_client

    # Get unit info
    unit = lesson.unit
    course = unit.course
    subject = course.title.split()[0] if course else "General"

    # Query knowledge base for rich context
    kb = CurriculumKnowledgeBase(institution_id=institution_id)
    context = kb.query_for_content_generation(
        lesson_title=lesson.title,
        lesson_objective=lesson.objective or "",
        unit_title=unit.title,
        subject=subject,
        grade_level=course.grade_level if course else "S1"
    )

    # Build curriculum context for LLM
    curriculum_context = _build_curriculum_context(context)

    prompt = f"""Create a tutoring session for this lesson:

LESSON: {lesson.title}
OBJECTIVE: {lesson.objective or 'Not specified'}
UNIT: {unit.title}
SUBJECT: {subject}
GRADE: {course.grade_level if course else 'Secondary'}

CURRICULUM CONTEXT:
{curriculum_context}

Create a tutoring session with these phases:

1. ENGAGE (1-2 steps): Hook the student with a question or scenario
2. EXPLORE (2-3 steps): Guide discovery through examples
3. EXPLAIN (2-3 steps): Direct instruction with clear explanations
4. PRACTICE (2-4 steps): Practice questions with feedback
5. EXIT TICKET (2-3 questions): Check understanding

For each step, provide:
- step_type: "teach", "question", or "reflect"
- content: What the tutor says/shows
- question: (for question type) The question to ask
- expected_answer: (for question type) What correct answer looks like
- hints: (for question type) Array of hints if student struggles"""

    system_prompt = "You are an expert tutor creating engaging, pedagogically sound tutoring content."

    client, config = _get_instructor_client()
    if client:
        try:
            create_kwargs = dict(
                response_model=LessonContentResult,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
                max_retries=3,
            )
            if config and config.provider == 'google':
                create_kwargs['generation_config'] = {'max_tokens': 8192}
            else:
                create_kwargs['max_tokens'] = 8192

            result = client.chat.completions.create(**create_kwargs)
            return {
                "success": True,
                "lesson_id": lesson.id,
                "steps": [s.model_dump() for s in result.steps],
                "exit_ticket": [q.model_dump() for q in result.exit_ticket],
                "key_vocabulary": result.key_vocabulary,
                "teaching_tips": result.teaching_tips,
                "curriculum_context_used": len(context.chunks) > 0,
            }
        except Exception as e:
            logger.error(f"Instructor lesson content generation failed: {e}")
            return {
                "success": False,
                "error": str(e),
            }
    else:
        # Fallback: raw LLM + JSON repair
        logger.warning("Instructor unavailable for lesson content, using raw LLM")
        model_config = ModelConfig.get_for('generation')
        if not model_config:
            raise ValueError("No active LLM model configured")
        llm_client = get_llm_client(model_config)

        response = llm_client.generate(
            messages=[{"role": "user", "content": prompt + "\n\nReturn ONLY valid JSON."}],
            system_prompt=system_prompt,
            max_tokens=8192,
        )
        content = _clean_json_response(response.content.strip())
        try:
            parsed = json.loads(content)
        except json.JSONDecodeError as e:
            parsed = _try_fix_json(content)
            if parsed is None:
                return {
                    "success": False,
                    "error": str(e),
                    "raw_response": content[:500],
                }

        return {
            "success": True,
            "lesson_id": lesson.id,
            "steps": parsed.get('steps', []),
            "exit_ticket": parsed.get('exit_ticket', []),
            "key_vocabulary": parsed.get('key_vocabulary', []),
            "teaching_tips": parsed.get('teaching_tips', []),
            "curriculum_context_used": len(context.chunks) > 0,
        }


def _build_curriculum_context(context) -> str:
    """Build a context string from KB query results."""
    parts = []
    
    if context.teaching_strategies:
        parts.append("TEACHING STRATEGIES from curriculum:")
        for s in context.teaching_strategies[:5]:
            parts.append(f"- {s}")
    
    if context.objectives:
        parts.append("\nRELATED OBJECTIVES:")
        for o in context.objectives[:5]:
            parts.append(f"- {o}")
    
    if context.chunks:
        parts.append("\nRELEVANT CONTENT:")
        for chunk in context.chunks[:3]:
            section = chunk.get('section', '')
            content = chunk.get('content', '')[:300]
            parts.append(f"[{section}] {content}...")
    
    return "\n".join(parts) if parts else "No specific curriculum context available."


# ============================================================================
# STEP 5: TUTORING (Get live session context)
# ============================================================================

def get_tutoring_context(
    lesson,  # Lesson model instance
    student_message: str = None,
    current_step: int = None
) -> Dict:
    """
    Get curriculum context for a live tutoring session.
    
    Called by the TutorEngine to get:
    - Teaching strategies
    - Relevant curriculum content
    - Expected misconceptions
    - Scaffolding approaches
    
    Args:
        lesson: Current lesson
        student_message: Student's latest message
        current_step: Current step index
    
    Returns:
        Dict with context for the AI tutor
    """
    from apps.curriculum.knowledge_base import CurriculumKnowledgeBase
    
    # Get institution from lesson
    institution_id = None
    if hasattr(lesson, 'unit') and hasattr(lesson.unit, 'course'):
        course = lesson.unit.course
        if hasattr(course, 'institution_id'):
            institution_id = course.institution_id
        elif hasattr(course, 'institution'):
            institution_id = course.institution.id
    
    if not institution_id:
        # Return default context
        return _default_tutoring_context(lesson)
    
    # Query knowledge base
    kb = CurriculumKnowledgeBase(institution_id=institution_id)
    context = kb.query_for_tutoring(
        lesson=lesson,
        student_message=student_message
    )
    
    return {
        "teaching_strategies": context.teaching_strategies,
        "curriculum_objectives": context.objectives,
        "related_content": [c.get('content', '')[:200] for c in context.chunks[:3]],
        "context_summary": context.context_summary,
    }


def _default_tutoring_context(lesson) -> Dict:
    """Default context when KB is not available."""
    return {
        "teaching_strategies": [
            "Break down complex concepts into smaller steps",
            "Use concrete examples before abstract concepts",
            "Check for understanding frequently",
            "Provide positive reinforcement",
            "Connect to prior knowledge"
        ],
        "curriculum_objectives": [lesson.objective] if lesson.objective else [],
        "related_content": [],
        "context_summary": f"Teaching: {lesson.title}",
    }


# ============================================================================
# MAIN PIPELINE PROCESSOR
# ============================================================================

def process_curriculum_upload(upload_id: int, skip_review: bool = False) -> Dict:
    """
    Main curriculum processing pipeline.
    
    Steps:
    1. PARSE: Extract text
    2. VECTORIZE: Index into knowledge base
    3. GENERATE LESSONS: Create lesson structure (with review)
    4. CREATE RECORDS: Save to database
    
    Args:
        upload_id: CurriculumUpload record ID
        skip_review: Skip teacher review step
    
    Returns:
        Processing result
    """
    from apps.dashboard.models import CurriculumUpload
    
    upload = CurriculumUpload.objects.get(id=upload_id)
    
    try:
        upload.status = 'processing'
        upload.current_step = 1
        upload.processing_log = ""
        upload.add_log("🚀 Starting curriculum pipeline...")
        upload.save()
        
        from apps.accounts.models import Institution
        institution_id = upload.institution_id or Institution.get_global().id
        
        # ================================================================
        # STEP 1: PARSE
        # ================================================================
        upload.add_log("📄 Step 1: Extracting text from document...")
        upload.add_log(f"   File: {upload.file_path}")
        
        text, file_type = extract_text_from_file(upload.file_path)
        upload.extracted_text_length = len(text)
        upload.add_log(f"   ✓ Extracted {len(text):,} characters ({file_type})")
        
        if len(text) < 100:
            raise ValueError("Could not extract meaningful text. Document may be scanned images.")
        
        # Show preview
        preview = text[:300].replace('\n', ' ')
        upload.add_log(f"   Preview: {preview}...")
        upload.save()
        
        # ================================================================
        # STEP 2: VECTORIZE
        # ================================================================
        upload.current_step = 2
        upload.add_log("🔢 Step 2: Vectorizing curriculum into knowledge base...")
        upload.save()

        try:
            from apps.curriculum.knowledge_base import CurriculumKnowledgeBase
            kb = CurriculumKnowledgeBase(institution_id=institution_id)

            index_result = kb.index_curriculum_document(
                file_path=upload.file_path,
                subject=upload.subject_name,
                grade_level=upload.grade_level or 'S1',
                curriculum_upload_id=upload.id
            )

            upload.add_log(f"   ✓ Created {index_result.get('chunks_created', 0)} chunks")
            upload.add_log(f"   ✓ Indexed into vector database")
        except ImportError as e:
            upload.add_log(f"   ⚠️ Vector DB not available: {e}")
            upload.add_log("   Continuing without vectorization...")
        except Exception as e:
            upload.add_log(f"   ⚠️ Vectorization error: {e}")
            upload.add_log("   Continuing without vectorization...")

        upload.save()

        # ================================================================
        # STEP 3: GENERATE LESSONS
        # ================================================================
        upload.current_step = 3
        upload.add_log("📚 Step 3: Generating lesson structure...")
        upload.save()
        print(f"[Pipeline] Step 3: Generating lesson structure for {upload.subject_name} {upload.grade_level}", flush=True)

        try:
            structure = generate_lesson_structure(
                subject=upload.subject_name,
                grade_level=upload.grade_level or 'S1',
                institution_id=institution_id,
                extracted_text=text
            )
            print(f"[Pipeline] Step 3 complete: {len(structure.get('units',[]))} units", flush=True)
            
            units_count = len(structure.get('units', []))
            lessons_count = structure.get('total_lessons', 0)
            
            upload.add_log(f"   ✓ Found {units_count} units with {lessons_count} lessons")
            
            # Log unit details
            for unit in structure.get('units', [])[:5]:
                upload.add_log(f"      📁 {unit['title']}: {len(unit.get('lessons', []))} lessons")
            
            # Enrich with KB-extracted competencies (format-agnostic)
            try:
                competencies = extract_competencies_from_kb(
                    subject=upload.subject_name,
                    grade_level=upload.grade_level or 'S1',
                    institution_id=institution_id,
                )
                if competencies:
                    upload.add_log(f"   ✓ Extracted {len(competencies)} competencies from knowledge base")
                    # Merge competencies as enabling_objectives into units/lessons
                    _merge_competencies_into_structure(structure, competencies)
            except Exception as e:
                upload.add_log(f"   ⚠ Competency extraction skipped: {e}")

            upload.parsed_data = structure
            upload.save()
            
        except Exception as e:
            upload.add_log(f"   ❌ Lesson generation failed: {e}")
            upload.status = 'failed'
            upload.error_message = str(e)
            upload.save()
            raise
        
        # Check if we got any content
        if lessons_count == 0:
            upload.add_log("⚠️ No lessons extracted. Document format may not be recognized.")
            upload.status = 'review'
            upload.add_log("⏸️ Please review and provide feedback.")
            upload.save()
            
            return {
                'success': True,
                'status': 'review',
                'units_count': 0,
                'lessons_count': 0,
                'message': 'No lessons found. Please check document format.',
            }
        
        # ================================================================
        # PAUSE FOR REVIEW (unless skipped)
        # ================================================================
        if not skip_review:
            upload.status = 'review'
            upload.add_log("⏸️ Waiting for teacher review...")
            upload.save()
            
            return {
                'success': True,
                'status': 'review',
                'units_count': units_count,
                'lessons_count': lessons_count,
                'message': 'Please review the curriculum structure.',
            }
        
        # Skip to completion
        return complete_curriculum_upload(upload_id)
        
    except Exception as e:
        logger.exception(f"Pipeline failed: {e}")
        upload.status = 'failed'
        upload.error_message = str(e)
        upload.add_log(f"❌ Error: {e}")
        upload.save()
        raise


def complete_curriculum_upload(upload_id: int, feedback: str = "") -> Dict:
    """
    Complete the curriculum upload by creating database records.
    
    Called after teacher approves the structure.
    """
    from apps.dashboard.models import CurriculumUpload
    from apps.curriculum.models import Course, Unit, Lesson, LessonStep
    
    upload = CurriculumUpload.objects.get(id=upload_id)
    
    try:
        upload.status = 'processing'
        upload.current_step = 4
        
        if feedback:
            upload.teacher_feedback = feedback
        
        upload.add_log("💾 Step 4: Creating curriculum records...")
        upload.save()
        
        structure = upload.parsed_data
        if not structure:
            raise ValueError("No parsed data available")

        from apps.dashboard.models import TeachingMaterialUpload

        subject = structure.get('subject', upload.subject_name)
        grade_level = upload.grade_level or 'S1'

        # One upload = one grade level = one course
        course_title = f"{subject} {grade_level}"

        course, created = Course.objects.update_or_create(
            institution=upload.institution,
            title=course_title,
            defaults={
                'description': f"{subject} curriculum for {grade_level}",
                'grade_level': grade_level,
                'is_published': False,
            }
        )

        upload.created_course = course
        upload.add_log(f"   {'Created' if created else 'Updated'} course: {course.title}")

        # Link teaching materials to this course
        linked = TeachingMaterialUpload.objects.filter(
            curriculum_upload=upload, course__isnull=True
        ).update(course=course)
        if linked:
            upload.add_log(f"   Linked {linked} teaching material(s) to course")

        units_created = 0
        lessons_created = 0

        # Create units and lessons
        for unit_idx, unit_data in enumerate(structure.get('units', [])):
            unit, u_created = Unit.objects.update_or_create(
                course=course,
                title=unit_data['title'],
                defaults={
                    'description': unit_data.get('description', ''),
                    'order_index': unit_idx,
                    'grade_level': unit_data.get('grade_level', ''),
                    'terminal_objectives': unit_data.get('terminal_objectives', []),
                    'enabling_objectives': unit_data.get('enabling_objectives', []),
                }
            )
            
            if u_created:
                units_created += 1
            
            upload.add_log(f"   📁 {unit.title}")
            
            for lesson_idx, lesson_data in enumerate(unit_data.get('lessons', [])):
                lesson, l_created = Lesson.objects.update_or_create(
                    unit=unit,
                    title=lesson_data['title'],
                    defaults={
                        'objective': lesson_data.get('objective', ''),
                        'order_index': lesson_idx,
                        'estimated_minutes': 20,
                        'is_published': False,
                        'enabling_objectives': lesson_data.get('enabling_objectives', []),
                        'metadata': {
                            'key_concepts': lesson_data.get('key_concepts', []),
                            'teaching_steps': lesson_data.get('teaching_steps', []),
                            'from_curriculum_upload': upload.id,
                        }
                    }
                )
                
                if l_created:
                    lessons_created += 1
                    
                    # Create initial teaching step
                    LessonStep.objects.get_or_create(
                        lesson=lesson,
                        order_index=0,
                        defaults={
                            'step_type': 'teach',
                            'teacher_script': f"Today we will learn about: {lesson.objective or lesson.title}",
                        }
                    )
        
        upload.units_created = units_created
        upload.lessons_created = lessons_created
        upload.status = 'completed'
        upload.completed_at = timezone.now()
        upload.add_log(f"   ✓ Created {units_created} units, {lessons_created} lessons")
        upload.add_log(f"✅ Complete! Course '{course.title}' is ready.")
        upload.save()

        return {
            'success': True,
            'status': 'completed',
            'course_id': course.id,
            'course_name': course.title,
            'units_created': units_created,
            'lessons_created': lessons_created,
        }
        
    except Exception as e:
        logger.exception(f"Completion failed: {e}")
        upload.status = 'failed'
        upload.error_message = str(e)
        upload.add_log(f"❌ Error: {e}")
        upload.save()
        raise