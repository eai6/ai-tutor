"""
Seychelles Curriculum Parser

Extracts structured curriculum data from:
1. Mathematics Curriculum (text/markdown format)
2. Geography Syllabus (requires OCR - images)

This parser does NOT rely on AI for structure extraction.
It uses pattern matching to extract curriculum data directly.
"""

import re
import os
import json
import logging
from typing import Dict, List, Optional, Tuple, Literal
from dataclasses import dataclass, asdict

from pydantic import BaseModel, Field
from django.utils import timezone

logger = logging.getLogger(__name__)


# =============================================================================
# TYPED EXCEPTIONS
# =============================================================================

class OCRFailure(Exception):
    """Vision-OCR fallback failed in a classifiable way.

    Carries a stable `reason` slug (used by the materials pipeline to render
    actionable UI + drive routing), plus a free-text `detail` for the log.
    """

    REASONS = (
        'no_config',         # No active ModelConfig for the generation purpose
        'rate_limit',        # 429 / sustained backoff
        'auth',              # 401 / API key invalid
        'context_too_large', # 400 from provider — page batch exceeds context window
        'timeout',           # connection / request timeout
        'oversized_page',    # all pages still over MAX_IMAGE_BYTES after fallback
        'no_pages',          # zero renderable pages
        'empty_response',    # provider returned no text
        'all_providers_failed',  # every provider in the fallback chain failed
        'unknown',           # uncategorised — see detail
    )

    def __init__(self, reason: str, detail: str = ''):
        if reason not in self.REASONS:
            reason = 'unknown'
        self.reason = reason
        self.detail = detail
        super().__init__(f"OCR failed ({reason}): {detail}" if detail else f"OCR failed ({reason})")


# =============================================================================
# STRUCTURED OUTPUT SCHEMAS
# =============================================================================

class FigureDescription(BaseModel):
    """A single figure extracted from a PDF page."""
    page_number: int = Field(description="The page number where this figure appears")
    figure_number: str = Field(description="The figure label, e.g. 'Figure 3.2' or 'unlabeled'")
    figure_type: Literal["diagram", "chart", "graph", "map", "illustration", "photo", "table"] = Field(
        description="Type of visual element"
    )
    description: str = Field(description="Detailed 2-4 sentence description specific enough to recreate the figure")
    educational_context: str = Field(description="What concept this figure teaches or illustrates")


class FigureExtractionResult(BaseModel):
    """List of figures extracted from PDF pages."""
    figures: List[FigureDescription] = Field(
        default_factory=list,
        description="List of figures found. Empty if no figures on the pages.",
    )


# ============================================================================
# DATA CLASSES
# ============================================================================

@dataclass
class ParsedCurriculum:
    """Complete parsed curriculum."""
    subject: str
    grade_level: str
    cycle: str
    description: str
    units: List[Dict]
    teaching_strategies: List[str]
    assessment_methods: List[str]


# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def extract_text_from_file(file_path: str, progress_cb=None) -> Tuple[str, str]:
    """
    Extract text from curriculum file.

    Returns: (text, file_type)

    progress_cb is forwarded to extract_from_pdf for materials uploads that
    want per-batch progress updates from the vision-OCR fallback. Other file
    types ignore it (no streaming work to report on).
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext in ['.txt', '.md']:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return _strip_nul(text), 'text'

    elif ext == '.docx':
        return _strip_nul(extract_from_docx(file_path)), 'docx'

    elif ext == '.doc':
        # Legacy binary Word format (OLE2). python-docx only handles .docx
        # (zip-based OOXML) — reading the binary as text scatters NUL bytes
        # through the string and Postgres rejects it on insert. Fail fast
        # with an actionable message rather than silently producing garbage.
        raise ValueError(
            "Legacy .doc files are not supported (this is the binary Word "
            "format from Word 97-2003). Please convert to .docx or .pdf "
            "and re-upload — most word processors offer 'Save As → PDF' "
            "or 'Save As → Word Document (.docx)'."
        )

    elif ext == '.pdf':
        return _strip_nul(extract_from_pdf(file_path, progress_cb=progress_cb)), 'pdf'

    elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.tiff', '.tif']:
        return _strip_nul(extract_from_image(file_path)), 'image'

    else:
        # Try reading as text anyway
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return _strip_nul(text), 'text'
        except:
            raise ValueError(f"Unsupported file type: {ext}")


def _strip_nul(text: str) -> str:
    """Strip NUL (0x00) bytes from extracted text.

    PostgreSQL's `text` type forbids NUL bytes — they raise
    ``ValueError: A string literal cannot contain NUL (0x00) characters``
    on insert. Any text extracted from a binary file (legacy .doc, scanned
    PDF with embedded binary streams, etc.) may contain NUL. Strip them
    universally as a defense-in-depth measure even though specific
    extractors should be doing the right thing — the cost is one
    str.replace pass and the upside is that one missed extraction edge
    case doesn't blow up the whole pipeline.
    """
    if not text:
        return text
    return text.replace('\x00', '')


def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"python-docx failed: {e}, trying as text")
        # File might already be text (exported from Google Docs)
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


def extract_from_pdf(file_path: str, progress_cb=None) -> str:
    """Extract text from PDF, with multimodal LLM fallback for scanned docs.

    OCRFailure (typed) propagates so the materials pipeline can record a
    structured error. Generic exceptions are still swallowed → "" to
    preserve historical behavior for other callers (curriculum upload).

    Args:
        file_path: PDF path
        progress_cb: optional ``(pages_processed, pages_total, phase)`` callback
            forwarded to ``_extract_pdf_with_vision`` so materials uploads can
            update their `pages_processed` / `phase` fields per batch.
    """
    try:
        import fitz
        doc = fitz.open(file_path)

        # First pass: embedded text
        text = ""
        for page in doc:
            text += page.get_text()

        if len(text.strip()) >= 100:
            doc.close()
            return text

        # Multimodal LLM fallback
        logger.info(f"Low text ({len(text.strip())} chars), trying LLM vision: {file_path}")
        try:
            llm_text = _extract_pdf_with_vision(doc, progress_cb=progress_cb)
        except OCRFailure:
            doc.close()
            raise   # propagate typed failures so callers can render structured errors
        except Exception as e:
            logger.warning(f"LLM vision extraction failed (uncategorised): {e}")
            doc.close()
            return text
        doc.close()
        return llm_text if len(llm_text.strip()) > len(text.strip()) else text

    except OCRFailure:
        raise
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


# ============================================================================
# FIGURE EXTRACTION
# ============================================================================

def _has_meaningful_figures(page) -> bool:
    """
    Check if a PDF page contains meaningful figures (not tiny icons).

    Uses PyMuPDF to detect:
    - Embedded raster images larger than 100x100 pixels
    - Vector drawings with >10 operations (suggests diagram, not just borders)
    """
    # Check for embedded raster images (filter tiny icons < 100x100)
    for img_info in page.get_images(full=True):
        width, height = img_info[2], img_info[3]
        if width > 100 and height > 100:
            return True
    # Check for vector drawings (>10 ops suggests a diagram, not just borders)
    if len(page.get_drawings()) > 10:
        return True
    return False


def extract_figures_from_pdf(file_path: str, institution_id: int = None) -> List[Dict]:
    """
    Extract figure descriptions from a PDF using LLM vision.

    Only pages with meaningful figures (detected via PyMuPDF) are sent to the
    LLM for description, controlling cost.

    Args:
        file_path: Path to the PDF file
        institution_id: Optional institution ID for LLM config

    Returns:
        List of dicts with keys: page_number, figure_number, figure_type,
        description, educational_context, page_image_bytes, page_image_media_type
    """
    try:
        import fitz
    except ImportError:
        logger.warning("PyMuPDF (fitz) not installed, cannot extract figures")
        return []

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        logger.error(f"Could not open PDF for figure extraction: {e}")
        return []

    # Identify pages with meaningful figures
    MAX_IMAGE_BYTES = 4_500_000  # Stay under Anthropic's 5MB limit
    pages_with_figures = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        if _has_meaningful_figures(page):
            # Render page at 100 DPI for vision analysis
            pix = page.get_pixmap(dpi=100)
            image_bytes = pix.tobytes("png")
            media_type = 'image/png'

            # If PNG exceeds size limit, fall back to JPEG
            if len(image_bytes) > MAX_IMAGE_BYTES:
                image_bytes = pix.tobytes("jpeg")
                media_type = 'image/jpeg'

            # If still too large, reduce DPI
            if len(image_bytes) > MAX_IMAGE_BYTES:
                pix = page.get_pixmap(dpi=72)
                image_bytes = pix.tobytes("jpeg")
                media_type = 'image/jpeg'

            if len(image_bytes) > MAX_IMAGE_BYTES:
                logger.warning(f"Page {page_num + 1} image still too large ({len(image_bytes)} bytes), skipping")
                continue

            pages_with_figures.append({
                'page_number': page_num + 1,
                'image_bytes': image_bytes,
                'media_type': media_type,
            })

    doc.close()

    if not pages_with_figures:
        logger.info(f"No meaningful figures found in {file_path}")
        return []

    logger.info(f"Found {len(pages_with_figures)} pages with figures in {file_path}")

    # Process one page at a time with delay to respect rate limits (30k tokens/min)
    import time
    all_figures = []
    for i, page_data in enumerate(pages_with_figures):
        max_retries = 3
        for attempt in range(max_retries):
            try:
                batch_figures = _batch_extract_figures_with_vision([page_data])
                all_figures.extend(batch_figures)
                # Pause between pages to stay under rate limits
                if i < len(pages_with_figures) - 1:
                    time.sleep(5)
                break
            except Exception as e:
                if '429' in str(e) and attempt < max_retries - 1:
                    wait = 30 * (attempt + 1)
                    logger.warning(f"Rate limited on page {page_data['page_number']}, waiting {wait}s...")
                    time.sleep(wait)
                else:
                    logger.error(f"Figure extraction failed for page {page_data['page_number']}: {e}")
                    break

    logger.info(f"Extracted {len(all_figures)} figure descriptions from {file_path}")
    return all_figures


def _batch_extract_figures_with_vision(pages_data: List[Dict]) -> List[Dict]:
    """
    Use LLM vision to extract figure descriptions from rendered page images.

    Uses instructor for structured output with Anthropic Haiku.

    Args:
        pages_data: List of dicts with page_number, image_bytes, media_type

    Returns:
        List of figure dicts with descriptions and metadata
    """
    import base64
    import instructor
    import anthropic

    # Use Haiku directly for figure extraction — cheapest, fastest, 50k token/min
    # limit (vs 30k for Sonnet/Opus). Fully multimodal and good enough for
    # describing figures. We bypass ModelConfig to avoid needing a DB entry.
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY not set for figure extraction")

    client = instructor.from_anthropic(anthropic.Anthropic(api_key=api_key))

    # Build multimodal message with all page images
    content_parts = [
        {
            "type": "text",
            "text": (
                "Analyze the following PDF page images and extract ALL figures, diagrams, "
                "charts, maps, illustrations, and visual elements (NOT decorative borders or "
                "page numbers).\n\n"
                "figure_type must be one of: diagram, chart, graph, map, illustration, photo, table\n"
                "If a figure has no label, use 'unlabeled' for figure_number.\n"
                "If no figures are found on a page, omit that page from results."
            ),
        }
    ]

    for page_data in pages_data:
        b64_image = base64.b64encode(page_data['image_bytes']).decode('utf-8')
        content_parts.append({
            "type": "text",
            "text": f"\n--- Page {page_data['page_number']} ---",
        })
        content_parts.append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": page_data['media_type'],
                "data": b64_image,
            },
        })

    result = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4096,
        system="You are an expert at analyzing educational documents. Extract figure descriptions precisely.",
        messages=[{"role": "user", "content": content_parts}],
        response_model=FigureExtractionResult,
        max_retries=2,
    )

    # Attach page image bytes to each figure for later storage
    page_data_by_num = {p['page_number']: p for p in pages_data}
    results = []
    for fig in result.figures:
        fig_dict = fig.model_dump()
        page_num = fig_dict.get('page_number')
        page_info = page_data_by_num.get(page_num)
        if page_info:
            fig_dict['page_image_bytes'] = page_info['image_bytes']
            fig_dict['page_image_media_type'] = page_info['media_type']
        results.append(fig_dict)

    return results


def _classify_llm_error(exc: Exception) -> str:
    """Map a provider exception to an OCRFailure.REASONS slug.

    Pattern-matches by exception class name + message — avoids importing
    every provider SDK at the top of the parser. Returns 'unknown' for
    anything we haven't seen yet so the detail string can carry context.
    """
    name = type(exc).__name__
    msg = str(exc).lower()
    if 'ratelimit' in name.lower() or '429' in msg or 'rate limit' in msg:
        return 'rate_limit'
    if 'auth' in name.lower() or '401' in msg or 'invalid api key' in msg or 'unauthorized' in msg:
        return 'auth'
    if 'timeout' in name.lower() or 'timed out' in msg:
        return 'timeout'
    if 'context' in msg and ('length' in msg or 'window' in msg or 'limit' in msg):
        return 'context_too_large'
    if '400' in msg and 'too large' in msg:
        return 'context_too_large'
    return 'unknown'


_VISION_MAX_IMAGE_BYTES = 4_500_000  # stay under Anthropic's 5MB image limit
_VISION_BATCH_SIZE = 10              # pages per LLM call
_VISION_MAX_WORKERS = 5              # concurrent in-flight batches


def _render_page_for_vision(page) -> Optional[Tuple[str, str]]:
    """Render one PyMuPDF page → (base64_str, media_type) or None if oversized.

    Cascades 200-DPI PNG → 200-DPI JPEG → 120-DPI JPEG, dropping the page only
    if even the smallest variant exceeds the 4.5 MB Anthropic image limit.
    """
    import base64
    pix = page.get_pixmap(dpi=200)
    img_bytes = pix.tobytes("png")
    media_type = "image/png"
    if len(img_bytes) > _VISION_MAX_IMAGE_BYTES:
        img_bytes = pix.tobytes("jpeg", jpg_quality=80)
        media_type = "image/jpeg"
    if len(img_bytes) > _VISION_MAX_IMAGE_BYTES:
        pix = page.get_pixmap(dpi=120)
        img_bytes = pix.tobytes("jpeg", jpg_quality=80)
    if len(img_bytes) > _VISION_MAX_IMAGE_BYTES:
        return None
    return base64.b64encode(img_bytes).decode("utf-8"), media_type


def _extract_pdf_with_vision(
    doc,
    progress_cb=None,
    start_page: int = 0,
):
    """
    Render PDF pages to images and use multimodal LLM to extract content.

    Provider strategy: uses ``apps.curriculum.vision_ocr`` for pluggable
    multi-provider fallback. Per batch:
      1. Try the primary provider (active ModelConfig for purpose='generation')
      2. On any failure, fall through to the next provider in the chain
         (typically a different vendor — e.g. Gemini → Claude → GPT)
      3. If ALL providers fail for that batch, mark the batch skipped
         and continue with the rest. The rest of the document still gets
         indexed; one bad batch doesn't kill the whole run.

    Streams page rendering (one batch at a time) and dispatches batches
    concurrently (up to _VISION_MAX_WORKERS in flight) so memory stays bounded
    AND wall-clock drops ~5× vs sequential.

    Args:
        doc: PyMuPDF document handle (caller owns close())
        progress_cb: Optional callback ``(pages_processed, pages_total, phase)``
            invoked after each batch completes.
        start_page: Resume hint (P3). Pages [0, start_page) are skipped.

    Raises:
        OCRFailure: only when EVERY batch failed across EVERY provider, OR
            when there are no providers configured, OR when there are no
            renderable pages. Partial success → returns concatenated text.
    """
    from concurrent.futures import ThreadPoolExecutor, FIRST_COMPLETED, wait
    from apps.curriculum.vision_ocr import (
        RenderedPage,
        get_vision_provider_chain,
        extract_text_with_fallback,
    )

    providers = get_vision_provider_chain()
    if not providers:
        raise OCRFailure(
            'no_config',
            "No active ModelConfig found across purposes (generation, judge, "
            "tutoring, exit_tickets) — cannot run vision OCR.",
        )

    pages_total = len(doc)
    if pages_total == 0:
        raise OCRFailure('no_pages', "PDF rendered zero pages.")

    if start_page >= pages_total:
        # Caller has already processed everything — nothing to do.
        return ""

    system_prompt = (
        "You are a document text extraction assistant. "
        "Extract ALL text, labels, titles, and describe visual elements "
        "(maps, diagrams, charts) from the provided document pages. "
        "Return only the extracted content, no commentary."
    )
    extraction_prompt = (
        "Extract ALL text, labels, titles, and describe visual elements "
        "(maps, diagrams, charts) from these document pages."
    )

    # Build batch boundaries [start, end) over the unprocessed page range.
    batches = [
        (i, min(i + _VISION_BATCH_SIZE, pages_total))
        for i in range(start_page, pages_total, _VISION_BATCH_SIZE)
    ]
    total_batches = len(batches)
    results: List[Optional[str]] = [None] * total_batches
    skipped_batches: List[Tuple[int, str]] = []   # [(batch_idx, reason)]
    pages_done = start_page
    skipped_oversized_total = 0

    def _process_one_batch(batch_idx: int, pages: List[RenderedPage]) -> Tuple[int, Optional[str], Optional[str]]:
        """Run the provider chain on one batch. Runs in worker thread.

        Returns (batch_idx, text_or_None, error_reason_or_None).
        Never raises — the orchestrator decides skip-vs-fail at the end.
        """
        if not pages:
            return batch_idx, "", None   # whole batch was oversized; treat as empty

        result = extract_text_with_fallback(
            pages=pages,
            providers=providers,
            system_prompt=system_prompt,
            extraction_prompt=extraction_prompt,
        )
        if result.success:
            return batch_idx, result.text, None
        # All providers failed for this batch
        start, end = batches[batch_idx]
        reason = (
            f"batch {batch_idx + 1}/{total_batches} (pages {start + 1}-{end}): "
            f"all providers failed; last={result.provider}/{result.model_name} "
            f"({result.error_reason}: {result.error_detail[:200]})"
        )
        return batch_idx, None, reason

    def _render_batch(batch_idx: int) -> Tuple[List[RenderedPage], int]:
        """Render pages for batch_idx (main thread; PyMuPDF not thread-safe)."""
        start, end = batches[batch_idx]
        rendered: List[RenderedPage] = []
        skipped = 0
        for page_idx in range(start, end):
            r = _render_page_for_vision(doc[page_idx])
            if r is None:
                logger.warning(
                    f"Page {page_idx + 1} oversized after fallback compression — skipping"
                )
                skipped += 1
            else:
                rendered.append(RenderedPage(b64=r[0], media_type=r[1]))
        return rendered, skipped

    # Sliding window: keep at most _VISION_MAX_WORKERS batches in flight.
    with ThreadPoolExecutor(max_workers=_VISION_MAX_WORKERS) as executor:
        in_flight: Dict[object, Tuple[int, int]] = {}   # future -> (batch_idx, batch_pages)
        next_to_submit = 0

        def _submit_next():
            nonlocal next_to_submit, skipped_oversized_total
            while next_to_submit < total_batches and len(in_flight) < _VISION_MAX_WORKERS:
                bi = next_to_submit
                rendered, skipped = _render_batch(bi)
                skipped_oversized_total += skipped
                start, end = batches[bi]
                fut = executor.submit(_process_one_batch, bi, rendered)
                in_flight[fut] = (bi, end - start)
                next_to_submit += 1

        _submit_next()

        while in_flight:
            done, _pending = wait(list(in_flight.keys()), return_when=FIRST_COMPLETED)
            for done_fut in done:
                bi, batch_pages = in_flight.pop(done_fut)
                try:
                    _bi, text, error_reason = done_fut.result()
                except Exception as exc:
                    # _process_one_batch is no-raise by design; this would
                    # be an orchestrator bug, not a provider failure.
                    skipped_batches.append((bi, f"orchestrator-bug: {exc!r}"))
                    text = None
                else:
                    if text is None:
                        skipped_batches.append((bi, error_reason or 'unknown'))
                        logger.warning(f"OCR batch skipped: {error_reason}")
                    else:
                        results[bi] = text
                pages_done += batch_pages
                if progress_cb is not None:
                    try:
                        progress_cb(pages_done, pages_total, f"vision_ocr_p{pages_done}_of_{pages_total}")
                    except Exception as cb_exc:
                        logger.warning(f"progress_cb raised — ignoring: {cb_exc}")
            _submit_next()

    # Decide outcome:
    #   - At least one batch succeeded → return concatenated text +
    #     log the skipped count so the materials UI can surface it
    #   - Every batch failed → raise with the most informative reason
    successful = [r for r in results if r and r.strip()]
    if successful:
        if skipped_batches:
            logger.warning(
                "Vision OCR completed with %d/%d batches skipped after all "
                "providers failed: %s",
                len(skipped_batches), total_batches,
                '; '.join(f"#{bi+1}" for bi, _ in skipped_batches[:5]),
            )
        return "\n\n".join(r for r in results if r)

    # Nothing succeeded — figure out the best reason to surface
    if skipped_oversized_total == sum(end - start for start, end in batches):
        raise OCRFailure(
            'oversized_page',
            f"All {skipped_oversized_total} page(s) exceeded {_VISION_MAX_IMAGE_BYTES} bytes "
            f"after fallback to 120 DPI JPEG.",
        )
    if skipped_batches:
        # Use the first skipped batch's reason as the headline; it's
        # already prefixed with "batch N/M (pages X-Y): all providers failed; last=..."
        raise OCRFailure('all_providers_failed', skipped_batches[0][1])
    raise OCRFailure(
        'empty_response',
        f"Vision LLM returned no text across {total_batches} batch(es).",
    )


def extract_curriculum_with_vision(file_path: str, subject: str, grade_level: str) -> Optional['ParsedCurriculum']:
    """
    Use LLM vision to extract complete curriculum structure from a PDF.

    Renders each page to an image and sends batches to a multimodal LLM,
    asking it to extract units, terminal objectives, enabling objectives,
    teaching strategies, and lesson structure — reading the actual formatted
    document as a human would.

    This is more accurate than regex-based text extraction because it
    preserves numbering, table structure, and multi-column layouts.
    """
    import base64
    import fitz
    from apps.llm.models import ModelConfig
    from apps.llm.client import get_llm_client

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        logger.error(f"Could not open PDF for vision extraction: {e}")
        return None

    config = ModelConfig.get_for('generation')
    if not config:
        logger.warning("No LLM model configured for vision extraction")
        return None

    client = get_llm_client(config)
    is_anthropic = config.provider == 'anthropic'

    # Render pages to images
    MAX_IMAGE_BYTES = 4_500_000
    page_images = []
    for page_num, page in enumerate(doc):
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("jpeg", jpg_quality=85)
        media_type = "image/jpeg"

        if len(img_bytes) > MAX_IMAGE_BYTES:
            pix = page.get_pixmap(dpi=100)
            img_bytes = pix.tobytes("jpeg", jpg_quality=75)

        if len(img_bytes) > MAX_IMAGE_BYTES:
            continue

        page_images.append({
            'page_num': page_num + 1,
            'b64': base64.b64encode(img_bytes).decode('utf-8'),
            'media_type': media_type,
        })

    if not page_images:
        return None

    is_math = subject.lower() in ('mathematics', 'math', 'maths')
    print(f"[VisionExtractor] {len(page_images)} pages rendered for {subject} {grade_level} (math={is_math})", flush=True)

    # Process pages in batches — extract structure from each batch
    all_units = []
    batch_size = 8  # 8 pages per LLM call

    system_prompt = (
        "You are a curriculum analysis expert. You analyze educational documents "
        "and extract their complete structure with perfect accuracy. "
        "You MUST extract every numbered item, every bullet point, every table entry. "
        "Return ONLY valid JSON, no explanation."
    )

    for batch_start in range(0, len(page_images), batch_size):
        batch = page_images[batch_start:batch_start + batch_size]
        page_range = f"pages {batch[0]['page_num']}-{batch[-1]['page_num']}"
        print(f"[VisionExtractor] Processing {page_range}...", flush=True)

        is_math = subject.lower() in ('mathematics', 'math', 'maths')

        if is_math:
            extraction_prompt = f"""Analyze these MATHEMATICS curriculum pages and extract the complete structure.

SUBJECT: {subject}

This is a math termly plan organized by TOPICS and WEEKS. Each topic has TWO columns:
- CORE (SET 3+): The foundational objectives every student must achieve
- EXTENDED (SETS 1&2): The deeper/harder objectives for higher-ability students

For EACH topic on these pages, extract:
1. unit_title: The strand code and topic name (e.g., "Measures (M4) - Metric Measures")
2. grade_level: The secondary level (S1, S2, S3, etc.)
3. terminal_objectives: The CORE (SET 3+) objectives — extract EVERY bullet point from the Core column
4. enabling_objectives: The EXTENDED (SETS 1&2) objectives — extract EVERY bullet point from the Extended column
5. teaching_strategies: Assessment methods listed
6. resources: Textbooks and materials listed
7. weeks: How many weeks allocated (e.g., "2 weeks")

IMPORTANT: Keep Core and Extended SEPARATE:
- terminal_objectives = CORE column only
- enabling_objectives = EXTENDED column only
Do NOT merge them. Do NOT skip any bullet points.

Return a JSON array:
[{{
    "unit_title": "Measures (M4) - Metric Measures",
    "grade_level": "S3",
    "weeks": "2 weeks",
    "terminal_objectives": [
        "List units for area",
        "Estimate area",
        "List units of volume",
        "Estimate volume and capacity",
        "Convert units of volume and capacity",
        "Solve simple problems involving the area, volume and capacity"
    ],
    "enabling_objectives": [
        "List units for area",
        "Convert units of area",
        "Convert units of volume and capacity",
        "Convert units and volume",
        "Solve problems involving area, volume, capacity and conversion of units of measures"
    ],
    "teaching_strategies": ["verbal questions", "class discussion", "end of topic tests"],
    "resources": ["Maths In Action Bk 2", "Complete Mathematics for Cambridge Secondary 1"]
}}]

Return ONLY valid JSON."""
        else:
            extraction_prompt = f"""Analyze these curriculum document pages and extract ALL units with their complete structure.

SUBJECT: {subject}

Extract units from ALL grade levels visible on these pages. For each unit, include a "grade_level" field indicating which secondary level it belongs to (look for "Secondary One" = "S1", "Secondary Two" = "S2", "Secondary Three" = "S3", etc. Also "Cycle 4" covers S1-S3, "Cycle 5/IGCSE" covers S4-S5).

For EACH unit you find on these pages, extract:
1. unit_title: The exact unit title (e.g., "Development and Trade")
2. terminal_objectives: ALL numbered terminal objectives — extract EVERY single one, preserving exact text
3. enabling_objectives: ALL bullet-point enabling objectives from the Teaching/Learning Scheme table — extract EVERY one
4. teaching_strategies: Methods listed in the Teaching/Learning Strategies column
5. resources: Resources listed in the Resources column
6. assessment: Assessment methods listed

Return a JSON array of units:
[{{
    "unit_title": "Unit 16: Development and Trade",
    "grade_level": "S3",
    "terminal_objectives": [
        "Know the terminologies associated with development",
        "Develop an understanding of the world pattern of development",
        ...every single one...
    ],
    "enabling_objectives": [
        "Define the terms Development, Globalization, MEDC, NIC, LEDC",
        "Describe how the influence can give an indication...",
        ...every single one from the table...
    ],
    "teaching_strategies": ["Discussions", "Written activities", ...],
    "resources": ["The New Wider World", "Geography in Place 1", ...],
    "assessment": ["Structured source-based written questions", ...]
}}]

CRITICAL: Do NOT summarize or skip any objectives. Extract the EXACT text of every numbered item and every bullet point. If a terminal objective says "10. Understand how world trade works" extract "Understand how world trade works" (without the number).

Return ONLY valid JSON."""

        # Build multimodal message — images FIRST, then text prompt (Anthropic format)
        content_blocks = []
        for pg in batch:
            if is_anthropic:
                content_blocks.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": pg['media_type'],
                        "data": pg['b64'],
                    }
                })
            else:
                content_blocks.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{pg['media_type']};base64,{pg['b64']}"},
                })
        content_blocks.append({"type": "text", "text": extraction_prompt})

        try:
            response = client.generate(
                messages=[{"role": "user", "content": content_blocks}],
                system_prompt=system_prompt,
                max_tokens=8000,
            )

            # Parse JSON response
            from apps.llm.json_utils import parse_llm_json
            units = parse_llm_json(response.content, expect_array=True)
            if units and isinstance(units, list):
                all_units.extend(units)
                print(f"[VisionExtractor] {page_range}: extracted {len(units)} units", flush=True)
            else:
                print(f"[VisionExtractor] {page_range}: no units found in response", flush=True)

        except Exception as e:
            print(f"[VisionExtractor] {page_range} failed: {e}", flush=True)
            continue

    if not all_units:
        return None

    # Filter to target grade level
    # Deduplicate units by title
    import re as _re
    seen = {}
    deduped = []
    for u in all_units:
        title = u.get('unit_title', '')
        clean = _re.sub(r'[^a-z0-9 ]', '', title.lower())[:35]
        if clean in seen:
            # Merge objectives
            existing = seen[clean]
            for to in u.get('terminal_objectives', []):
                if to not in existing.get('terminal_objectives', []):
                    existing.setdefault('terminal_objectives', []).append(to)
            for eo in u.get('enabling_objectives', []):
                if eo not in existing.get('enabling_objectives', []):
                    existing.setdefault('enabling_objectives', []).append(eo)
        else:
            seen[clean] = u
            deduped.append(u)

    # Filter to requested grade levels, then build ParsedCurriculum
    from apps.curriculum.utils import determine_cycles, parse_grade_level_string
    requested_grades = parse_grade_level_string(grade_level)
    cycles = determine_cycles(grade_level)
    cycle = "/".join(cycles)

    # Build set of acceptable grade labels for filtering
    # e.g. grade_level="S3" → accept S3, Cycle 4
    # e.g. grade_level="S1,S2,S3,S4,S5" → accept all
    accept_grades = set(g.upper() for g in requested_grades) if requested_grades else None
    accept_cycles = set(cycles)  # e.g. {'4'} or {'4','5'}

    def _grade_matches(unit_grade_str: str) -> bool:
        """Check if a unit's grade matches the requested grades."""
        if not accept_grades:
            return True  # No filter — accept all
        if not unit_grade_str:
            return True  # No grade tagged — assume it matches

        ug = unit_grade_str.upper().strip()
        # Direct match: "S3" in {"S3"}
        if ug in accept_grades:
            return True
        # Cycle match: "Cycle 4" with accept_cycles={'4'}
        import re as _re2
        cycle_match = _re2.search(r'cycle\s*(\d+)', ug, _re2.IGNORECASE)
        if cycle_match and cycle_match.group(1) in accept_cycles:
            return True
        # IGCSE match: only if Cycle 5 is accepted
        if 'igcse' in ug.lower() and '5' in accept_cycles:
            return True
        # Multi-grade label: "S1-S3" or "S1,S2,S3"
        unit_grades = parse_grade_level_string(unit_grade_str)
        if unit_grades and accept_grades.intersection(g.upper() for g in unit_grades):
            return True
        return False

    filtered = [u for u in deduped if _grade_matches(u.get('grade_level', ''))]
    print(f"[VisionExtractor] Grade filter: {len(deduped)} total → {len(filtered)} matching {grade_level}", flush=True)

    units = []
    for u in filtered:
        title = u.get('unit_title', 'Untitled Unit')
        unit_grade = u.get('grade_level', grade_level) or grade_level
        # Remove "Unit N:" prefix for cleaner titles
        title = _re.sub(r'^Unit\s+\d+\s*:\s*', '', title).strip()

        tos = u.get('terminal_objectives', [])
        eos = u.get('enabling_objectives', [])

        # Unified teaching-objectives rule (2026-04-27): one lesson per
        # teaching objective. We flatten the unit's terminal + enabling
        # objectives into a single ordered, deduplicated list — TOs first
        # (broad outcomes), then EOs (granular skills). Each item becomes
        # one 20-minute lesson that drills that single objective intensely.
        seen = set()
        teaching_objectives = []
        for obj in (tos or []) + (eos or []):
            if not obj:
                continue
            key = ' '.join(str(obj).split()).lower()
            if key in seen:
                continue
            seen.add(key)
            teaching_objectives.append(str(obj).strip())

        lessons = []
        for objective in teaching_objectives:
            lessons.append({
                'title': create_lesson_title(objective),
                'objective': objective,
                # Each lesson owns exactly ONE teaching objective. Stored
                # in `enabling_objectives` for backward compatibility with
                # the rest of the pipeline (content generator, exit-ticket
                # generator, etc.).
                'enabling_objectives': [objective],
                'teaching_steps': [],
                'teaching_strategies': u.get('teaching_strategies', []),
                'resources': u.get('resources', []),
                'assessment_methods': u.get('assessment', []),
                'order': len(lessons) + 1,
            })

        units.append({
            'number': len(units) + 1,
            'title': title,
            'grade_level': unit_grade,
            'duration': '',
            'introduction': '',
            'terminal_objectives': tos,
            'enabling_objectives': eos,
            'lessons': lessons,
        })

    print(f"[VisionExtractor] Final: {len(units)} units, "
          f"{sum(len(u['lessons']) for u in units)} lessons, "
          f"{sum(len(u['terminal_objectives']) for u in units)} TOs, "
          f"{sum(len(u['enabling_objectives']) for u in units)} EOs", flush=True)

    return ParsedCurriculum(
        subject=subject,
        grade_level=grade_level,
        cycle=cycle,
        description=f"{subject} curriculum for {grade_level} (extracted via LLM vision)",
        units=units,
        teaching_strategies=list(set(
            s for u in deduped for s in u.get('teaching_strategies', [])
        ))[:10],
        assessment_methods=list(set(
            a for u in deduped for a in u.get('assessment', [])
        ))[:10],
    )


def extract_from_image(file_path: str) -> str:
    """Extract text/content from an image file using multimodal LLM."""
    import base64
    from apps.llm.models import ModelConfig
    from apps.llm.client import get_llm_client

    config = ModelConfig.get_for('generation')
    if not config:
        logger.warning("No active LLM configured, cannot extract from image")
        return ""

    # Determine media type from extension
    ext = os.path.splitext(file_path)[1].lower()
    media_types = {
        '.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.gif': 'image/gif', '.webp': 'image/webp', '.bmp': 'image/bmp',
        '.tiff': 'image/tiff', '.tif': 'image/tiff',
    }
    media_type = media_types.get(ext, 'image/png')

    with open(file_path, 'rb') as f:
        img_bytes = f.read()

    # Downsize if over 4.5MB (Anthropic limit is 5MB)
    if len(img_bytes) > 4_500_000:
        try:
            from PIL import Image
            import io
            img = Image.open(io.BytesIO(img_bytes))
            # Scale down to ~75% until under limit
            while len(img_bytes) > 4_500_000:
                new_size = (int(img.width * 0.75), int(img.height * 0.75))
                img = img.resize(new_size, Image.LANCZOS)
                buf = io.BytesIO()
                img.save(buf, format='JPEG', quality=85)
                img_bytes = buf.getvalue()
                media_type = 'image/jpeg'
        except ImportError:
            logger.warning("Pillow not installed, cannot resize large image")
            return ""

    img_b64 = base64.b64encode(img_bytes).decode('utf-8')

    is_anthropic = config.provider == ModelConfig.Provider.ANTHROPIC

    if is_anthropic:
        image_block = {
            "type": "image",
            "source": {"type": "base64", "media_type": media_type, "data": img_b64},
        }
    else:
        image_block = {
            "type": "image_url",
            "image_url": {"url": f"data:{media_type};base64,{img_b64}"},
        }

    messages = [{"role": "user", "content": [
        image_block,
        {"type": "text", "text": (
            "Extract ALL text, labels, titles, and describe visual elements "
            "(maps, diagrams, charts) from this image."
        )},
    ]}]

    client = get_llm_client(config)

    try:
        response = client.generate(
            messages=messages,
            system_prompt=(
                "You are a document text extraction assistant. "
                "Extract ALL text, labels, titles, and describe visual elements. "
                "Return only the extracted content, no commentary."
            ),
            max_tokens=4096,
        )
        return response.content
    except Exception as e:
        logger.warning(f"LLM image extraction failed: {e}")
        return ""


# ============================================================================
# DETECT SUBJECT TYPE
# ============================================================================

def detect_subject(text: str, provided_subject: str = "") -> str:
    """Detect the subject from text content using weighted keyword scoring.

    More robust than simple keyword matching — uses multiple keywords per subject
    with scoring to handle documents that mention multiple subjects.
    """
    if provided_subject:
        return provided_subject

    text_lower = text[:10000].lower()  # Only scan first 10k chars for efficiency

    # Weighted keyword sets — score-based detection handles ambiguous documents
    subject_keywords = {
        'Mathematics': [
            'mathematics', 'algebra', 'arithmetic', 'equation', 'fraction',
            'geometry', 'trigonometry', 'calculus', 'polynomial', 'quadratic',
            'integer', 'decimal', 'percentage', 'multiplication', 'division',
            'perimeter', 'pythagoras', 'histogram', 'probability',
        ],
        'Geography': [
            'geography', 'map skills', 'settlement', 'population', 'climate',
            'topographic', 'contour', 'erosion', 'tectonic', 'latitude',
            'longitude', 'hemisphere', 'migration', 'urbanisation', 'development',
        ],
        'Biology': [
            'biology', 'organism', 'cell', 'photosynthesis', 'ecosystem',
            'genetics', 'evolution', 'anatomy', 'respiration', 'reproduction',
        ],
        'Physics': [
            'physics', 'mechanics', 'velocity', 'acceleration', 'force',
            'energy', 'electricity', 'magnetism', 'waves', 'thermodynamics',
        ],
        'Chemistry': [
            'chemistry', 'element', 'compound', 'molecule', 'reaction',
            'acid', 'periodic table', 'bonding', 'oxidation', 'titration',
        ],
        'Science': [
            'science', 'scientific method', 'hypothesis', 'experiment',
            'observation', 'laboratory',
        ],
    }

    scores = {}
    for subject, keywords in subject_keywords.items():
        scores[subject] = sum(1 for kw in keywords if kw in text_lower)

    best_subject = max(scores, key=scores.get)
    if scores[best_subject] >= 2:
        return best_subject
    return 'General'


# ============================================================================
# MATHEMATICS CURRICULUM PARSER
# ============================================================================

def parse_mathematics_curriculum(text: str, grade_level: str = "S1") -> ParsedCurriculum:
    """
    Parse Seychelles Mathematics curriculum text.

    The curriculum is organized by:
    - **Knowledge** (K-codes): Concept understanding per strand per cycle
    - **Skills** (S-codes): Computation, Measuring, Communicating, Reasoning, Problem Solving
    - **Attitudes** (A-codes): Learning behaviors
    - **Terminal Objectives**: By cycle, organized by topic
    - **Scope & Sequence**: Topics per strand per cycle

    Output is organized by **strand** (as Units) with **sub-strand topic groups** (as Lessons).
    Each lesson gets enabling_objectives from K/S/A codes + terminal objectives.
    """
    import re as _re
    from apps.curriculum.utils import determine_cycles
    cycles = determine_cycles(grade_level)
    cycle = "/".join(cycles)

    # Map grade level to cycle column for K/S/A extraction
    # Cycle 4 = S1-S2, Cycle 5 = S3-S5
    target_cycles = set()
    gl = grade_level.upper()
    if gl in ('S1', 'S2'):
        target_cycles = {'4', 'CYCLE 4'}
    elif gl in ('S3', 'S4', 'S5'):
        target_cycles = {'5', 'CYCLE 5'}
    else:
        target_cycles = {'4', '5', 'CYCLE 4', 'CYCLE 5'}

    # ── Step 1: Extract coded objectives (K/S/A codes) ──
    coded_objectives = {'K': [], 'S': [], 'A': []}
    lines = text.split('\n')
    for line in lines:
        clean = line.strip()
        # Match patterns like "K408 ...", "S401 ...", "A301 ..."
        code_match = _re.match(r'^([KSA])(\d)(\d{2})\s+(.+)', clean)
        if code_match:
            prefix = code_match.group(1)      # K, S, or A
            cycle_num = code_match.group(2)    # 1-5
            obj_text = code_match.group(4).strip().rstrip('.')
            full_code = f"{prefix}{code_match.group(2)}{code_match.group(3)}"
            if cycle_num in target_cycles or f'CYCLE {cycle_num}' in target_cycles:
                if len(obj_text) > 10:
                    coded_objectives[prefix].append({
                        'code': full_code,
                        'text': obj_text,
                        'full': f"{full_code}: {obj_text}",
                    })

    # ── Step 2: Extract terminal objectives for target cycle ──
    terminal_objectives = {
        'number': [], 'measures': [], 'fractions': [],
        'shape': [], 'statistics': [], 'algebra': [],
    }
    current_topic = None
    in_target_cycle = False

    cycle_headers = {
        '4': ['cycle 4', 'cycle four'],
        '5': ['cycle 5', 'cycle five'],
    }
    target_cycle_keys = []
    for tc in target_cycles:
        tc_num = tc.replace('CYCLE ', '')
        target_cycle_keys.extend(cycle_headers.get(tc_num, [f'cycle {tc_num}']))

    for line in lines:
        clean = line.strip()
        lower = clean.lower()

        # Detect cycle section
        if any(ck in lower for ck in ['cycle 1', 'cycle 2', 'cycle 3', 'cycle 4', 'cycle 5']):
            in_target_cycle = any(ck in lower for ck in target_cycle_keys)
            continue

        if not in_target_cycle:
            continue

        # Detect topic headers (match sub-headers in terminal objectives section)
        topic_map = {
            'whole number': 'number', 'number': 'number',
            'addition': 'number', 'subtraction': 'number',
            'multiplication': 'number', 'division': 'number',
            'ratio': 'number', 'proportion': 'number', 'percentage': 'fractions',
            'measurement': 'measures', 'measure': 'measures',
            'length': 'measures', 'mass': 'measures', 'money': 'measures',
            'time': 'measures', 'capacity': 'measures', 'temperature': 'measures',
            'decimal measure': 'measures',
            'fraction': 'fractions', 'decimal': 'fractions',
            'shape': 'shape', 'space': 'shape', 'angle': 'shape',
            'symmetry': 'shape', 'transformation': 'shape', 'position': 'shape',
            '2-d': 'shape', '3-d': 'shape', 'movement': 'shape',
            'statistic': 'statistics', 'data': 'statistics', 'handling data': 'statistics',
            'probability': 'statistics',
            'algebra': 'algebra', 'equation': 'algebra', 'pattern': 'algebra',
        }
        if len(clean) < 80:
            for keyword, topic in topic_map.items():
                if keyword in lower:
                    current_topic = topic
                    break

        # Extract bullet-point objectives (handle \uf0b7 Windows bullet, •, -, etc.)
        if current_topic and (clean.startswith('\u2022') or clean.startswith('-') or clean.startswith('•') or clean.startswith('\uf0b7')):
            obj = clean.lstrip('•-\u2022\uf0b7 ').strip().rstrip('.')
            if len(obj) > 15 and not obj.startswith('---'):
                terminal_objectives[current_topic].append(obj)

    # ── Step 3: Build strand-based units ──
    STRANDS = [
        {
            'key': 'number',
            'title': 'Number',
            'sub_strands': [
                ('Whole Numbers and Place Value', ['place value', 'read', 'write', 'order', 'numeral']),
                ('Operations with Whole Numbers', ['add', 'subtract', 'multiply', 'divide', 'operation', 'mental']),
                ('Fractions, Decimals and Percentages', ['fraction', 'decimal', 'percentage', 'mixed number', 'equivalent']),
                ('Ratio and Proportion', ['ratio', 'proportion', 'scale']),
            ],
            'topics': ['number', 'fractions'],
        },
        {
            'key': 'algebra',
            'title': 'Algebra',
            'sub_strands': [
                ('Patterns and Sequences', ['pattern', 'sequence', 'term', 'generaliz']),
                ('Expressions and Equations', ['expression', 'equation', 'formula', 'solve', 'variable', 'letter']),
                ('Functions and Graphs', ['function', 'graph', 'coordinate', 'plot']),
            ],
            'topics': ['algebra'],
        },
        {
            'key': 'shape',
            'title': 'Shape and Space',
            'sub_strands': [
                ('2-D and 3-D Shapes', ['shape', 'triangle', 'rectangle', 'circle', 'polygon', 'cube', 'prism', 'net']),
                ('Angles and Measurement', ['angle', 'protractor', 'degree', 'bearing']),
                ('Symmetry and Transformation', ['symmetry', 'reflect', 'rotate', 'translate', 'transform', 'enlarg']),
                ('Coordinates and Position', ['coordinate', 'grid', 'position', 'location']),
            ],
            'topics': ['shape'],
        },
        {
            'key': 'measures',
            'title': 'Measures',
            'sub_strands': [
                ('Money', ['money', 'rupee', 'currency', 'profit', 'loss', 'interest', 'wage']),
                ('Time and Temperature', ['time', 'clock', 'calendar', 'temperature', 'timetable']),
                ('Length, Mass and Capacity', ['length', 'mass', 'capacity', 'weight', 'metre', 'kilogram', 'litre', 'convert']),
                ('Area, Perimeter and Volume', ['area', 'perimeter', 'volume', 'surface']),
            ],
            'topics': ['measures'],
        },
        {
            'key': 'handling_data',
            'title': 'Handling Data',
            'sub_strands': [
                ('Data Collection and Graphs', ['data', 'collect', 'graph', 'chart', 'table', 'bar', 'pie', 'histogram', 'frequency']),
                ('Probability', ['probability', 'likely', 'certain', 'chance', 'random']),
            ],
            'topics': ['statistics'],
        },
    ]

    # Collect all enabling objectives from K/S codes
    all_eo = [o['full'] for o in coded_objectives['K']] + [o['full'] for o in coded_objectives['S']]

    units = []
    for strand in STRANDS:
        # Collect terminal objectives for this strand's topics
        strand_terminal = []
        for topic_key in strand['topics']:
            strand_terminal.extend(terminal_objectives.get(topic_key, []))

        # Match coded objectives to this strand by keyword
        strand_enabling = []
        strand_keywords = set()
        for _, keywords in strand['sub_strands']:
            strand_keywords.update(keywords)

        for eo in all_eo:
            eo_lower = eo.lower()
            if any(kw in eo_lower for kw in strand_keywords):
                strand_enabling.append(eo)

        # Build lessons from sub-strands
        lessons = []
        for sub_title, sub_keywords in strand['sub_strands']:
            # Match terminal objectives to this sub-strand
            sub_objectives = []
            for obj in strand_terminal:
                obj_lower = obj.lower()
                if any(kw in obj_lower for kw in sub_keywords):
                    sub_objectives.append(obj)

            # Match enabling objectives to this sub-strand
            sub_enabling = []
            for eo in strand_enabling:
                eo_lower = eo.lower()
                if any(kw in eo_lower for kw in sub_keywords):
                    sub_enabling.append(eo)

            # Also create enabling objectives from terminal objectives if no K/S codes matched
            if not sub_enabling and sub_objectives:
                sub_enabling = sub_objectives[:5]

            if sub_objectives or sub_enabling:
                lessons.append({
                    "title": sub_title,
                    "objective": sub_objectives[0] if sub_objectives else f"Master {sub_title.lower()} concepts",
                    "enabling_objectives": sub_enabling or sub_objectives[:5],
                    "teaching_strategies": ["Worked examples", "Practice exercises", "Problem solving"],
                    "resources": ["Textbook", "Whiteboard", "Calculator"],
                    "assessment_methods": ["Written exercises", "Problem-solving tasks"],
                    "estimated_minutes": 20,
                    "order": len(lessons) + 1,
                })

        if lessons:
            units.append({
                "number": len(units) + 1,
                "title": f"{grade_level}: {strand['title']}",
                "grade_level": grade_level,
                "duration": "Multiple periods",
                "introduction": f"{strand['title']} strand for {grade_level} (Cycle {cycle})",
                "terminal_objectives": strand_terminal[:15],
                "enabling_objectives": strand_enabling[:20],
                "lessons": lessons,
            })

    teaching_strategies = [
        "Worked examples with step-by-step solutions",
        "Mental computation practice",
        "Problem solving in real-world contexts",
        "Group work and collaborative learning",
        "Use of manipulatives and visual aids",
        "Practice exercises with graduated difficulty",
        "Application to Seychelles context (SCR, local measurements)",
    ]

    assessment_methods = [
        "Written exercises",
        "Problem-solving tasks",
        "Mental math tests",
        "Practical investigations",
    ]

    return ParsedCurriculum(
        subject="Mathematics",
        grade_level=grade_level,
        cycle=cycle,
        description=f"Mathematics curriculum for Seychelles secondary schools (Cycle {cycle})",
        units=units,
        teaching_strategies=teaching_strategies,
        assessment_methods=assessment_methods,
    )


def create_lessons_from_objectives(objectives: List[str], unit_title: str) -> List[Dict]:
    """Create lesson structures from terminal objectives."""
    lessons = []
    
    for i, objective in enumerate(objectives):
        # Create lesson title from objective
        title = create_lesson_title(objective)
        
        lessons.append({
            "title": title,
            "objective": objective,
            "enabling_objectives": create_enabling_objectives(objective),
            "teaching_strategies": ["Worked examples", "Practice exercises", "Discussion"],
            "resources": get_resources_for_topic(unit_title),
            "assessment_methods": ["Written exercises", "Oral questioning"],
            "estimated_minutes": 20,
            "order": i + 1
        })
    
    return lessons


def create_lesson_title(objective: str) -> str:
    """Create a student-friendly lesson title from an objective."""
    # Remove common prefixes
    prefixes = [
        "demonstrate the understanding of",
        "demonstrate understanding of",
        "understand and use",
        "use with confidence",
        "appreciate, discuss and express ideas about",
        "work out",
        "apply",
        "solve problems involving",
        "distinguish between",
        "recognise and name",
        "draw and measure",
        "use the",
        "form",
        "determine",
        "construct and solve",
        "develop",
        "know",
        "make",
        "choose",
    ]
    
    title = objective
    obj_lower = objective.lower()
    
    for prefix in prefixes:
        if obj_lower.startswith(prefix):
            title = objective[len(prefix):].strip()
            break
    
    # Capitalize first letter
    if title:
        title = title[0].upper() + title[1:]
    
    # Limit length
    if len(title) > 60:
        title = title[:57] + "..."
    
    return title or objective[:60]


def create_enabling_objectives(terminal_objective: str) -> List[str]:
    """Break a terminal objective into smaller enabling objectives."""
    parts = []
    
    # If objective mentions multiple skills, split them
    if " and " in terminal_objective.lower():
        segments = re.split(r'\s+and\s+', terminal_objective, flags=re.IGNORECASE)
        for seg in segments[:3]:  # Max 3 segments
            seg = seg.strip()
            if seg and len(seg) > 10:
                parts.append(seg)
    
    if not parts:
        parts.append(terminal_objective)
    
    # Add standard enabling objectives
    return parts[:5]  # Limit to 5


def get_resources_for_topic(unit_name: str) -> List[str]:
    """Get appropriate resources for a topic."""
    base = ["Textbook", "Workbook", "Whiteboard"]
    
    extras = {
        "Number": ["Calculator", "Number line"],
        "Measures": ["Rulers", "Measuring tape", "Scales"],
        "Shape and Space": ["Protractor", "Compass", "Graph paper"],
        "Algebra": ["Algebra tiles", "Graphing calculator"],
        "Handling Data": ["Graph paper", "Dice", "Survey forms"],
    }
    
    return base + extras.get(unit_name, [])


# ============================================================================
# GEOGRAPHY CURRICULUM PARSER
# ============================================================================

def parse_geography_curriculum(text: str, grade_level: str = "S1") -> ParsedCurriculum:
    """
    Parse Seychelles Geography curriculum text.

    The geography syllabus has two formats:
    - Cycle 4 (S1-S2): Unit blocks with terminal objectives + Teaching/Learning Scheme
      tables containing enabling objectives in the first column.
    - Cycle 5 (S3-S5): IGCSE themes with "Candidates should be able to:" objectives.

    Each unit block structure:
        Unit N: Title
        Duration: X periods
        Introduction: ...
        Terminal objectives: 1. ... 2. ...
        Teaching/Learning Scheme:
          Enabling objectives | Teaching Strategies | Resources | Assessment
    """
    import re as _re
    from apps.curriculum.utils import determine_cycles
    cycles = determine_cycles(grade_level)
    cycle = "/".join(cycles)

    gl = grade_level.upper()
    # Cycle 4 = S1, S2, S3 (detailed syllabus with Teaching/Learning Scheme)
    # Cycle 5 = S4, S5 (IGCSE programme with themes)
    is_cycle_5_only = gl in ('S4', 'S5')
    in_igcse_section = False  # Track when we enter IGCSE section

    lines = text.split('\n')
    units = []
    current_unit = None
    current_section = None  # 'intro', 'terminal', 'enabling', 'assessment'
    current_term_grade = grade_level

    # Bullet chars used in the document
    BULLETS = {'\uf0b7', '•', '-', '\u2022'}

    for i, raw_line in enumerate(lines):
        line = raw_line.strip()
        if not line:
            continue

        line_lower = line.lower()

        # ── Detect IGCSE/Cycle 5 section boundary ──
        # Only match "CYCLE 5" as a standalone header, not mentions within unit text
        if _re.match(r'^CYCLE\s+5', line, _re.IGNORECASE) or 'igcse programme' in line_lower:
            in_igcse_section = True
            if not is_cycle_5_only:
                # For S1-S3: stop processing when we hit the IGCSE section header
                break  # Exit the loop entirely — we have all Cycle 4 content

        # ── Detect grade/term headers ──
        term_match = _re.match(r'Secondary\s+(One|Two|Three|Four|Five)', line, _re.IGNORECASE)
        if term_match:
            grade_map = {'one': 'S1', 'two': 'S2', 'three': 'S3', 'four': 'S4', 'five': 'S5'}
            current_term_grade = grade_map.get(term_match.group(1).lower(), grade_level)

        # Skip content not for our target grade
        # For S1-S3: only process Cycle 4 content (not IGCSE)
        if not is_cycle_5_only and in_igcse_section:
            continue
        # For S4-S5: only process IGCSE content
        if is_cycle_5_only and not in_igcse_section:
            # Still detect unit headers to skip Cycle 4 units
            if _re.match(r'Unit\s+\d+\s*:', line):
                continue

        # ── Detect unit headers ──
        unit_match = _re.match(r'Unit\s+(\d+)\s*:\s*(.+)', line)
        if not unit_match and (is_cycle_5_only or in_igcse_section):
            # IGCSE theme format: "Theme 1: Population and settlement" or "1.1 Population dynamics"
            unit_match = _re.match(r'(?:Theme\s+)?(\d+(?:\.\d+)?)\s*[:.]\s*(.+)', line)
            if unit_match and len(unit_match.group(2)) < 5:
                unit_match = None  # Skip short matches like page numbers

        if unit_match:
            # Save previous unit
            if current_unit and (current_unit['terminal_objectives'] or current_unit['enabling_objectives']):
                units.append(current_unit)

            current_unit = {
                'number': len(units) + 1,
                'title': unit_match.group(2).strip().rstrip('.'),
                'grade_level': current_term_grade,
                'duration': '',
                'introduction': '',
                'terminal_objectives': [],
                'enabling_objectives': [],
                'teaching_strategies': [],
                'resources': [],
                'assessment_methods': [],
                'lessons': [],
            }
            current_section = 'intro'
            continue

        if not current_unit:
            continue

        # ── Detect duration ──
        dur_match = _re.match(r'Duration\s*:\s*(.+)', line, _re.IGNORECASE)
        if dur_match:
            current_unit['duration'] = dur_match.group(1).strip()
            continue

        # ── Detect section transitions ──
        if 'terminal objective' in line_lower:
            current_section = 'terminal'
            continue
        if 'enabling objective' in line_lower:
            current_section = 'enabling'
            continue
        if 'teaching/learning scheme' in line_lower or 'teaching/learning strategies' in line_lower:
            current_section = 'enabling'  # The scheme table starts with enabling objectives column
            continue
        if 'unit assessment' in line_lower or 'assessment:' in line_lower:
            current_section = 'assessment'
            continue
        if 'candidates should be able to' in line_lower:
            current_section = 'enabling'  # IGCSE format
            continue
        if line_lower.startswith('resources') or line_lower.startswith('assessment'):
            # Column headers in the Teaching/Learning Scheme table — skip
            continue

        # ── Extract content based on section ──
        if current_section == 'intro' and len(line) > 20:
            current_unit['introduction'] += ' ' + line

        elif current_section == 'terminal':
            # Numbered objectives: "1. Understand that..."
            obj_match = _re.match(r'\d+\.?\s+(.+)', line)
            if obj_match:
                obj = obj_match.group(1).strip().rstrip('.')
                if len(obj) > 10:
                    current_unit['terminal_objectives'].append(obj)
            elif line[0] in BULLETS:
                obj = line.lstrip(''.join(BULLETS) + ' ').strip().rstrip('.')
                if len(obj) > 10:
                    current_unit['terminal_objectives'].append(obj)

        elif current_section == 'enabling':
            # Bullet-point enabling objectives
            if line[0] in BULLETS:
                obj = line.lstrip(''.join(BULLETS) + ' ').strip().rstrip('.')
                if len(obj) > 10:
                    current_unit['enabling_objectives'].append(obj)
            elif len(line) > 15 and not any(kw in line_lower for kw in [
                'teacher', 'oral questioning', 'discussion', 'textbook',
                'written activities', 'group work', 'explanation', 'brainstorm',
                'handout', 'atlas', 'globe', 'video', 'map interpretation',
                'informal', 'structured', 'source-based', 'photographs',
                'exposition', 'demonstration', 'practical', 'newspaper',
                'magazine', 'relevant document', 'handbook', 'teaching',
                'learning strategies', 'visit:', 'meteorological',
                'blue economy', 'documentary', 'chart', 'analysis of',
                'interpretation of', 'videos', 'case study', 'research',
                'field', 'investigate', 'project', 'presentation',
            ]):
                # Continuation of previous objective or standalone objective
                if current_unit['enabling_objectives']:
                    # Append to last objective if it looks like continuation
                    last = current_unit['enabling_objectives'][-1]
                    if not last.endswith('.') and len(line) < 80:
                        current_unit['enabling_objectives'][-1] = last + ' ' + line
                    elif len(line) > 20:
                        current_unit['enabling_objectives'].append(line.rstrip('.'))
                elif len(line) > 20:
                    current_unit['enabling_objectives'].append(line.rstrip('.'))

        elif current_section == 'assessment':
            if len(line) > 15:
                current_unit['assessment_methods'].append(line.rstrip('.'))

    # Save last unit
    if current_unit and (current_unit['terminal_objectives'] or current_unit['enabling_objectives']):
        units.append(current_unit)

    # ── Deduplicate units (same unit appears twice: intro block + Teaching/Learning Scheme) ──
    # Strategy: merge by normalized title, prefer the version with more enabling objectives
    deduped = []
    seen_titles = {}
    for u in units:
        # Normalize: lowercase, strip punctuation and extra spaces, first 35 chars
        clean_title = _re.sub(r'[^a-z0-9 ]', '', u['title'].lower()).strip()
        clean_title = _re.sub(r'\s+', ' ', clean_title)[:35]
        if clean_title in seen_titles:
            # Merge into the existing unit (keep the one with more EOs)
            existing = seen_titles[clean_title]
            for eo in u.get('enabling_objectives', []):
                if eo not in existing.get('enabling_objectives', []):
                    existing.setdefault('enabling_objectives', []).append(eo)
            for to in u.get('terminal_objectives', []):
                if to not in existing.get('terminal_objectives', []):
                    existing.setdefault('terminal_objectives', []).append(to)
        else:
            seen_titles[clean_title] = u
            deduped.append(u)
    units = deduped

    # ── Filter units by target grade level ──
    if gl in ('S1', 'S2', 'S3', 'S4', 'S5'):
        filtered = [u for u in units if u.get('grade_level', '').upper() == gl]
        if filtered:
            units = filtered

    # ── Post-process: clean enabling objectives ──
    ACTION_VERBS = {
        'define', 'describe', 'explain', 'list', 'state', 'identify', 'name',
        'compare', 'distinguish', 'classify', 'demonstrate', 'apply', 'use',
        'calculate', 'solve', 'give', 'know', 'understand', 'recognise',
        'recognize', 'locate', 'illustrate', 'suggest', 'evaluate', 'discuss',
        'outline', 'show', 'draw', 'measure', 'construct', 'assess',
        'develop', 'determine', 'differentiate', 'interpret', 'analyse',
        'analyze', 'investigate', 'with', 'select', 'read', 'work',
    }
    for unit in units:
        cleaned = []
        for eo in unit.get('enabling_objectives', []):
            eo = eo.strip()
            if len(eo) < 15:
                continue
            first_word = eo.split()[0].lower().rstrip('s') if eo else ''
            # Keep if starts with an action verb or is clearly an objective
            if first_word in ACTION_VERBS or any(v in eo.lower() for v in ['should be able', 'will be able']):
                cleaned.append(eo)
        unit['enabling_objectives'] = cleaned

    # ── Build lessons from TERMINAL OBJECTIVES (primary) ──
    # Terminal objectives are the assessment targets — always clearly defined.
    # Enabling objectives supplement but are harder to extract reliably.
    for unit in units:
        # Use terminal objectives as the primary lesson driver
        tos = unit.get('terminal_objectives', [])
        eos = unit.get('enabling_objectives', [])

        # Terminal objectives are the lesson backbone
        # Enabling objectives provide additional granularity
        if tos:
            lesson_objectives = tos
        elif eos:
            lesson_objectives = eos
        else:
            continue

        # Store all objectives on the unit for competency tracking
        # Terminal objectives = what we assess, enabling = supplementary detail
        unit['enabling_objectives'] = lesson_objectives + [
            e for e in eos if e not in lesson_objectives
        ]

        # Group terminal objectives into lessons (1 TO per 20-minute lesson)
        # Attach relevant enabling objectives as teaching steps for content generation
        raw_eos = unit.get('enabling_objectives', [])  # The raw extracted EOs
        lessons = []
        chunk_size = 1  # One TO per lesson for focused 20-min sessions
        for start in range(0, len(lesson_objectives), chunk_size):
            chunk = lesson_objectives[start:start + chunk_size]
            if not chunk:
                continue

            # Match enabling objectives to this lesson's terminal objectives by keyword overlap
            lesson_eos = []
            chunk_words = set()
            for to in chunk:
                chunk_words.update(w.lower() for w in to.split() if len(w) > 3)

            for eo in raw_eos:
                eo_words = set(w.lower() for w in eo.split() if len(w) > 3)
                if len(chunk_words & eo_words) >= 2:
                    lesson_eos.append(eo)

            # Use first objective as lesson title seed
            title = create_lesson_title(chunk[0])
            lessons.append({
                'title': title,
                'objective': chunk[0],
                # Terminal objectives = what we assess (the lesson targets)
                'enabling_objectives': chunk,
                # Enabling objectives = teaching steps (inform content generation)
                'teaching_steps': lesson_eos,
                'teaching_strategies': ['Explanation', 'Discussion', 'Practical activities'],
                'resources': ['Textbook', 'Atlas', 'Maps'],
                'assessment_methods': ['Structured source-based written questions', 'Oral questioning'],
                'estimated_minutes': 20,
                'order': len(lessons) + 1,
            })

        unit['lessons'] = lessons

    teaching_strategies = [
        'Teacher exposition and explanation',
        'Oral questioning and discussion',
        'Map interpretation and analysis',
        'Group work and collaborative learning',
        'Written activities and exercises',
        'Use of photographs and visual sources',
        'Fieldwork and observation',
        'Data interpretation (graphs, tables, statistics)',
    ]

    assessment_methods = [
        'Structured source-based written questions',
        'Informal assessments based on oral questioning',
        'Written exercises and assignments',
        'Map and data interpretation tasks',
    ]

    return ParsedCurriculum(
        subject='Geography',
        grade_level=grade_level,
        cycle=cycle,
        description=f"Geography curriculum for Seychelles secondary schools (Cycle {cycle})",
        units=units,
        teaching_strategies=teaching_strategies,
        assessment_methods=assessment_methods,
    )


# ============================================================================
# LLM-BASED CURRICULUM PARSER (Robust)
# ============================================================================

def parse_curriculum_with_llm(text: str, subject: str, grade_level: str, institution_id: int = None) -> ParsedCurriculum:
    """
    Use LLM to parse curriculum structure from text.
    This is more robust than regex-based parsing.

    If institution_id is provided, queries the knowledge base for teaching
    material context to help align unit/lesson structure with textbooks.
    """
    from apps.llm.models import ModelConfig
    from apps.llm.client import get_llm_client

    # Get LLM client
    model_config = ModelConfig.get_for('generation')
    if not model_config:
        logger.warning("No LLM configured, falling back to regex parser")
        return parse_generic_curriculum(text, subject, grade_level)

    llm_client = get_llm_client(model_config)

    # Query knowledge base for teaching material context if available
    kb_context_str = ""
    if institution_id:
        try:
            from apps.curriculum.knowledge_base import CurriculumKnowledgeBase
            kb = CurriculumKnowledgeBase(institution_id=institution_id)
            kb_result = kb.query_for_content_generation(
                lesson_title=subject,
                lesson_objective=f"{subject} curriculum structure",
                unit_title="",
                subject=subject,
                grade_level=grade_level,
                n_results=8,
            )
            if kb_result and kb_result.chunks:
                excerpts = "\n\n".join(
                    f"--- From {c.get('metadata', {}).get('material_title', 'teaching material')} ---\n{c.get('content', '')[:400]}"
                    for c in kb_result.chunks[:6]
                    if c.get('content', '').strip()
                )
                if excerpts:
                    kb_context_str = f"""
REFERENCE MATERIAL FROM UPLOADED TEXTBOOKS/TEACHING RESOURCES:
The following excerpts are from textbooks and materials used at this school.
Align unit and lesson names with the terminology and structure used in these materials where appropriate.

{excerpts}
"""
        except Exception as e:
            logger.warning(f"KB query for curriculum parsing failed: {e}")

    # Truncate text if too long (keep first and last parts for context)
    max_chars = 30000
    if len(text) > max_chars:
        # Keep first 20k and last 10k
        text = text[:20000] + "\n\n[...middle section truncated...]\n\n" + text[-10000:]

    from apps.curriculum.utils import determine_cycles
    cycles = determine_cycles(grade_level)
    cycle = "/".join(cycles)

    prompt = f"""Analyze this curriculum document and extract its structure.

DOCUMENT TEXT:
{text}

CONTEXT:
- Subject: {subject}
- Grade Level: {grade_level} (Cycle {cycle})
- This is a Seychelles secondary school curriculum
{kb_context_str}
TASK:
Extract the curriculum structure as JSON with this format:
{{
    "units": [
        {{
            "title": "Unit title",
            "grade_level": "S2",
            "terminal_objectives": ["Broad outcome 1", "Broad outcome 2"],
            "enabling_objectives": ["Define term X", "State that Y", "Explain why Z"],
            "lessons": [
                {{
                    "title": "Lesson title (short, clear name)",
                    "objective": "What students will learn/be able to do",
                    "enabling_objectives": ["Define term X", "State that Y"]
                }}
            ]
        }}
    ]
}}

GUIDELINES:
1. Look for natural divisions in the curriculum (chapters, units, topics, strands, themes)
2. Each unit should have 3-15 lessons
3. Each lesson should cover ONE main concept or skill
4. Lesson titles should be clear and student-friendly (not "Objective 1.2")
5. If you find terminal objectives or learning outcomes, extract them as unit terminal_objectives
6. Extract enabling objectives (granular teaching steps with action verbs: define, state,
   explain, describe, identify, compare) and assign each to the relevant lesson
7. If the document has numbered sections, use those as units
8. Extract as many lessons as you can find - don't skip content
9. Each unit MUST have a grade_level field (e.g. "S1", "S2", "S3") based on the document content
10. If the document covers multiple grade levels, create separate units per grade

Return ONLY valid JSON, no explanation or markdown."""

    try:
        response = llm_client.generate(
            prompt=prompt,
            system_prompt="You are a curriculum parsing assistant. Extract structured curriculum data from documents. Return only valid JSON.",
            max_tokens=8000,
            temperature=0.1,
        )
        
        # Parse the JSON response
        content = response.get('content', '').strip()
        
        # Clean up common issues
        if content.startswith('```'):
            content = content.split('```')[1]
            if content.startswith('json'):
                content = content[4:]
        content = content.strip()
        
        parsed = json.loads(content)
        
        # Convert to our format
        units = []
        for i, unit_data in enumerate(parsed.get('units', [])):
            lessons = []
            for j, lesson_data in enumerate(unit_data.get('lessons', [])):
                lessons.append({
                    "title": lesson_data.get('title', f'Lesson {j+1}'),
                    "objective": lesson_data.get('objective', ''),
                    "enabling_objectives": lesson_data.get('enabling_objectives', []),
                    "teaching_strategies": ["Discussion", "Practice", "Examples"],
                    "resources": ["Textbook", "Whiteboard"],
                    "assessment_methods": ["Written work", "Oral questioning"],
                    "estimated_minutes": 20,
                    "order": j + 1
                })

            if lessons:  # Only add units that have lessons
                units.append({
                    "number": i + 1,
                    "title": unit_data.get('title', f'Unit {i+1}'),
                    "grade_level": unit_data.get('grade_level', grade_level),
                    "duration": "Multiple periods",
                    "introduction": f"{unit_data.get('title', '')} for {subject}",
                    "terminal_objectives": unit_data.get('terminal_objectives',
                        [l['objective'] for l in lessons if l['objective']]),
                    "enabling_objectives": unit_data.get('enabling_objectives', []),
                    "lessons": lessons
                })
        
        return ParsedCurriculum(
            subject=subject,
            grade_level=grade_level,
            cycle=cycle,
            description=f"{subject} curriculum for {grade_level}",
            units=units,
            teaching_strategies=["Discussion", "Practical work", "Group activities", "Problem solving"],
            assessment_methods=["Written tests", "Projects", "Oral questioning", "Practical tasks"]
        )
        
    except json.JSONDecodeError as e:
        logger.error(f"LLM returned invalid JSON: {e}")
        logger.error(f"Response was: {content[:500]}...")
        # Fall back to regex parser
        return parse_generic_curriculum(text, subject, grade_level)
    except Exception as e:
        logger.error(f"LLM parsing failed: {e}")
        return parse_generic_curriculum(text, subject, grade_level)


# ============================================================================
# FALLBACK: REGEX-BASED PARSER
# ============================================================================

def parse_generic_curriculum(text: str, subject: str, grade_level: str) -> ParsedCurriculum:
    """
    Generic curriculum parser using flexible text extraction.
    Handles various document formats.
    """
    from apps.curriculum.utils import determine_cycles
    cycles = determine_cycles(grade_level)
    cycle = "/".join(cycles)

    units = []
    current_unit = None
    objectives = []
    
    lines = text.split('\n')
    
    # Patterns that indicate a section/unit header
    def is_header(line):
        line = line.strip()
        if not line or len(line) < 5:
            return False
        # **Bold headers**
        if line.startswith('**') and line.endswith('**'):
            return True
        # ALL CAPS headers
        if line.isupper() and 5 < len(line) < 60:
            return True
        # Numbered sections like "1. Introduction" or "Unit 1:"
        if re.match(r'^(Unit\s+)?\d+[\.\:\)]\s+\w', line, re.IGNORECASE):
            return True
        # Headers ending with colon
        if line.endswith(':') and len(line) < 50 and not line.startswith('-'):
            return True
        # Markdown headers
        if line.startswith('#'):
            return True
        return False
    
    # Patterns that indicate an objective/learning point
    def is_objective(line):
        line = line.strip()
        if not line or len(line) < 10:
            return False
        # Bullet points
        if line.startswith('-') or line.startswith('•') or line.startswith('*'):
            return True
        # Numbered lists
        if re.match(r'^\d+[\.\)]\s+\w', line):
            return True
        # Lettered lists
        if re.match(r'^[a-zA-Z][\.\)]\s+\w', line):
            return True
        return False
    
    def clean_header(line):
        """Clean up header text."""
        line = line.strip()
        line = line.strip('*#:')
        line = re.sub(r'^(Unit\s+)?\d+[\.\:\)]\s*', '', line, flags=re.IGNORECASE)
        return line.strip()
    
    def clean_objective(line):
        """Clean up objective text."""
        line = line.strip()
        # Remove bullet/number prefix
        line = re.sub(r'^[-•*]\s*', '', line)
        line = re.sub(r'^\d+[\.\)]\s*', '', line)
        line = re.sub(r'^[a-zA-Z][\.\)]\s*', '', line)
        return line.strip()
    
    for line in lines:
        line_clean = line.strip()
        
        if is_header(line_clean):
            # Save previous unit if it has objectives
            if current_unit and objectives:
                units.append({
                    "number": len(units) + 1,
                    "title": current_unit,
                    "duration": "Multiple periods",
                    "introduction": f"{current_unit} for {subject}",
                    "terminal_objectives": objectives[:30],  # Limit per unit
                    "lessons": create_lessons_from_objectives(objectives[:30], current_unit)
                })
            
            # Start new unit
            current_unit = clean_header(line_clean)
            objectives = []
        
        elif is_objective(line_clean):
            obj = clean_objective(line_clean)
            # Filter garbage
            if (obj and 
                len(obj) > 10 and 
                not obj.startswith('---') and
                not all(c in '-=|+_' for c in obj.replace(' ', ''))):
                objectives.append(obj)
    
    # Save last unit
    if current_unit and objectives:
        units.append({
            "number": len(units) + 1,
            "title": current_unit,
            "duration": "Multiple periods",
            "introduction": f"{current_unit} for {subject}",
            "terminal_objectives": objectives[:30],
            "lessons": create_lessons_from_objectives(objectives[:30], current_unit)
        })
    
    # If no units found, try to extract ANY content as lessons
    if not units:
        # Look for sentences that could be objectives
        all_objectives = []
        
        for line in lines:
            line_clean = line.strip()
            
            # Skip very short or very long lines
            if len(line_clean) < 15 or len(line_clean) > 300:
                continue
            
            # Skip lines that look like metadata
            if any(skip in line_clean.lower() for skip in 
                   ['page', 'copyright', 'table of contents', 'index', 'chapter']):
                continue
            
            # Check if it looks like an objective (contains action verbs)
            action_verbs = ['understand', 'explain', 'describe', 'identify', 'analyze',
                           'apply', 'evaluate', 'create', 'define', 'list', 'compare',
                           'demonstrate', 'develop', 'recognize', 'use', 'know', 'learn']
            
            line_lower = line_clean.lower()
            if any(verb in line_lower for verb in action_verbs):
                # Clean it up
                obj = clean_objective(line_clean)
                if obj and len(obj) > 15:
                    all_objectives.append(obj)
        
        # Also try bullet points one more time with looser criteria
        if not all_objectives:
            for line in lines:
                line_clean = line.strip()
                if is_objective(line_clean):
                    obj = clean_objective(line_clean)
                    if obj and len(obj) > 10:
                        all_objectives.append(obj)
        
        # Create a single unit if we found anything
        if all_objectives:
            # Remove duplicates while preserving order
            seen = set()
            unique_objectives = []
            for obj in all_objectives:
                if obj.lower() not in seen:
                    seen.add(obj.lower())
                    unique_objectives.append(obj)
            
            units.append({
                "number": 1,
                "title": f"{subject} Fundamentals",
                "duration": "Multiple periods",
                "introduction": f"Core concepts for {subject}",
                "terminal_objectives": unique_objectives[:30],
                "lessons": create_lessons_from_objectives(unique_objectives[:30], subject)
            })
    
    return ParsedCurriculum(
        subject=subject,
        grade_level=grade_level,
        cycle=cycle,
        description=f"{subject} curriculum for {grade_level}",
        units=units,
        teaching_strategies=["Discussion", "Practical work", "Group activities", "Field observation"],
        assessment_methods=["Written tests", "Projects", "Oral questioning", "Practical assessment"]
    )


# ============================================================================
# MAIN PARSING FUNCTION
# ============================================================================

def parse_curriculum_file(file_path: str, subject: str, grade_level: str) -> Dict:
    """
    Main entry point for parsing curriculum files.
    
    This does NOT use AI - it extracts structure directly from text.
    """
    text, file_type = extract_text_from_file(file_path)
    
    if not text or len(text) < 100:
        raise ValueError("Could not extract meaningful text from file")
    
    # Detect subject if not provided
    detected_subject = detect_subject(text, subject)
    
    # Parse based on subject
    if 'math' in detected_subject.lower():
        curriculum = parse_mathematics_curriculum(text, grade_level)
    else:
        curriculum = parse_generic_curriculum(text, detected_subject, grade_level)
    
    return {
        "curriculum": asdict(curriculum),
        "source_file": file_path,
        "extraction_method": file_type,
    }


# ============================================================================
# DATABASE INTEGRATION
# ============================================================================

def create_curriculum_from_structure(structure: Dict, institution, upload=None) -> Dict:
    """
    Create Course, Units, and Lessons from parsed structure.
    """
    from apps.curriculum.models import Course, Unit, Lesson, LessonStep
    
    from apps.curriculum.utils import format_grade_display
    subject_name = structure.get('subject', 'General')
    grade = structure.get('grade_level', '')
    grade_display = format_grade_display(grade)

    # Create course
    course_name = f"{subject_name} {grade_display}"

    course, created = Course.objects.update_or_create(
        institution=institution,
        title=course_name,
        defaults={
            'grade_level': grade,
            'description': structure.get('description', ''),
            'is_published': False,
        }
    )
    
    if upload:
        upload.created_course = course
        upload.add_log(f"{'Created' if created else 'Updated'} course: {course.title}")
    
    lessons_created = 0
    units_created = 0
    
    # Create Units and Lessons
    for unit_data in structure.get('units', []):
        unit, u_created = Unit.objects.update_or_create(
            course=course,
            title=unit_data.get('title', 'Unnamed Unit'),
            defaults={
                'description': unit_data.get('introduction', ''),
                'order_index': unit_data.get('number', 0),
            }
        )
        
        if u_created:
            units_created += 1
        
        if upload:
            upload.add_log(f"  {'Created' if u_created else 'Updated'} unit: {unit.title}")
        
        # Create lessons
        for lesson_data in unit_data.get('lessons', []):
            lesson_metadata = {
                'enabling_objectives': lesson_data.get('enabling_objectives', []),
                'teaching_strategies': lesson_data.get('teaching_strategies', []),
                'resources': lesson_data.get('resources', []),
                'assessment_methods': lesson_data.get('assessment_methods', []),
            }
            
            lesson, l_created = Lesson.objects.update_or_create(
                unit=unit,
                title=lesson_data.get('title', 'Unnamed Lesson'),
                defaults={
                    'objective': lesson_data.get('objective', ''),
                    'estimated_minutes': lesson_data.get('estimated_minutes', 20),
                    'order_index': lesson_data.get('order', 0),
                    'is_published': False,
                    'metadata': lesson_metadata,
                }
            )
            
            if l_created:
                lessons_created += 1
                
                # Create a basic teach step
                LessonStep.objects.get_or_create(
                    lesson=lesson,
                    order_index=0,
                    defaults={
                        'step_type': 'teach',
                        'teacher_script': f"Today we will learn about: {lesson.objective}",
                    }
                )
                
                if upload:
                    upload.add_log(f"    Created lesson: {lesson.title}")
    
    if upload:
        upload.lessons_created = lessons_created
        upload.save()
    
    return {
        'course_id': course.id,
        'course_name': course.title,
        'units_created': units_created,
        'lessons_created': lessons_created,
    }


# ============================================================================
# MAIN PROCESSING FUNCTION (called by dashboard)
# ============================================================================

def process_curriculum_upload(upload_id: int, skip_review: bool = False) -> Dict:
    """
    Process a curriculum upload with optional teacher review.
    
    Flow:
    1. Extract text from document
    2. Parse curriculum structure
    3. (If not skip_review) Set status to 'review' and wait for approval
    4. Create database records
    
    Args:
        upload_id: ID of the CurriculumUpload record
        skip_review: If True, skip the review step and create records immediately
    """
    from apps.dashboard.models import CurriculumUpload
    
    upload = CurriculumUpload.objects.get(id=upload_id)
    
    try:
        upload.status = 'processing'
        upload.current_step = 1
        upload.processing_log = ""
        upload.add_log("🚀 Starting curriculum processing...")
        upload.save()
        
        # Step 1: Extract text
        upload.add_log("📄 Step 1: Extracting text from document...")
        upload.add_log(f"   File: {upload.file_path}")
        
        text, file_type = extract_text_from_file(upload.file_path)
        upload.extracted_text_length = len(text)
        upload.add_log(f"   ✓ Extracted {len(text):,} characters ({file_type})")
        
        # Show preview of extracted text for debugging
        if text:
            preview = text[:500].replace('\n', ' ')[:200]
            upload.add_log(f"   Preview: {preview}...")
        
        upload.save()
        
        if len(text) < 100:
            raise ValueError("Could not extract meaningful text from file. The document may be scanned images or in an unsupported format.")
        
        # Step 2: Parse curriculum structure
        upload.current_step = 2
        upload.add_log("📚 Step 2: Parsing curriculum structure...")
        upload.save()
        
        detected_subject = detect_subject(text, upload.subject_name)
        upload.add_log(f"   Subject detected: {detected_subject}")
        
        # Try LLM-based parsing first (more robust)
        try:
            upload.add_log("   Using AI to analyze document structure...")
            curriculum = parse_curriculum_with_llm(
                text, detected_subject, upload.grade_level or '',
                institution_id=upload.institution_id,
            )
            upload.add_log("   ✓ AI parsing complete")
        except Exception as e:
            upload.add_log(f"   ⚠️ AI parsing failed: {e}")
            upload.add_log("   Falling back to pattern-based parsing...")
            # Fall back to regex-based parsing
            if 'math' in detected_subject.lower():
                curriculum = parse_mathematics_curriculum(text, upload.grade_level or '')
            else:
                curriculum = parse_generic_curriculum(text, detected_subject, upload.grade_level or '')
        
        structure = asdict(curriculum)
        
        units_count = len(structure.get('units', []))
        lessons_count = sum(len(u.get('lessons', [])) for u in structure.get('units', []))
        
        upload.add_log(f"   ✓ Found {units_count} units with {lessons_count} lessons")
        
        # Log some details about what was found
        for unit in structure.get('units', [])[:3]:  # Show first 3 units
            upload.add_log(f"      📁 {unit.get('title')}: {len(unit.get('lessons', []))} lessons")
        
        upload.parsed_data = structure
        upload.save()
        
        # If no content found, show warning but still allow proceeding
        if lessons_count == 0:
            upload.add_log("⚠️ No lessons extracted. The document format may not be recognized.")
            upload.add_log(f"   Document had {len(text):,} characters of text.")
            upload.add_log("   Try uploading a document with clear sections and bullet points.")
            # Still go to review so teacher can see what happened
            upload.status = 'review'
            upload.add_log("⏸️ Please review - no content was extracted.")
            upload.save()
            
            return {
                'success': True,
                'status': 'review',
                'units_count': 0,
                'lessons_count': 0,
                'message': 'No lessons extracted. Please check document format.',
            }
        
        # If review is required, stop here and wait for teacher approval
        if not skip_review:
            upload.status = 'review'
            upload.add_log("⏸️ Waiting for teacher review...")
            upload.save()
            
            return {
                'success': True,
                'status': 'review',
                'units_count': units_count,
                'lessons_count': lessons_count,
                'message': 'Please review the parsed curriculum structure.',
            }
        
        # Step 3: Create database records
        return complete_curriculum_upload(upload_id)
        
    except Exception as e:
        logger.exception(f"Curriculum processing failed: {e}")
        upload.status = 'failed'
        upload.error_message = str(e)
        upload.add_log(f"❌ Error: {e}")
        upload.save()
        raise


def complete_curriculum_upload(upload_id: int, feedback: str = "") -> Dict:
    """
    Complete the curriculum upload by creating database records.
    Called after teacher approves the parsed structure.
    """
    from apps.dashboard.models import CurriculumUpload
    from django.utils import timezone
    
    upload = CurriculumUpload.objects.get(id=upload_id)
    
    try:
        upload.status = 'processing'
        upload.current_step = 3
        
        if feedback:
            upload.teacher_feedback = feedback
        
        upload.add_log("💾 Step 3: Creating curriculum records...")
        upload.save()
        
        structure = upload.parsed_data
        if not structure:
            raise ValueError("No parsed data available. Please re-process the document.")
        
        # Create database records
        result = create_curriculum_from_structure(
            structure=structure,
            institution=upload.institution,
            upload=upload
        )
        
        upload.units_created = result.get('units_created', 0)
        upload.lessons_created = result.get('lessons_created', 0)
        upload.add_log(f"   ✓ Created {result['units_created']} units, {result['lessons_created']} lessons")
        
        # Mark complete
        upload.status = 'completed'
        upload.completed_at = timezone.now()
        upload.add_log(f"✅ Complete! Course '{result['course_name']}' is ready.")
        upload.save()
        
        return {
            'success': True,
            'status': 'completed',
            'course_id': result['course_id'],
            'course_name': result['course_name'],
            'units_created': result['units_created'],
            'lessons_created': result['lessons_created'],
        }
        
    except Exception as e:
        logger.exception(f"Curriculum completion failed: {e}")
        upload.status = 'failed'
        upload.error_message = str(e)
        upload.add_log(f"❌ Error: {e}")
        upload.save()
        raise