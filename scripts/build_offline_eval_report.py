"""Generate Offline_Model_Evaluation_Report.docx — a 2-4 page report on the
offline / benchmark tutor-model evaluation (30 models).

Run:
    python scripts/build_offline_eval_report.py [--output <path>]
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = REPO_ROOT / 'offline_eval' / 'Offline_Model_Evaluation_Report.docx'


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_meta(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def add_para(doc, text, *, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, *, lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if lead:
        r = p.add_run(lead + ' ')
        r.bold = True
    p.add_run(text)
    return p


def add_table(doc, header, rows, *, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = str(val)
    return table


def build(output_path: Path) -> Path:
    doc = Document()

    # ----- Title block -----
    add_title(doc, 'Offline & Benchmark Model Evaluation')
    add_subtitle(doc, 'Selecting an open-source tutor model for offline, low-connectivity deployment')
    add_meta(doc, 'Nyansapo Labs — AI Tutor   ·   30 models scored   ·   June 2026')
    add_meta(doc, 'Branch: pixeldesignlabs-dev-portuguese   ·   Context: Mozambique / Tanzania pilots')
    doc.add_paragraph()

    # ----- 1. Executive Summary -----
    doc.add_heading('1. Executive Summary', level=1)
    add_para(doc,
        'The AI Tutor pilot runs on a hosted Anthropic model. For data residency and '
        'use in low-connectivity schools, we need a tutor model that can run locally — '
        'on a phone/tablet or a modest school server — with no cloud dependency. This '
        'study measured how well 23 open-source models and 7 proprietary benchmarks '
        'drive the real production tutoring engine, scored by the same trusted '
        'Anthropic-based harness across an identical 60-scenario test set.')
    add_para(doc,
        'The proprietary ceiling is Claude Opus 4.7 at a 90% pass rate. The best '
        'open-source model, Qwen2.5-14B, reaches 55% — about 61% of that ceiling — '
        'while a 7B Qwen runs on a laptop at 52% with the highest teaching-quality '
        'score of any open model (0.71), rivalling Gemini 3.5 Flash. A 3B Qwen that '
        'fits a phone/tablet scores 45%. The headline: a capable offline tutor is '
        'viable today, the Qwen2.5 family leads at every size, and the remaining gap '
        'to the cloud ceiling is realistically closable with tutor-specific tuning. '
        'These are stock models — no prompt or fine-tuning has been applied yet.')

    # ----- 2. Objective & Context -----
    doc.add_heading('2. Objective & Context', level=1)
    add_para(doc,
        'Two pilots (Mozambique, Tanzania) operate in settings where reliable '
        'internet cannot be assumed and where data-residency expectations favour '
        'on-device inference. The question this evaluation answers is concrete: which '
        'locally-runnable model best drives our actual tutoring pedagogy, and how far '
        'is it from the quality our cloud model delivers today? We deliberately ranked '
        'models on the real engine rather than a generic benchmark, because tool-use '
        'reliability and pedagogical behaviour — not raw language ability — determine '
        'tutoring quality.')

    # ----- 3. Methodology -----
    doc.add_heading('3. Methodology', level=1)
    add_para(doc,
        'Each model under test was swapped in as the tutor while the judge and the '
        'student-simulator were held constant on our production Anthropic models — a '
        'fixed, high-quality yardstick. Every model ran the identical set of 60 '
        'single-turn lesson scenarios (math and reading, across five student '
        'personas) and was scored on a pass/fail verdict plus a 0-1 teaching-quality '
        'rubric.')
    add_para(doc, 'Cross-family judging.', bold=True)
    add_para(doc,
        'The pass/fail grader excludes the tutor\'s own vendor, so no model ever '
        'grades itself — the headline pass rate is cross-family for all 30 models. '
        'Only secondary rubric/label layers can be same-vendor when a Claude or '
        'Gemini model is the tutor; those columns should be read with that in mind.')
    add_para(doc, 'Engine changes required to run open models.', bold=True)
    add_para(doc,
        'Two changes were needed (the production Anthropic path is unchanged). First, '
        'the engine was hard-wired to the Anthropic SDK; it now routes any provider '
        'through our pluggable client layer, so local models (via Ollama) can drive '
        'the same tool-based pedagogy. Second, several open models emit their tool '
        'calls as text rather than through the structured channel; the client now '
        'parses those leaks. This second fix mattered: GLM-4 would have scored ~0% as '
        'a pure artifact before it, and recovered to 43% once parsed correctly.')
    add_para(doc, 'Hardware tiers.', bold=True)
    add_para(doc,
        'Small models (<=9B) ran on a local 8 GB CPU-only laptop; 7-14B models ran on '
        'a free Google Colab T4 GPU; the proprietary benchmarks ran via API. The same '
        'harness and scenarios were used throughout, so scores are directly '
        'comparable across tiers.')

    # ----- 4. Results -----
    doc.add_heading('4. Results', level=1)

    doc.add_heading('4.1 Proprietary benchmark ceiling', level=2)
    add_table(doc,
        ['Model', 'Vendor', 'Pass rate', 'Rubric (0-1)'],
        [
            ['claude-opus-4-7', 'Anthropic', '90%', '0.88'],
            ['claude-haiku-4-5', 'Anthropic', '82%', '0.86'],
            ['claude-sonnet-4-6', 'Anthropic', '78%', '0.82'],
            ['gemini-2.5-flash', 'Google', '65%', '0.74'],
            ['gemini-3.1-pro-preview', 'Google', '58%', '0.71'],
            ['gemini-3.5-flash', 'Google', '50%', '0.66'],
            ['gemini-2.5-pro', 'Google', '43% *', '0.64'],
        ])
    add_para(doc,
        '* The Gemini Pro models score below the Flash models (2.5-flash 65% > '
        '3.1-pro 58% > 2.5-pro 43%), which is backwards. In id-probing the Pro models '
        'returned empty responses (thinking-mode), so these rows are most likely '
        'depressed by a harness interaction rather than true capability — see Section 6.',
        italic=True)

    doc.add_heading('4.2 Open-source models (deployment candidates)', level=2)
    add_table(doc,
        ['Model', 'Params', 'Device tier', 'Pass rate', 'Rubric'],
        [
            ['qwen2.5:14b', '14B', 'GPU server', '55%', '0.66'],
            ['mistral-nemo:12b', '12B', 'GPU server', '53%', '0.67'],
            ['qwen2.5:7b', '7B', 'GPU laptop', '52%', '0.71'],
            ['qwen2.5:3b', '3B', 'phone / tablet', '45%', '0.61'],
            ['glm4:9b', '9B', 'GPU laptop', '43%', '0.60'],
            ['granite3.1-dense:8b', '8B', 'GPU laptop', '33%', '0.56'],
            ['llama3.1:8b', '8B', 'GPU laptop', '33%', '0.53'],
            ['mistral:7b', '7B', 'GPU laptop', '32%', '0.54'],
            ['qwen2.5:1.5b', '1.5B', 'phone / tablet', '28%', '0.50'],
            ['(13 more, 28% down to 0%)', '0.5-10B', 'mixed', '<=28%', '<=0.46'],
        ])
    add_para(doc,
        'Three models scored 0% — falcon3:10b, phi4, and gemma2:2b. This is a '
        'tool-protocol failure, not a teaching failure: gemma2 has no tool capability, '
        'and falcon3/phi4 leak tool calls in a format the parser does not yet handle '
        '(the same trap GLM-4 was rescued from).',
        italic=True)

    doc.add_heading('4.3 The combined picture', level=2)
    add_para(doc,
        'Placed on one scale, the field separates into three bands: a Claude frontier '
        '(78-90%), a Gemini/best-OSS middle band where Gemini Flash, Qwen2.5-14B/7B '
        'and Mistral-Nemo all cluster around 50-65%, and a long tail of smaller or '
        'tool-incompatible models. The most important comparison for this project is '
        'that the best laptop-class open model (Qwen2.5-7B, 52%, rubric 0.71) sits '
        'level with Gemini 3.5 Flash (50%, 0.66) — an open 7B matching a cloud model '
        'on teaching quality.')

    # ----- 5. Key Findings -----
    doc.add_heading('5. Key Findings', level=1)
    add_bullet(doc,
        'a ~3B model is viable on phones/tablets today (45%), and a 7B on a school '
        'server is meaningfully better (52%) at the best teaching quality of any open '
        'model.', lead='Offline tutoring is viable.')
    add_bullet(doc,
        'Qwen2.5 leads at 3B, 7B, and 14B. Scaling helps with diminishing returns '
        '(45% -> 52% -> 55%), and the 7B has the single best rubric score of all open '
        'models.', lead='Qwen2.5 is the family to back.')
    add_bullet(doc,
        'Mistral-Nemo-12B (53%) beats Llama-3.1-8B (33%) by 20 points; the largest '
        'model that fit the 8 GB laptop sat mid-pack. Choosing the right family beats '
        'simply going bigger.', lead='Family matters more than size.')
    add_bullet(doc,
        'Opus 4.7 sets a 90% ceiling; the best offline model reaches ~61% of that '
        'pass rate. The gap is real but realistically closable with tuning.',
        lead='The cloud gap is closable.')
    add_bullet(doc,
        'Math reasoning and persona/tone adaptation are the two most common failure '
        'categories across nearly every model — the obvious first targets for prompt '
        'tuning.', lead='Two universal weak spots.')

    # ----- 6. What Can Be Improved -----
    doc.add_heading('6. What Can Be Improved', level=1)
    add_para(doc, 'Recover models under-measured by harness artifacts.', bold=True)
    add_para(doc,
        'Three capable models are currently scored below their real ability: phi4 and '
        'falcon3 (0%, tool-call leak format) and the Gemini Pro models '
        '(thinking-mode). Each should be given a short diagnostic run that captures '
        'its raw output, after which the parser/handler can be extended and the model '
        're-scored — exactly the path that recovered GLM-4 from 0% to 43%.')
    add_para(doc, 'Tutor-specific tuning.', bold=True)
    add_para(doc,
        'All numbers here are stock models with no prompt or fine-tuning. The leading '
        'offline candidate (Qwen2.5-7B/3B) should be prompt-tuned on its two weak '
        'spots — math and persona handling — and re-measured against the ceiling. '
        'This is the most likely lever to close the gap to the cloud models.')
    add_para(doc, 'Broaden the evaluation.', bold=True)
    add_bullet(doc,
        'Add multi-turn scenarios — the current run is single-turn only; session-level '
        'failures (loops, premature advance) are not yet measured per model.')
    add_bullet(doc,
        'Test the larger open tier (32B-70B: Qwen2.5-32B/72B, Command-R, Mixtral, '
        'Llama-3.3-70B) on an A100 to see whether scaling past 14B is worth a heavier '
        'server.')
    add_bullet(doc,
        'Grow scenario coverage toward the full failure-category vocabulary and add '
        'USD cost tracking per model.')
    add_para(doc, 'Tighten methodology for the benchmark rows.', bold=True)
    add_para(doc,
        'The headline pass rate is already cross-family judged. To make the secondary '
        'rubric fully vendor-neutral for the Claude/Gemini tutor rows, the rubric '
        'judge could be pinned to a third vendor (e.g. OpenAI) for those runs.')

    # ----- 7. Recommendation -----
    doc.add_heading('7. Recommendation', level=1)
    add_table(doc,
        ['Deployment target', 'Recommended model', 'Pass rate'],
        [
            ['Phone / tablet (on-device)', 'qwen2.5:3b', '45%'],
            ['Laptop / school server', 'qwen2.5:7b', '52% (rubric 0.71)'],
            ['Higher-spec server', 'qwen2.5:14b', '55%'],
            ['Cloud reference (today\'s pilot)', 'claude-opus-4-7', '90%'],
        ])
    add_para(doc,
        'Back the Qwen2.5 family for offline deployment, sized to the target device. '
        'Treat 45-55% as a starting baseline, prioritise a prompt-tuning pass on math '
        'and persona handling, and recover the three artifact-suppressed models before '
        'finalising the shortlist.')

    # ----- 8. Reproducibility -----
    doc.add_heading('8. Reproducibility', level=1)
    add_para(doc,
        'The full harness, model lists, runners, and per-model result files live under '
        'offline_eval/ in the repository; the leaderboard is regenerated with '
        'offline_eval/aggregate.py. Open models run via run_matrix.sh (Ollama) and the '
        'Colab notebook colab_eval.ipynb; cloud benchmarks via run_cloud.sh. Detailed '
        'findings accompany this report in offline_eval/FINDINGS_offline_model_eval.md.')

    doc.save(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('--output', type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    out = build(args.output)
    print(f"Wrote {out} ({out.stat().st_size} bytes)")
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
