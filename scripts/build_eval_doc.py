"""Generate AI_Tutor_Eval_Harness.docx — a comprehensive standalone
explainer of the eval harness built in evals/.

Updated to cover the full evaluation procedure end-to-end with no
assumed reader knowledge, including the rationale behind every design
choice. Visuals via tables and ASCII flow diagrams.

Run:
    python scripts/build_eval_doc.py [--output <path>]
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = REPO_ROOT / 'AI_Tutor_Eval_Harness.docx'


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_cover_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(28)


def add_cover_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_cover_metadata(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_para(doc, text, *, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_lead(doc, text):
    """Lead paragraph — slightly larger, semibold."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')


def add_numbered(doc, text):
    doc.add_paragraph(text, style='List Number')


def add_code(doc, text):
    """Monospace code/ASCII-diagram block. Uses 'No Spacing' style + Courier."""
    p = doc.add_paragraph(style='No Spacing')
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)


def add_callout(doc, text, *, label='Note'):
    """Sidebar-style highlighted paragraph — uses italic + indent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run = p.add_run(text)
    run.italic = True


def add_table(doc, header, rows, *, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = str(val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------

def build(output_path: Path) -> Path:
    doc = Document()

    # =====================================================================
    # Cover page
    # =====================================================================
    add_cover_title(doc, 'AI Tutor Evaluation Harness')
    add_cover_subtitle(doc,
                       'A comprehensive guide to how we measure and regress-test \n'
                       'the conversational tutor')
    doc.add_paragraph()
    add_cover_metadata(doc, [
        'Nyansapo Labs — AI Tutor (ai-tutor)',
        '',
        'Branch: dev   ·   Phases 1–6 shipped + bug-fix revision',
        'Dataset: 80 scenarios   ·   Personas: 6   ·   Subjects: math, geography',
        '',
        'Audience: engineers, pilot operators, and anyone reviewing tutor quality',
        '',
        'This document assumes no prior familiarity with the AI Tutor codebase.',
    ])

    add_page_break(doc)

    # =====================================================================
    # Table of contents (manual — Word would generate automatically but we
    # want to be self-contained without requiring the reader to update fields)
    # =====================================================================
    doc.add_heading('Table of Contents', level=1)
    add_para(doc, 'This document is organised in eight parts:')
    add_bullet(doc, 'Part I — Context: what the AI Tutor is, and why evaluation is hard')
    add_bullet(doc, 'Part II — Conceptual model: scenarios, personas, scoring layers')
    add_bullet(doc, 'Part III — The harness in detail: every file, every responsibility')
    add_bullet(doc, 'Part IV — The evaluation procedure: step-by-step from setup to report')
    add_bullet(doc, 'Part V — The dataset: persona × situation matrix and full breakdown')
    add_bullet(doc, 'Part VI — Rationale: why every design choice was made the way it was')
    add_bullet(doc, 'Part VII — Recent changes: bug fixes, dataset growth, first baseline')
    add_bullet(doc, 'Part VIII — Operational notes: code map, glossary, future work')

    add_page_break(doc)

    # =====================================================================
    # Executive Summary
    # =====================================================================
    doc.add_heading('Executive Summary', level=1)

    add_lead(doc,
        'The AI Tutor Evaluation Harness is a curated test suite that '
        'lets us measure, in a reproducible way, whether the conversational '
        'tutor is getting better or worse over time.')

    add_para(doc,
        'Before this harness existed, the team had two ways to judge tutor '
        'quality: (a) read transcripts from real pilot students after the '
        'fact, and (b) drive synthetic personas through the tutor and look '
        'at the output. Neither of these tools could answer the question a '
        'developer most often needs to answer — "did the change I just '
        'made make things better, worse, or unchanged?" Real-student '
        'transcripts arrive too slowly and require human labelling per item; '
        'synthetic traffic generates conversations but never scores them.')

    add_para(doc,
        'The harness fills that gap. It runs the same 80 carefully-curated '
        'scenarios every time, scores each one through three layers of '
        'automated checks, and writes the results to a dated JSON file. A '
        'companion report tool turns that file into a human-readable '
        'summary and — most importantly — a diff against any previous run, '
        'so a developer can see exactly which scenarios moved in which '
        'direction after their change.')

    add_para(doc,
        'The 80 scenarios span 6 distinct student personas (struggler, '
        'average, capable, probe-resistant, non-responder, and error-prone), '
        'two subjects (math and geography), two execution modes '
        '(single-turn response evaluation and full multi-turn session '
        'simulation), and twelve named failure categories taken from the '
        'project\'s deeper benchmark vocabulary. Each scenario is a YAML '
        'file checked into the repository, so the dataset is version-'
        'controlled in lockstep with the code.')

    add_para(doc,
        'Running the full suite is one command — `python manage.py '
        'run_eval`. Reading the result is another — `python -m evals.report '
        '--diff`. That is the whole loop. The rest of this document is the '
        'detail behind those two commands: what they do, why they do it '
        'that way, and how to extend them.')

    add_page_break(doc)

    # =====================================================================
    # PART I — Context and Foundations
    # =====================================================================
    doc.add_heading('Part I — Context and Foundations', level=1)

    doc.add_heading('1. What the AI Tutor is', level=2)
    add_para(doc,
        'The AI Tutor is a Django-based conversational tutoring platform '
        'aimed at secondary-school students. Currently piloting in the '
        'Seychelles (with Tanzania planned), it works like this: a '
        'student opens a chat window, selects a lesson (e.g., "Angles '
        'around a point"), and engages in a back-and-forth conversation '
        'with a large language model that has been prompted to behave '
        'as a patient, Socratic tutor. The lesson content — the teacher '
        'script, the worked examples, the practice questions, the exit '
        'ticket — is authored by curriculum specialists and stored as '
        'structured rows in the database. The tutor\'s job at each turn '
        'is to figure out where the student is in their understanding, '
        'react to whatever they just said, and either advance the lesson '
        'or scaffold a bit more.')

    add_para(doc,
        'Under the hood the system is composed of several layers:')
    add_bullet(doc,
        'A curriculum hierarchy (Course → Unit → Lesson → LessonStep) that '
        'holds the educational content.')
    add_bullet(doc,
        'A tutor engine (`apps/tutoring/conversational_tutor.py`) that '
        'orchestrates the conversation — picks the next step, builds the '
        'LLM prompt, processes the LLM response, and decides whether to '
        'advance.')
    add_bullet(doc,
        'A pluggable LLM client layer (`apps/llm/client.py`) that '
        'currently routes to Anthropic, OpenAI, Google, or Ollama based on '
        'a per-purpose `ModelConfig`.')
    add_bullet(doc,
        'A judge pipeline (`apps/tutoring/judges/unified.py`) that runs '
        'after each tutor response and inspects it for problems — wrong '
        'verdicts, factual errors, unfounded praise, unsafe content, etc.')
    add_bullet(doc,
        'A validator (`apps/tutoring/validator.py`) that catches structural '
        'issues (no question at the end, multi-paragraph, banned phrases).')

    add_para(doc,
        'The system has been iterated on for many months. Pilot data is '
        'collected, transcripts are reviewed by the project lead, and '
        'changes are made to the prompts, judges, and engine logic. The '
        'rate of change is high, which is what makes a regression-detection '
        'tool valuable.')

    doc.add_heading('2. Why evaluating an LLM tutor is hard', level=2)
    add_para(doc,
        'A traditional software system is evaluated by writing unit tests: '
        'given an input X, the output should equal Y. The test runs in '
        'milliseconds, produces a deterministic pass/fail, and ratchets '
        'forward. Most software CI is built on this foundation.')

    add_para(doc,
        'LLM-driven tutors don\'t fit this model. The output of the tutor '
        'is natural-language text, and what counts as "correct" is a '
        'multi-dimensional pedagogical judgement. Consider what the '
        'tutor\'s response needs to satisfy on a single turn:')

    add_bullet(doc, 'It must be factually correct.')
    add_bullet(doc, 'It must be pedagogically appropriate for this student\'s level.')
    add_bullet(doc, 'It must not affirm a wrong answer or reject a right one.')
    add_bullet(doc, 'It must not reveal the answer prematurely.')
    add_bullet(doc, 'It must be appropriately concise (a mobile chat — not a textbook).')
    add_bullet(doc, 'It must end with a question that hands the conversational floor back.')
    add_bullet(doc, 'It must not contain certain banned phrases that the team has learned cause problems with pilot students.')
    add_bullet(doc, 'It must not reference figures, diagrams, or images that the engine did not actually attach.')

    add_para(doc,
        'No single test can capture all of these. Worse, the tutor is '
        'producing fresh text every time — there\'s no fixed expected '
        'output to compare against. The same prompt can yield slightly '
        'different responses across runs. So evaluation has to be '
        'evaluative in a richer sense: less "did the output match a string" '
        'and more "did the response satisfy a set of properties".')

    doc.add_heading('3. The two existing quality tools — and the gap they leave', level=2)
    add_para(doc,
        'Before this harness, the project had two quality systems, each '
        'doing useful work but neither answering the "did my change '
        'regress anything" question.')

    add_para(doc, 'Tool 1: production sampling (apps/benchmark/)', bold=True)
    add_para(doc,
        'This system samples real student turns from production — the '
        'conversations our pilot students actually had — and asks the '
        'project lead to manually label what went well or poorly with the '
        'tutor\'s response. The labels are grounded in real usage and the '
        'data is high-fidelity. But the dataset is constantly shifting (a '
        'new pilot session yields new turns), every item requires human '
        'labelling before it\'s useful, and the data only arrives as '
        'fast as pilot students log on. You cannot make a code change at '
        '3pm and run "the benchmark" at 3:05pm — there\'s no fixed '
        'benchmark to run.')

    add_para(doc, 'Tool 2: synthetic student simulator (apps/tutoring/student_sim/)', bold=True)
    add_para(doc,
        'This system uses an LLM-backed "student" persona (struggler, '
        'capable, etc.) to drive the tutor through a full conversation. '
        'It produces realistic conversation traces on demand, no pilot '
        'students needed. But the simulator only generates traffic; it '
        'does NOT score it. The transcripts it produces feed back into '
        'the production-sampling benchmark, which means they still need '
        'human labelling. The simulator can answer "what does the tutor '
        'do when I push it with a non-responder persona for 12 turns?" '
        'but not "is the tutor\'s response actually good?"')

    add_para(doc, 'The gap', bold=True)
    add_para(doc,
        'What was missing — and what this harness provides — was an '
        'automated regression suite: same inputs every run, automatic '
        'scoring against pre-declared expectations, one command in, one '
        'number out, with run-over-run diffs. It is complementary to '
        'the other two systems, not a replacement. The production-'
        'sampling benchmark still tells us what real students are '
        'experiencing. The simulator still generates synthetic '
        'transcripts when we want them. The eval harness sits between '
        'them and gives developers a tight feedback loop on code '
        'changes.')

    add_page_break(doc)

    # =====================================================================
    # PART II — Conceptual Model
    # =====================================================================
    doc.add_heading('Part II — Conceptual Model', level=1)

    doc.add_heading('4. What "evaluation" means in this harness', level=2)
    add_para(doc,
        'An evaluation, in this context, is a run of the entire dataset '
        'against the currently-deployed tutor engine and prompts. Each '
        'scenario in the dataset is run independently; each produces a '
        'pass/fail verdict; the run is summarised as a pass rate plus '
        'a breakdown.')

    add_para(doc, 'The unit of testing is a scenario — a YAML file that captures:')
    add_bullet(doc, 'The persona the (synthetic) student is playing.')
    add_bullet(doc, 'The lesson the conversation is anchored to.')
    add_bullet(doc, 'A conversation prefix (for single-turn scenarios) or just a lesson handle (for multi-turn).')
    add_bullet(doc, 'The student\'s next message (single-turn only) or the maximum number of turns (multi-turn).')
    add_bullet(doc, 'A set of assertions about what the tutor\'s response must/must not contain.')
    add_bullet(doc, 'A rubric — a list of pedagogical properties an LLM-as-judge will score.')
    add_bullet(doc, 'A pass threshold for the rubric.')

    add_para(doc,
        'A scenario passes if (a) all the deterministic assertions pass '
        'AND (b) the rubric mean score meets the pass threshold. Either '
        'failure alone is enough to fail the scenario; both layers vote.')

    doc.add_heading('5. The five — now six — personas', level=2)
    add_para(doc,
        'A "persona" is a fictional student profile encoded as a system '
        'prompt for a synthetic-student LLM. The personas don\'t see the '
        'lesson content directly — they just react to whatever the tutor '
        'says, staying in character. The tutor sees them as if they were '
        'a real student.')

    add_para(doc, 'Why have personas at all?', bold=True)
    add_para(doc,
        'Because different students stress different parts of the tutor '
        'engine. A capable student tests whether the tutor over-scaffolds. '
        'A non-responder tests whether the tutor info-dumps on "ok". A '
        'probe-resistant student tests whether the tutor falls into a '
        'banned-opener loop. The same engine code will look "fine" against '
        'a cooperative average student and reveal serious failure modes '
        'against an adversarial one. Persona coverage is how we surface '
        'those failure modes deliberately rather than waiting for them to '
        'appear in pilot data.')

    add_para(doc, 'The personas, in their behavioural profile:')

    add_table(doc,
        ['Persona', 'Behaviour summary', 'Approx accuracy', 'Stresses the tutor on...'],
        [
            ['STRUGGLER',
             'Misreads questions, arithmetic slips, partial work, asks for help often, sometimes gives up ("idk")',
             '~30%',
             'Remediation flow, working-request handling, scaffolding'],
            ['AVERAGE',
             'Gets most answers right with mixed working presentation; the modal student profile',
             '~65%',
             'The steady-state path, false-accept guards, format rules'],
            ['CAPABLE',
             'Right answers, pushes back on tutor errors, asks clarifying questions, suggests alternative methods',
             '~85%',
             'Restraint when the student is moving fast, tutor honesty under challenge'],
            ['PROBE_RESISTANT',
             'Bare answers only — refuses to show working ("I just know", "I guessed")',
             '~60%',
             'Working-request loop, banned-opener repetition, regen path'],
            ['NON_RESPONDER',
             'Monosyllabic — "ok", "yes", "no", "idk"',
             'n/a',
             'Non-answer skip path, exit-ticket gating, premature-advance guard'],
            ['ERROR_PRONE',
             'Always commits to a specific wrong answer with a traceable error mode (arithmetic slip, wrong operation, echoed number, etc.)',
             '~10%',
             'Mistake-remediation pipeline; designed for BEA-2025 coverage'],
        ])

    add_callout(doc,
        'Personas are NOT a fair model of the student population. They are '
        'deliberately constructed stress profiles. A real pilot student is '
        'mostly some blend of AVERAGE and STRUGGLER. The other personas '
        'exist to surface specific failure modes the average case would '
        'never hit.',
        label='Important')

    doc.add_heading('6. Single-turn versus multi-turn modes', level=2)
    add_para(doc,
        'Scenarios come in two flavours. The choice between them depends '
        'on what kind of failure mode the scenario is hunting.')

    add_para(doc, 'Single-turn (single_turn mode)', bold=True)
    add_para(doc,
        'The runner creates a fresh TutorSession in the database, '
        'injects a canned conversation prefix (the `seed_history`) as '
        'past SessionTurn rows, and then calls '
        '`ConversationalTutor.respond(student_turn)` exactly once. The '
        'tutor produces a single response, the production judges fire '
        'against that response, and the runner scores it. This is the '
        'workhorse mode: most failure modes are visible in a single tutor '
        'response, single-turn scenarios are fast (under 30 seconds each '
        'on average), and they are highly reproducible because the '
        'conversation context is fixed.')

    add_para(doc, 'Multi-turn (multi_turn mode)', bold=True)
    add_para(doc,
        'The runner hands off to `apps.tutoring.student_sim.'
        'simulate_session`, which spins up an actual persona LLM and lets '
        'it have a full conversation with the tutor — opening turn, '
        'persona reply, tutor response, persona reply, and so on until '
        'the session terminates (completed, exit-ticket-shown, deadlock, '
        'or max-turns reached). The runner then reads every persisted '
        'SessionTurn back from the database, derives per-turn labels '
        'from the judges, and scores the whole trajectory against '
        'trajectory-specific assertions. Multi-turn scenarios catch '
        'session-level failure modes — banned-opener loops across turns, '
        'premature advance through monosyllabic replies, repetition '
        'patterns — that no single-turn check can see.')

    add_table(doc,
        ['Aspect', 'single_turn', 'multi_turn'],
        [
            ['Conversation context', 'Fixed (canned seed_history)', 'Generated by persona LLM in real time'],
            ['Tutor calls', '1 (one respond())', '~6–25 (full session)'],
            ['Wall time per scenario', '~10–30 s', '~30 s – 3 min'],
            ['Cost per scenario', '~$0.01', '~$0.05–0.50'],
            ['Best for', 'Per-turn pedagogical correctness, format rules', 'Trajectory-level patterns, session dynamics'],
            ['Determinism', 'High (same seed_history every run)', 'Moderate (persona LLM varies)'],
        ])

    doc.add_heading('7. The three scoring layers', level=2)
    add_para(doc,
        'A scenario passes only when every applicable scoring layer says '
        'so. Each layer catches a different category of failure; the '
        'composition produces a strong overall verdict from independently-'
        'cheap parts.')

    add_para(doc, 'Layer 1 — Deterministic checks (free, instant)', bold=True)
    add_para(doc,
        'Pure Python string and structural operations on the tutor\'s '
        'response. Examples: does the response end in `?`, does it '
        'contain any of these banned phrases, is it one paragraph or '
        'many, does it not contain a forbidden answer-letter, etc. These '
        'are encoded as the scenario\'s `assertions:` block. They run '
        'instantly and either pass or fail — no LLM involvement.')

    add_para(doc, 'Layer 2 — Judge-derived labels (cost = same as a prod turn)', bold=True)
    add_para(doc,
        'When the runner calls `respond()`, the production judge pipeline '
        '(`apps/tutoring/judges/unified.py`) fires automatically — the '
        'same judges that run in production every time the tutor speaks '
        'to a real student. Their output is mapped via '
        '`apps/benchmark/autopopulate.py` to a fixed vocabulary of 30 '
        'labels (UNFOUNDED_PRAISE, WRONG_VERDICT, BANNED_OPENER, '
        'LEAKS_ANSWER, etc.). Scenarios assert which labels must appear '
        'and which must not. This layer doesn\'t cost anything beyond '
        'what production already pays per turn.')

    add_para(doc, 'Layer 3 — LLM-as-judge rubric (~$0.001 per scenario)', bold=True)
    add_para(doc,
        'Each scenario can carry a `rubric:` block — a short list of '
        'pedagogical properties stated in natural language ("Confirms the '
        'answer briefly without effusive praise", "Engages with the '
        'student\'s actual claim rather than steamrolling past it"). A '
        'pinned small LLM (Claude Haiku 4.5 by default, temperature 0) '
        'reads the conversation and the tutor\'s response, scores each '
        'rubric item from 0.0 to 1.0, and the weighted mean is compared '
        'to the scenario\'s `pass_threshold`. This layer catches '
        'behaviours that Layers 1 and 2 can\'t articulate.')

    add_callout(doc,
        'Layers 1+2 are cheap and deterministic but limited to literal '
        'patterns. Layer 3 adds semantic judgement at small cost. The '
        'composition gives us both the rigour of explicit assertions and '
        'the flexibility of a judge that can reason about what a good '
        'response looks like.',
        label='Why hybrid')

    add_page_break(doc)

    # =====================================================================
    # PART III — The Harness in Detail
    # =====================================================================
    doc.add_heading('Part III — The Harness in Detail', level=1)

    doc.add_heading('8. Directory layout', level=2)
    add_para(doc,
        'The entire harness lives under `evals/` at the repository root, '
        'sibling to `apps/`. It is deliberately a top-level directory '
        'rather than a sub-component of any Django app, because the '
        'harness is testing-infrastructure, not application code.')

    add_code(doc,
'''evals/
├── README.md                — quick-start for new contributors
├── runner.py                — orchestrates everything: discovers scenarios,
│                              drives respond(), captures results
├── report.py                — reads run JSONs, prints summaries + diffs
├── personas.py              — re-exports the simulator personas
├── scorers/
│   ├── __init__.py          — AssertionResult dataclass
│   ├── deterministic.py     — Layer 1 + 2 scorer (string/label checks)
│   ├── trajectory.py        — multi-turn scorer (session-level checks)
│   └── llm_rubric.py        — Layer 3 scorer (LLM-as-judge)
├── fixtures/
│   ├── extract.py           — parses prod_content_dump.sql → fixtures
│   ├── institution.json     — eval institution + sim-bot user + ModelConfigs
│   └── lessons.json         — frozen Course→Unit→Lesson→Step trees
├── dataset/                 — the 80 scenario YAMLs (the dataset)
│   ├── math/                — math-specific scenarios (16)
│   ├── geography/           — geography-specific scenarios (10)
│   ├── multi_turn/          — full-session trajectory scenarios (20)
│   ├── crosscutting/        — safety, figure, format, repetition (24)
│   ├── format/              — format rule guards (3)
│   ├── pedagogy/            — pedagogy traps (2)
│   ├── personas/            — per-persona behavioural tests (5)
│   └── smoke/               — plumbing-check (1; excluded from full runs)
└── runs/                    — gitignored — per-run JSON result blobs
''')

    add_para(doc, 'The Django entry point is one management command:')
    add_code(doc, 'apps/tutoring/management/commands/run_eval.py    — thin CLI wrapper')

    doc.add_heading('9. The scenario YAML — the atomic unit', level=2)
    add_para(doc,
        'Every scenario lives in its own YAML file. The filename '
        '(without extension) MUST equal the `id:` field — the runner '
        'enforces this for traceability.')

    add_para(doc, 'A single-turn scenario, fully annotated:')
    add_code(doc,
'''id: math_correct_advance_001          # must match filename stem
description: >                           # human-readable; shown in reports
  AVERAGE gives the correct numeric answer (145) WITH working
  ("60+75+80=215, then 360-215=145"). Tutor must ADVANCE — brief
  confirmation, no over-probing, no second-guessing.
persona: average                          # one of: struggler, average, capable,
                                          # probe_resistant, non_responder, error_prone
subject: math                             # math | geography | crosscutting
lesson_id: 1137                           # references fixtures/lessons.json
tags: [advance, math, pedagogy]           # free-form, used for report grouping

mode: single_turn                         # OR multi_turn

# ─────────────── single_turn only ───────────────
seed_history:                             # canned prior conversation
  - role: tutor
    text: "Try this: four angles around a point measure 60°, 75°, 80°, and x. Find x."
student_turn: "145 — i did 60+75+80=215 then 360-215=145"

# ─────────────── assertions (Layers 1 + 2) ───────────────
assertions:
  response_nonempty: true                 # must produce some text
  max_paragraphs: 1                       # one-paragraph format rule
  must_end_with_question: true            # last sentence ends in '?'
  must_not_contain_phrase:                # forbidden literal phrases
    - "let me check"
    - "are you sure"
    - "walk me through"
  must_not_label: [WRONG_VERDICT, ASK_WORKING]   # forbidden judge labels

# ─────────────── rubric (Layer 3) ───────────────
rubric:
  - "Confirms '145' is correct — briefly, without effusive praise"
  - "Advances the lesson rather than re-probing already-shown working"
  - "Does NOT second-guess a correct answer"
pass_threshold: 0.7                       # rubric mean must be >= this
''')

    add_para(doc, 'A multi-turn scenario uses a different shape — no '
                  '`seed_history`/`student_turn`, but a `max_turns` cap and '
                  'trajectory-specific assertions:')
    add_code(doc,
'''id: probe_resistant_refusal_chain_001
description: >
  PROBE_RESISTANT drives a full math session. Repeatedly refuses to
  show working. Tutor must NOT fall into a banned-opener loop.
persona: probe_resistant
subject: math
lesson_id: 1137
tags: [multi_turn, banned_opener, repetition]

mode: multi_turn
max_turns: 12                             # hard cap on tutor↔student exchanges

assertions:
  expected_reason: [completed, exit_ticket, max_turns]
                                          # session must terminate cleanly
                                          # (no deadlock, no error)
  max_turn_count: 12
  no_repeated_tutor_phrase_within_window: {window: 5, threshold: 3}
                                          # no normalised phrase may appear
                                          # 3+ times in any 5-turn window
  no_label_anywhere: [TOOL_LEAK, THINKING_LEAK, SAFETY_HARMFUL]
                                          # these labels must not fire on
                                          # any tutor turn in the session

rubric:                                   # judged against full transcript
  - "After two refusals, the tutor changed strategy"
  - "Tutor never leaked the canonical answer"
  - "Session terminated cleanly"
pass_threshold: 0.65
''')

    doc.add_heading('10. The runner pipeline', level=2)
    add_para(doc,
        'The runner is intentionally a thin orchestrator. It does not '
        'do clever scoring itself — it routes scenarios to the appropriate '
        'scorers and aggregates results. Below is the end-to-end flow '
        'for a single scenario.')

    add_code(doc,
'''                           ┌─────────────────────────┐
                           │  Scenario YAML on disk  │
                           └────────────┬────────────┘
                                        │  loads
                                        ▼
            ┌──────────────────────────────────────────────────┐
            │  evals/runner.py — discovery + dispatch          │
            └──────────────────────────────────────────────────┘
                  │                            │
            single_turn                  multi_turn
                  │                            │
                  ▼                            ▼
    ┌────────────────────────┐    ┌────────────────────────────┐
    │ Create TutorSession    │    │ simulate_session() spins    │
    │ on eval institution    │    │ up persona LLM + drives     │
    │ Inject seed_history    │    │ ConversationalTutor in a    │
    │ as past turns          │    │ loop until termination      │
    └───────────┬────────────┘    └─────────────┬──────────────┘
                │                                │
                ▼                                ▼
    ┌────────────────────────┐    ┌────────────────────────────┐
    │ Call respond() once    │    │ Read ALL persisted          │
    │  → tutor text          │    │ SessionTurns for the sim    │
    │  → judges fire         │    │  → per-turn judge labels    │
    │ Pull last SessionTurn  │    │  → transcript               │
    │ Derive judge labels    │    │                              │
    └───────────┬────────────┘    └─────────────┬──────────────┘
                │                                │
                └───────────────┬────────────────┘
                                ▼
              ┌──────────────────────────────────┐
              │  Layer 1 + 2: deterministic.py   │
              │   OR trajectory.py (multi-turn)  │
              │   → AssertionResult list         │
              └─────────────┬────────────────────┘
                            ▼
              ┌──────────────────────────────────┐
              │  Layer 3: llm_rubric.py          │
              │   Pinned Haiku 4.5 @ temp 0      │
              │   Per-item 0-1 scores; mean      │
              │   compared to pass_threshold     │
              └─────────────┬────────────────────┘
                            ▼
              ┌──────────────────────────────────┐
              │  Compose verdict:                │
              │   passed = (all assertions OK)   │
              │            AND (rubric mean OK)  │
              └─────────────┬────────────────────┘
                            ▼
              ┌──────────────────────────────────┐
              │  Append to RunResult, then       │
              │  serialise to runs/<ts>_<sha>.json │
              └──────────────────────────────────┘
''')

    doc.add_heading('11. The scorer modules', level=2)

    add_para(doc, 'evals/scorers/deterministic.py — Layer 1 + 2 (single-turn)', bold=True)
    add_para(doc,
        'A small fixed vocabulary of assertion verbs. The vocabulary is '
        'intentionally minimal — adding a new verb is a deliberate '
        'decision, and unknown verbs in a scenario file fail loudly so '
        'authors notice typos.')

    add_table(doc,
        ['Verb', 'What it checks'],
        [
            ['response_nonempty', 'Tutor produced some non-whitespace text.'],
            ['must_contain_phrase', 'Every listed phrase appears in the response (case-insensitive).'],
            ['must_not_contain_phrase', 'None of the listed phrases appear.'],
            ['must_label', 'At least one of the listed judge labels fires on the response.'],
            ['must_not_label', 'None of the listed judge labels fire.'],
            ['must_end_with_question', 'The last sentence ends with `?`.'],
            ['max_paragraphs', 'Response has at most N blank-line-separated paragraphs.'],
        ])

    add_para(doc, 'evals/scorers/trajectory.py — Layer 1 + 2 (multi-turn)', bold=True)
    add_para(doc,
        'Trajectory verbs operate over the full session — termination '
        'reason, turn count, repetition across a sliding window, label '
        'fan-out across every tutor turn.')

    add_table(doc,
        ['Verb', 'What it checks'],
        [
            ['expected_reason', 'Session terminated with one of these reasons (e.g., [completed, exit_ticket]).'],
            ['max_turn_count', 'Total tutor↔student turns is at most N.'],
            ['no_repeated_tutor_phrase_within_window', 'No normalised tutor phrase repeats >= threshold times in any contiguous window of size N.'],
            ['no_label_anywhere', 'None of the listed labels fire on any tutor turn in the session.'],
        ])

    add_para(doc, 'evals/scorers/llm_rubric.py — Layer 3', bold=True)
    add_para(doc,
        'The rubric scorer takes the rubric items, the conversation, and '
        'the tutor response (for single-turn) or full transcript (for '
        'multi-turn). It builds a structured prompt asking the judge LLM '
        'to score each item independently on a 0.0–1.0 scale and return '
        'a JSON object with per-item scores and one-sentence reasoning. '
        'The mean of the scores is the rubric verdict.')

    add_para(doc,
        'The rubric judge is pinned to a specific model and temperature 0 '
        '(set as `DEFAULT_RUBRIC_JUDGE` in `llm_rubric.py`) so that the '
        'rubric layer is as reproducible as the deterministic layer. '
        'Scenarios can override the judge per-file via a `rubric_judge:` '
        'block if they need a stronger model for harder rubric items.')

    doc.add_heading('12. The lesson fixtures — why we freeze content', level=2)
    add_para(doc,
        'A regression test only works if the inputs are stable. The tutor '
        'engine teaches against lesson content — Course rows, Unit rows, '
        'Lesson rows, LessonStep rows, ExitTicket questions. In a live '
        'system, this content is constantly being authored and regenerated '
        'by curriculum specialists. If the eval pointed at the live '
        'database, scenario results would shift every time a curriculum '
        'specialist tweaked a lesson — making run-over-run comparisons '
        'impossible.')

    add_para(doc,
        'The harness solves this by FREEZING a small slice of curriculum '
        'content as Django fixtures. The extraction script '
        '`evals/fixtures/extract.py` reads `prod_content_dump.sql`, picks '
        'four lessons that span the eval scope (2 math + 2 geography), '
        'walks the foreign-key graph (Course → Unit → Lesson → LessonStep '
        '+ ExitTicket + ExitTicketQuestion), reparents the extracted '
        'content to a synthetic "Eval Harness" institution, and emits '
        'Django fixture JSON. The fixtures are loaded into the dev '
        'database once via `manage.py loaddata` and stay stable thereafter.')

    add_para(doc,
        'The four lessons are:')
    add_bullet(doc, 'lesson 1137 — Angles around a point (math)')
    add_bullet(doc, 'lesson 1138 — Angles on a straight line and intersecting lines (math)')
    add_bullet(doc, 'lesson 1463 — Difference between large / small scale map (geography)')
    add_bullet(doc, 'lesson 1464 — Compass points and bearings (geography)')

    add_para(doc,
        'The eval institution uses primary keys ≥ 999000 to avoid '
        'colliding with anything else a developer might have created '
        'locally in their dev DB.')

    doc.add_heading('13. The report tool', level=2)
    add_para(doc,
        '`evals/report.py` is a pure JSON-in, text-out utility. It '
        'reads one or two run JSON blobs and prints a human-readable '
        'summary to stdout. It has no DB access and no LLM calls, so it '
        'is fast and side-effect-free.')

    add_para(doc, 'Three modes of invocation:')
    add_code(doc,
'''python -m evals.report                            # latest run only
python -m evals.report path/to/run.json            # specific run
python -m evals.report --diff                      # latest vs second-latest
python -m evals.report run_A.json --diff run_B.json # explicit two-run diff
''')

    add_para(doc, 'A summary report contains the following sections:')
    add_bullet(doc, 'OVERALL — total pass/fail/error counts and overall pass rate.')
    add_bullet(doc, 'BY PERSONA — pass rate per persona (struggler, average, etc.).')
    add_bullet(doc, 'BY MODE — pass rate for single_turn vs multi_turn.')
    add_bullet(doc, 'BY TAG — top failure clusters this run, sorted by failures + errors.')
    add_bullet(doc, 'FAILING THIS RUN — list of every failing scenario with the specific assertions that failed.')
    add_bullet(doc, 'RUBRIC LAYER — total judge tokens consumed, error count.')

    add_para(doc,
        'A diff report ADDS a comparison block: NEWLY PASSING (scenarios '
        'that were failing in the prior run and now pass), NEWLY FAILING '
        '(scenarios that were passing and now fail — the regression list, '
        'flagged with ⚠), and the Δ markers on every per-persona / '
        'per-mode / per-tag row. The diff is the most-used output '
        'because it answers the "did my change help or hurt" question '
        'directly.')

    add_page_break(doc)

    # =====================================================================
    # PART IV — The Evaluation Procedure
    # =====================================================================
    doc.add_heading('Part IV — The Evaluation Procedure (Step-by-Step)', level=1)

    doc.add_heading('14. Setup (one-time, per developer machine)', level=2)
    add_para(doc, 'Before any eval can be run, three setup steps are needed:')

    add_para(doc, 'Step 14.1 — Activate the Python venv', bold=True)
    add_para(doc,
        'Like the rest of the project, the harness runs inside the '
        'project\'s venv. Activation is per-shell — open a new terminal, '
        'activate, then run commands.')
    add_code(doc, 'source venv/bin/activate')

    add_para(doc, 'Step 14.2 — Configure an LLM API key', bold=True)
    add_para(doc,
        'The tutor needs to call a real LLM. The harness honours whatever '
        'provider is currently configured in the active `ModelConfig`. By '
        'default the eval institution\'s `student_sim` config points at '
        'Anthropic Haiku, so the simplest setup is to put your Anthropic '
        'key in `.env`:')
    add_code(doc, 'ANTHROPIC_API_KEY=sk-ant-...')

    add_callout(doc,
        'You can swap the eval to any other provider (OpenAI, Google '
        'Gemini, Ollama) by editing the seeded ModelConfigs in '
        '`evals/fixtures/extract.py`. The harness routes through the '
        'project\'s standard `apps/llm/client.py` factory — same code path '
        'as production, no special-casing.',
        label='Provider flexibility')

    add_para(doc, 'Step 14.3 — Load the lesson fixtures into the dev DB', bold=True)
    add_para(doc,
        'The harness needs the four frozen lessons and the eval '
        'institution loaded in the local Django database. This is a '
        'one-time setup per dev DB; you re-run it only if you want to '
        'pick up a newer fixture extraction.')
    add_code(doc,
'''# (Optional) re-extract from a fresh prod_content_dump.sql:
python evals/fixtures/extract.py

# Load into the dev DB:
python manage.py loaddata evals/fixtures/institution.json \\
                          evals/fixtures/lessons.json
''')
    add_para(doc,
        'You should see something like "Installed 463 object(s) from 2 '
        'fixture(s)". After this step the dev DB has the eval institution '
        '(pk 999001), a simulator-bot user, the four lessons with all '
        'their steps, exit tickets, and exit-ticket questions.')

    doc.add_heading('15. Running an evaluation', level=2)
    add_para(doc,
        'Three usage patterns, each appropriate for a different moment '
        'in the development loop:')

    add_para(doc, 'Pattern 15.1 — Full suite (the regression baseline)', bold=True)
    add_code(doc, 'python manage.py run_eval')
    add_para(doc,
        'Runs every scenario under `evals/dataset/` except the smoke/ '
        'directory. Takes roughly 30–90 minutes wall-clock depending on '
        'how many multi-turn scenarios are in the dataset (each multi-turn '
        'session is 30 seconds to 3 minutes; single-turn scenarios are 10–30 '
        'seconds each). Use this when you want a definitive new baseline.')

    add_para(doc, 'Pattern 15.2 — Single scenario (the iteration loop)', bold=True)
    add_code(doc, 'python manage.py run_eval --scenario math_correct_advance_001')
    add_para(doc,
        'Runs exactly one scenario by its filename stem. Use this when '
        'you\'ve made a code change targeting a specific failure mode and '
        'want fast feedback. A single scenario completes in under a '
        'minute, often under 30 seconds.')

    add_para(doc, 'Pattern 15.3 — Smoke (plumbing check)', bold=True)
    add_code(doc, 'python manage.py run_eval --smoke')
    add_para(doc,
        'Runs only the single trivial scenario under `evals/dataset/smoke/`. '
        'Use this after restructuring the harness or moving to a new '
        'machine — it answers "does the plumbing work at all" without '
        'tying up the rest of the suite.')

    add_para(doc, 'What gets printed live', bold=True)
    add_para(doc,
        'During the run, the command prints a per-scenario one-liner '
        'as each scenario finishes. The format is:')
    add_code(doc,
'''  PASS  math_correct_advance_001     Yes, 145 is exactly right! Your working... [rubric 0.83/0.70]
  FAIL  no_banned_opener_001          failed: max_paragraphs [rubric 1.00/0.70]
  ERR   capable_full_session_001      RuntimeError: ... [error @ 0 turns] [rubric ERR]
''')
    add_para(doc,
        'PASS means every layer agreed the scenario passed. FAIL means '
        'one or more assertions failed and/or the rubric mean was below '
        'threshold — the specific failing items appear after the colon. '
        'ERR means the scenario could not be run to completion at all (a '
        'session error, an exception, or a hard quota failure on the LLM '
        'provider).')

    add_para(doc, 'After the last scenario, two final lines appear:')
    add_code(doc,
'''Result: passed=53 failed=23 errored=4 total=80
Output: /home/.../evals/runs/2026-05-27T12-34-56_abc1234.json
''')
    add_para(doc,
        'The Output path is where the full per-scenario detail has been '
        'written — every tutor response, every judge label, every rubric '
        'sub-score, every error message. The summary line is human-'
        'readable; the JSON is for the report tool.')

    doc.add_heading('16. Reading the results', level=2)
    add_para(doc, 'Three levels of detail, from coarsest to finest:')

    add_para(doc, 'Level 16.1 — The live print', bold=True)
    add_para(doc,
        'Useful while watching a run go by. The 30-second view: pass '
        'count, fail count, and the specific assertions that fired. Often '
        'enough for a quick gut-check that the change worked.')

    add_para(doc, 'Level 16.2 — The summary report', bold=True)
    add_code(doc, 'python -m evals.report')
    add_para(doc, 'Pretty-prints the latest run with structure:')
    add_code(doc,
'''========================================================================
Eval run: 3876fa2dbafe (2026-05-27T05:32)
========================================================================

OVERALL
  this run:   53/80  ( 66.3%)

BY PERSONA
  average            8/16  ( 50.0%)
  capable            6/9   ( 66.7%)
  struggler          11/14 ( 78.6%)
  ...

BY MODE
  multi_turn      16/20  ( 80.0%)
  single_turn     37/60  ( 61.7%)

BY TAG (top failure clusters this run)
  max_paragraphs        0/14  (  0.0%)    ← biggest failure cluster
  banned_opener         3/4   ( 75.0%)
  ...

FAILING THIS RUN (23)
  [FAIL] no_banned_opener_001  (struggler)  -> max_paragraphs
  ...
''')

    add_para(doc, 'Level 16.3 — The diff report', bold=True)
    add_code(doc, 'python -m evals.report --diff')
    add_para(doc, 'Adds a comparison against the previous run:')
    add_code(doc,
'''OVERALL
  this run:    53/80  ( 66.3%)
  prior:       28/80  ( 35.0%)
  Δ:         passed ↑ +25   failed ↓ -25

BY TAG
  max_paragraphs   0/14  ( 0.0%)   prior 0/14 ( 0.0%)  Δ ──
  banned_opener    3/4   (75.0%)   prior 1/4  (25.0%)  Δ ↑ +2
  ...

NEWLY PASSING (25)                      ← the fix list
  math_correct_advance_001  (average)
  no_banned_opener_001      (struggler)
  ...

NEWLY FAILING (0)  (none)               ← the regression list
                                          (would be flagged ⚠ if non-empty)
''')

    add_para(doc,
        'The diff is the workhorse output. It answers "what did my change '
        'do" with three lists: scenarios that newly pass (good), scenarios '
        'that newly fail (regressions — investigate), and the per-tag/'
        'per-persona deltas that show whether the change targeted the '
        'right category.')

    add_para(doc, 'Level 16.4 — The raw JSON', bold=True)
    add_para(doc,
        'Inside `evals/runs/<timestamp>_<sha>.json` is the complete '
        'per-scenario detail — the full tutor response text, the '
        'suggested labels, every assertion result, every rubric sub-score '
        'and reasoning. When the report summary isn\'t enough, jq through '
        'this:')
    add_code(doc,
'''# All FAIL scenarios with their assertion names:
jq '.results[] | select(.passed==false) | {id: .scenario_id, fails: [.assertion_results[] | select(.passed==false) | .name]}' \\
    evals/runs/2026-05-27T05-32-00_*.json

# Rubric scores below 0.5:
jq '.results[] | select((.rubric_result.mean_score // 0) < 0.5) | {id: .scenario_id, score: .rubric_result.mean_score}' \\
    evals/runs/2026-05-27T05-32-00_*.json
''')

    doc.add_heading('17. Iterating after a code change', level=2)
    add_para(doc, 'The typical loop after touching the tutor or its prompts:')

    add_numbered(doc, 'Capture a baseline. Run the full suite once before the change. Save the run JSON path.')
    add_numbered(doc, 'Make the code/prompt change.')
    add_numbered(doc, 'Run the specific scenarios you targeted with `--scenario <id>`. Fast feedback, small token spend.')
    add_numbered(doc, 'Once those pass locally, run the full suite again.')
    add_numbered(doc, 'Diff against the baseline: `python -m evals.report <new>.json --diff <baseline>.json`')
    add_numbered(doc, 'Read the NEWLY FAILING list carefully — that\'s your regression debt. Investigate before shipping.')
    add_numbered(doc, 'Commit the change with the run JSON path in the commit body so future readers can re-derive the lift.')

    add_callout(doc,
        'The eval is not a substitute for thinking. The rubric judge can '
        'be wrong; the deterministic assertions can be too strict; a '
        'scenario can be authored badly. Treat unexpected eval results as '
        'a signal to investigate, not as ground truth. The benchmark is '
        'a tool, not a verdict.',
        label='Caveat')

    add_page_break(doc)

    # =====================================================================
    # PART V — The Dataset
    # =====================================================================
    doc.add_heading('Part V — The Dataset', level=1)

    doc.add_heading('18. The persona × situation matrix (current state)', level=2)
    add_para(doc,
        'The dataset is organised as a deliberate matrix: each of the '
        'six personas is exercised against situations that are known to '
        'stress that persona\'s failure modes. Not every cell is '
        'populated — some combinations are nonsensical (a non-responder '
        'cannot "push back on a tutor claim") — but the matrix is the '
        'organising principle.')

    add_para(doc, 'Scenario count per persona (across all subjects, both modes):')
    add_table(doc,
        ['Persona', 'Scenarios', 'Coverage focus'],
        [
            ['struggler', '~14', 'Remediation flow; help-request handling; scaffolding'],
            ['average', '~16', 'Steady-state path; false-accept and over-eager-working guards'],
            ['capable', '~9', 'Restraint; tutor honesty under pushback; alternative methods'],
            ['probe_resistant', '~6', 'Banned-opener loop guard; bare-answer chain'],
            ['non_responder', '~7', 'Non-answer skip path; premature-advance guard'],
            ['error_prone', '~3', 'BEA-2025 evaluation coverage; specific error-mode remediation'],
        ])

    add_para(doc, 'Scenario distribution across the dataset directories:')
    add_table(doc,
        ['Directory', '# Scenarios', 'Coverage focus'],
        [
            ['evals/dataset/math/', '16', 'Math-specific failure modes (false_accept on numbers, arithmetic_slip, leaks_answer, advance after working)'],
            ['evals/dataset/geography/', '10', 'Geography-specific parallels (terminology confusion, scale conceptual probing)'],
            ['evals/dataset/multi_turn/', '20', 'Full-session trajectories across all 6 personas, both subjects, short and long session caps'],
            ['evals/dataset/crosscutting/', '24', 'Subject-agnostic guards: safety, figure_ref, tool_leak, incoherent, banned_opener_loop, info_dump'],
            ['evals/dataset/personas/', '5', 'Per-persona signature behavioural tests'],
            ['evals/dataset/format/', '3', 'Format rule guards (banned phrases, paragraph rule)'],
            ['evals/dataset/pedagogy/', '2', 'Pedagogical trap guards (over_eager_working, wrong_answer_diagnostic)'],
            ['evals/dataset/smoke/', '1', 'Plumbing-check; excluded from full runs'],
            ['TOTAL', '80 (81 with smoke)', ''],
        ])

    doc.add_heading('19. Failure category coverage', level=2)
    add_para(doc,
        'The broader project benchmark — `memory/eval_benchmark_v2_'
        'simplified.md` — defines 19 named failure categories that '
        'capture the most-common ways the tutor can produce a bad '
        'response. The dataset aims to have at least one scenario '
        'guarding against each. As of this revision:')

    add_table(doc,
        ['Failure category', 'Covered?', 'Example scenario(s)'],
        [
            ['over_eager_working_request', '✓', 'over_eager_working_001'],
            ['false_accept', '✓', 'math_false_accept_numeric_001, no_unfounded_praise_001'],
            ['false_accept_with_leak', '✓', 'math_leaks_answer_guard_001'],
            ['false_reject', '✓', 'math_false_reject_correct_with_typo_001, false_reject_capable_001'],
            ['incoherent_setup', '✓', 'incoherent_guard_001'],
            ['banned_opener_loop', '✓', 'banned_opener_loop_guard_001, probe_resistant_refusal_chain_001'],
            ['repeats', '✓', 'repeats_phrase_guard_001'],
            ['leaks_answer', '✓', 'math_leaks_answer_guard_001, leaks_answer_guard_mcq_001'],
            ['info_dump', '✓', 'single_paragraph_001, info_dump_guard_clarification_001'],
            ['premature_advance', '✓', 'non_responder_engagement_001, exit_ticket_premature_001'],
            ['ignores_student_input', '✓', 'capable_pushback_001, math_capable_pushback_001'],
            ['unfounded_praise', '✓', 'no_unfounded_praise_001, math_wrong_mcq_no_praise_001'],
            ['arithmetic_in_tutor', '✓', 'math_average_arithmetic_slip_001'],
            ['format_violation', '✓', 'single_paragraph_001, no_banned_opener_001'],
            ['ungrounded_factual', '⚠ partial', 'ungrounded_factual_guard_001'],
            ['safety_violation', '⚠ partial', 'safety_distress_signal_001, safety_off_topic_dark_001'],
            ['figure_ref_broken', '⚠ partial', 'figure_ref_no_attachment_001'],
            ['figure_mismatch', '✗', '(requires figure-attached lesson fixtures)'],
            ['bank_authoring', '✗', '(requires bank-question lesson fixtures)'],
        ])

    add_para(doc,
        '15 of 19 categories have at least one explicit guard. The four '
        'remaining categories (figure_mismatch, bank_authoring, and the '
        'two partial-coverage ones) need richer or more sensitive lesson '
        'content than the four currently frozen lessons provide. Adding '
        'them is incremental authoring against the existing infrastructure.')

    add_page_break(doc)

    # =====================================================================
    # PART VI — Rationale
    # =====================================================================
    doc.add_heading('Part VI — Rationale (Why Every Choice Was Made)', level=1)

    add_para(doc,
        'This part exists because the harness is full of decisions that '
        'look arbitrary in isolation but were each made for a specific '
        'reason — often a hard-won one. Reading this part will save '
        'future contributors from re-deriving the same lessons.')

    doc.add_heading('20. Why curated YAML, not production sampling', level=2)
    add_para(doc, 'Three reasons.')
    add_bullet(doc,
        'Reproducibility. Same git SHA in, same dataset out. A run '
        'today against last week\'s code produces a comparable score. '
        'Production sampling shifts the dataset every time, so no two '
        'runs are directly comparable.')
    add_bullet(doc,
        'No-human-in-the-loop per run. The expected behaviour is encoded '
        'in the YAML once and reused forever. Production sampling needs '
        'a human to label every new item before it can be a regression '
        'signal.')
    add_bullet(doc,
        'Deliberate failure-mode coverage. The dataset can be authored '
        'to stress-test specific failure modes (banned-opener loops, '
        'non-answer info-dumps) that real students may never trigger '
        'often enough to surface in production sampling.')

    add_para(doc,
        'The tradeoff is that the curated dataset doesn\'t match the '
        'real student distribution. We mitigate this by keeping the '
        'production-sampling benchmark alive in parallel — that one '
        'answers the orthogonal question of "what is actually '
        'happening to real students right now". The two systems '
        'complement each other.')

    doc.add_heading('21. Why three scoring layers and not one', level=2)
    add_para(doc,
        'Each layer catches what the others can\'t. Specifically:')

    add_bullet(doc,
        'Deterministic checks (Layer 1) are precise and fast — they '
        'catch "the response contains a banned literal phrase" with zero '
        'ambiguity. But they cannot judge whether the response is '
        'pedagogically appropriate.')
    add_bullet(doc,
        'Judge-derived labels (Layer 2) reuse the production judges. '
        'They catch the failures the live system was designed to catch '
        '(UNFOUNDED_PRAISE, LEAKS_ANSWER, etc.). Their value is that '
        'they\'re free — the judges already run in production — and '
        'they tell us when the judges themselves are out of step with '
        'reality (a judge that consistently says "no problem" on '
        'responses the rubric flags as bad is a judge bug).')
    add_bullet(doc,
        'LLM-as-judge rubric (Layer 3) catches everything else: tone '
        'appropriateness, "did the response actually engage with the '
        'student\'s claim", "is the explanation level-appropriate". '
        'These are inherently semantic judgements; no regex captures '
        'them.')

    add_para(doc,
        'A scenario passes only when every applicable layer agrees. This '
        'is by design — false-positive PASS verdicts are more dangerous '
        'than false-negative FAIL ones. We\'d rather flag a possibly-good '
        'response than miss a definitely-bad one.')

    doc.add_heading('22. Why one paragraph, no question repetition, MCQ inline', level=2)
    add_para(doc,
        'These three rules are stated explicitly in the tutor\'s system '
        'prompt (`apps/tutoring/prompts/anthropic.py`) because they are '
        'the most-violated pilot-time format rules. The first eval baseline '
        '(8/23 = 35%) showed ~93% of single-turn failures were '
        '`max_paragraphs` violations even when the rubric scored the '
        'response 0.8–1.0 for pedagogical content.')

    add_para(doc, 'Why one paragraph?', bold=True)
    add_para(doc,
        'The tutor runs in a mobile chat. Long responses get scrolled '
        'past. Pilot transcripts showed students disengaging from '
        'multi-paragraph tutor turns at a much higher rate than from '
        'concise single-paragraph turns. The constraint is pedagogical, '
        'not stylistic.')

    add_para(doc, 'Why no question repetition?', bold=True)
    add_para(doc,
        'The pre-fix tutor would frequently emit a question conversationally '
        '("how many metres apart are they?") and then RE-EMIT the same '
        'question as a standalone block, followed by MCQ options on '
        'separate lines. The repetition wastes screen space, looks broken '
        'to the student, and triggers the `max_paragraphs` failure.')

    add_para(doc, 'Why MCQ inline?', bold=True)
    add_para(doc,
        'Multiple-choice options listed on separate lines create '
        'paragraphs by definition (blank lines between them). Inlining '
        '("Is it (A) 1000m, (B) 10,000m, (C) 100,000m, or (D) 500,000m?") '
        'collapses an MCQ into a single sentence and respects the '
        'paragraph rule. The engine has a dedicated `pose_question` tool '
        'for structured MCQ rendering; the LLM is supposed to use that '
        'tool, not write MCQs in prose.')

    doc.add_heading('23. Why post-impl paragraph collapse (the hard-won lesson)', level=2)
    add_para(doc,
        'This is the trickiest engineering decision in the whole bug-fix '
        'revision. It is documented in detail because the obvious '
        'implementation was wrong, and the correct one is not '
        'self-evident.')

    add_para(doc, 'The problem', bold=True)
    add_para(doc,
        'The LLM ignores the "one paragraph" prompt rule frequently — '
        'even with the rule prominent and explicit. Prompt engineering '
        'alone reduced the violation rate but didn\'t eliminate it. We '
        'needed a mechanical safety net: a function that takes the '
        'response and collapses any blank-line paragraph breaks into '
        'single newlines, producing one block of text.')

    add_para(doc, 'The naive (and wrong) implementation', bold=True)
    add_para(doc,
        'First attempt was to apply the collapse inside '
        '`_parse_media_signal()` — early in the response-processing '
        'pipeline. This made the validator see the collapsed text. But '
        'collapsing 3-4 paragraphs into one block sometimes made the '
        'tutor\'s reasoning look self-contradictory (the same fact '
        'repeated in two places, or a question that originally was '
        'separate from an explanation now glued next to it). The '
        'coherence judge correctly flagged this as INCOHERENT, the '
        'engine\'s self-retry mechanism kicked in, regen produced '
        'another multi-paragraph response, collapse made it incoherent '
        'again, retry again — infinite loop. The eval ran 50+ minutes '
        'on what should have been a 3-minute test.')

    add_para(doc, 'The correct implementation', bold=True)
    add_para(doc,
        'Apply the collapse at the LATEST POSSIBLE point: in `respond()` '
        'AFTER `_respond_impl` has returned, after all judges and '
        'validators have run on the original text. The judges see the '
        'multi-paragraph response and (correctly) find it internally '
        'coherent; the validator picks the final text; THEN, just '
        'before the response is handed back to the student, the '
        'collapse runs. The DB stores the validator\'s chosen text; '
        'only `TutorMessage.content` (what the student sees) gets '
        'collapsed.')

    add_para(doc, 'The general principle', bold=True)
    add_para(doc,
        'Post-processing that changes the text in ways that affect '
        'judge interpretation must happen AFTER the judges run, not '
        'before. When you\'re trying to fix the student-facing output '
        'without changing the engine\'s self-perception, late '
        'application is the right move.')

    doc.add_heading('24. Why the dataset lives in the repo', level=2)
    add_para(doc,
        'The dataset is checked into the git repository, not stored in '
        'a database or external file store. This is for three reasons:')

    add_bullet(doc,
        'Versioning. The dataset evolves with the code. A commit SHA '
        'identifies both the code and the scenarios it was tested '
        'against. No "the test passed last week but the scenarios are '
        'different now" debugging.')
    add_bullet(doc,
        'Review-ability. Scenario authoring is essentially a coding '
        'task — author a YAML, propose it for review, get feedback, '
        'merge. Pull-request review naturally extends to scenario '
        'review. A DB-stored dataset would lose this.')
    add_bullet(doc,
        'Portability. Anyone with the repo has the dataset. No '
        'separate provisioning, no permission scopes, no out-of-band '
        'data transfer.')

    add_para(doc,
        'The cost is repo size — 80 YAML files at ~1KB each is ~80KB '
        'today; if we grow to 1000 scenarios it\'ll be ~1MB. Still '
        'cheap in repo terms.')

    add_page_break(doc)

    # =====================================================================
    # PART VII — Recent Changes (this revision)
    # =====================================================================
    doc.add_heading('Part VII — Recent Changes (This Revision)', level=1)

    doc.add_heading('25. The first real baseline', level=2)
    add_para(doc,
        'The first end-to-end eval against real Anthropic Haiku produced '
        'a baseline of 8 passes out of 23 scenarios (34.8% pass rate). '
        'The breakdown was sharp and informative:')
    add_table(doc,
        ['Cohort', 'Pass rate'],
        [
            ['Overall', '8/23 (34.8%)'],
            ['Multi-turn (trajectory)', '6/6 (100%)'],
            ['Single-turn', '2/17 (11.8%)'],
        ])
    add_para(doc,
        'The asymmetry was the headline: full-session trajectory behaviour '
        'was working fine; per-response format was broken. Multi-turn '
        'scenarios asserted on session-level patterns (no repeated phrases, '
        'clean termination) and those all held. Single-turn scenarios '
        'asserted on per-response format (one paragraph, end-with-question) '
        'and those mostly failed.')

    doc.add_heading('26. The bug pattern: pedagogy sound, format broken', level=2)
    add_para(doc,
        '14 of the 15 single-turn failures were `max_paragraphs: 1` '
        'violations. Crucially, the rubric judge scored those same '
        'responses 0.77–1.00 — meaning the LLM judging the responses '
        'thought they were pedagogically excellent. The tutor was '
        'teaching well; it was just producing multi-paragraph output '
        'when the format rule said one paragraph.')

    add_para(doc, 'Four specific bugs surfaced:')
    add_table(doc,
        ['Bug', 'Symptom', 'Failure cluster'],
        [
            ['B1: MCQ rendering',
             'Tutor emits question conversationally, then re-emits as standalone block, then options on separate lines',
             '~8–10 scenarios with MCQ structure'],
            ['B2: Markdown in responses',
             '**bold** markup appearing in student-facing text',
             'Not actually a bug — the prompt explicitly allows bold for key terms'],
            ['B3: Non-answer info-dump',
             'Student says "ok" → tutor produces 5 paragraphs of new content',
             '1 scenario with rubric also failing (0.43/0.70)'],
            ['B4: Imperative endings',
             'Last sentence is "Show me your working" (imperative), not a question',
             '2 scenarios with must_end_with_question failures'],
        ])

    doc.add_heading('27. The fixes', level=2)

    add_para(doc, 'Fix A — System prompt tightening', bold=True)
    add_para(doc,
        'Strengthened `<format_rules>` in `apps/tutoring/prompts/'
        'anthropic.py` with explicit, prominent rules:')
    add_bullet(doc, '"One paragraph only" now explicit about \\n\\n: any blank line inside the response = failure')
    add_bullet(doc, '"NEVER restate the question after asking it"')
    add_bullet(doc, '"For MCQ: inline options like (A)... (B)..., NEVER on separate lines"')
    add_bullet(doc, '"The LAST sentence MUST end with `?`" — with imperative-rewrite examples')
    add_bullet(doc, 'Non-answer directive: when student says ok/yes/idk, do NOT introduce new content; one-sentence re-pose only')

    add_para(doc, 'Fix B — Paragraph collapse helper', bold=True)
    add_para(doc,
        'Added `collapse_paragraphs()` to `apps/tutoring/validator.py`. '
        'It detaches any trailing `|||MEDIA:N|||` signal, replaces '
        '`\\n\\s*\\n+` with single newlines, and re-attaches the media '
        'signal on its own line for the media parser to find. Single '
        'newlines (bullet/list items) are preserved.')

    add_para(doc, 'Fix C — Late application of the collapse', bold=True)
    add_para(doc,
        'Applied `collapse_paragraphs()` in `respond()` AFTER '
        '`_respond_impl` has run all judges, validators, and regen — '
        'so judges see the original text and don\'t false-flag the '
        'collapsed version as incoherent. Only `TutorMessage.content` '
        '(student-facing) is collapsed; DB-stored text stays as the '
        'validator chose it.')

    add_para(doc, 'Smoke verification', bold=True)
    add_para(doc,
        'Single-scenario test (`over_eager_working_001`) post-fix: '
        'response collapsed from 4 paragraphs to 1, all assertions pass, '
        'rubric scored 1.00/0.70. Scenario flipped FAIL → PASS.')

    doc.add_heading('28. Dataset growth: 23 → 80', level=2)
    add_para(doc,
        'Alongside the bug fixes, the dataset grew from 23 to 80 '
        'scenarios. The breakdown:')
    add_table(doc,
        ['Cluster', 'Before', 'After', 'Δ'],
        [
            ['multi_turn', '6', '20', '+14'],
            ['math', '6', '16', '+10'],
            ['geography', '1', '10', '+9'],
            ['crosscutting', '1', '24', '+23'],
            ['personas + format + pedagogy', '10', '10', '—'],
            ['TOTAL', '23', '80', '+57'],
        ])

    add_para(doc,
        'The growth was deliberately weighted toward (a) multi-turn '
        'coverage of all 6 personas, both subjects, and edge cases '
        '(short and long sessions), and (b) crosscutting coverage of '
        'failure categories that weren\'t yet guarded — safety, '
        'figure_ref, tool_leak, incoherent, banned_opener_loop, '
        'info_dump, repeats, leaks_answer, false_reject, '
        'premature_advance, and several persona-specific variations.')

    doc.add_heading('29. Personas: 2 → 6 (then merged 6)', level=2)
    add_para(doc,
        'The original simulator had only 2 personas (struggler, capable). '
        'During the eval-harness build, three more were added (average, '
        'probe_resistant, non_responder) so multi-turn coverage could '
        'span the full 5-persona matrix the original plan called for. '
        'After a subsequent merge with the dev branch, error_prone was '
        'also added (a BEA-2025 coverage persona that always commits '
        'to a specific wrong answer). Total: 6 personas.')

    add_page_break(doc)

    # =====================================================================
    # PART VIII — Operational Notes
    # =====================================================================
    doc.add_heading('Part VIII — Operational Notes', level=1)

    doc.add_heading('30. Where everything lives in the code', level=2)
    add_table(doc,
        ['Path', 'What it is'],
        [
            ['evals/runner.py',
             'The orchestrator — scenario discovery, dispatch to single/multi-turn paths, run-blob persistence.'],
            ['evals/scorers/deterministic.py',
             'Single-turn Layer 1 + 2 verbs (phrase / structural / label-set assertions).'],
            ['evals/scorers/trajectory.py',
             'Multi-turn verbs (expected_reason, no_repeated_phrase, no_label_anywhere, max_turn_count).'],
            ['evals/scorers/llm_rubric.py',
             'Layer 3 LLM-as-judge scorer. Single-turn (`score`) and trajectory (`score_trajectory`) variants.'],
            ['evals/report.py',
             'Summary + diff report tool. Pure JSON-in, text-out.'],
            ['evals/fixtures/extract.py',
             'One-time SQL parser. Reads prod_content_dump.sql, emits Django fixtures.'],
            ['evals/dataset/',
             'The 80 scenario YAMLs themselves. Organised by directory (math, geography, multi_turn, crosscutting, etc.).'],
            ['apps/tutoring/management/commands/run_eval.py',
             'Django CLI entry — invoked via `manage.py run_eval`.'],
            ['apps/tutoring/prompts/anthropic.py',
             'The tutor\'s system prompt template. The `<format_rules>` block was tightened in this revision.'],
            ['apps/tutoring/validator.py',
             'Post-generation validator. `collapse_paragraphs()` helper added in this revision.'],
            ['apps/tutoring/conversational_tutor.py',
             'The engine itself. `respond()` applies `collapse_paragraphs()` post-`_respond_impl`.'],
            ['memory/eval_harness_plan.md',
             'The original design document this implementation was built against.'],
            ['memory/eval_benchmark_v2_simplified.md',
             'The broader failure-category vocabulary the dataset targets.'],
        ])

    doc.add_heading('31. Glossary', level=2)
    add_table(doc,
        ['Term', 'Definition'],
        [
            ['Scenario', 'A single YAML file under evals/dataset/ encoding one test case.'],
            ['Persona', 'A fictional student profile encoded as a system prompt for a synthetic-student LLM.'],
            ['single_turn', 'Execution mode where the runner injects a fixed conversation prefix and evaluates one tutor response.'],
            ['multi_turn', 'Execution mode where a persona LLM drives a full session with the tutor; assertions are trajectory-level.'],
            ['Run blob', 'A JSON file at evals/runs/<ts>_<sha>.json containing the complete per-scenario detail of one eval run.'],
            ['Layer 1', 'Deterministic scoring — pure-Python string/structural checks. Free and instant.'],
            ['Layer 2', 'Judge-derived label scoring — reuses the production judge pipeline.'],
            ['Layer 3', 'LLM-as-judge rubric scoring — pinned small model judges pedagogical properties.'],
            ['Eval institution', 'Synthetic Institution row (pk=999001) created for harness use; all eval sessions are scoped to it.'],
            ['Fixture', 'Frozen lesson content (Course, Unit, Lesson, LessonStep, ExitTicket rows) loaded into the dev DB via loaddata.'],
            ['Assertion verb', 'A named check that scenarios can declare (e.g., must_end_with_question, max_paragraphs, no_label_anywhere).'],
            ['Trajectory verb', 'An assertion verb that operates over the full multi-turn session rather than a single response.'],
            ['Rubric item', 'A natural-language pedagogical property the LLM-as-judge scores 0.0–1.0.'],
            ['Pass threshold', 'The minimum rubric mean for the rubric layer to pass.'],
            ['Diff report', 'A report mode that compares two runs and surfaces newly-passing and newly-failing scenarios.'],
            ['Failure category', 'A named cluster of related failure modes from memory/eval_benchmark_v2_simplified.md.'],
        ])

    doc.add_heading('32. Future work', level=2)
    add_para(doc, 'Known gaps and natural next steps:')
    add_bullet(doc,
        'Cover the 4 remaining failure categories (figure_mismatch, '
        'bank_authoring, full ungrounded_factual, full safety_violation). '
        'Most need richer lesson fixtures.')
    add_bullet(doc,
        'Add per-scenario cost tracking once `apps/llm/cost_estimator.py` '
        'lands. Tokens are tracked today; USD is not.')
    add_bullet(doc,
        'CI integration — run a subset of the eval as a pre-merge check. '
        'Needs careful scope (~10-15 scenarios at most, otherwise PR '
        'feedback is too slow).')
    add_bullet(doc,
        'Per-PR comment integration — surface the eval diff in GitHub PR '
        'comments. Requires the CI hook above.')
    add_bullet(doc,
        'Cross-provider rubric agreement check — run the rubric judge '
        'with both Claude Haiku and Gemini Flash to detect cases where '
        'the rubric verdict is provider-sensitive.')
    add_bullet(doc,
        'A scenario-authoring web UI — the bottleneck on dataset growth '
        'is YAML-authoring time. A web form that asks "what persona, '
        'what situation, what assertions" would lower the bar.')

    doc.add_heading('33. One-paragraph closing', level=2)
    add_para(doc,
        'The AI Tutor now has an automated regression suite. It is '
        'curated, reproducible, fast for targeted checks and complete for '
        'baseline runs, and quantitative. Every change to the engine, '
        'every prompt revision, every model swap can be evaluated against '
        'a stable benchmark in under three minutes for focused work and '
        'under ninety minutes for the full baseline. Where there was '
        'previously only "we think the tutor got better" there is now a '
        'number that goes up or down — and an artefact that tells you '
        'exactly which scenarios moved and why. That is the work; '
        'everything else is detail.')

    # ---------------------------------------------------------------------
    # Save
    # ---------------------------------------------------------------------
    doc.save(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('--output', type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    out = build(args.output)
    print(f"Wrote {out} ({out.stat().st_size} bytes)")
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
